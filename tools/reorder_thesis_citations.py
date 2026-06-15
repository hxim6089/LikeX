from __future__ import annotations

import re
from pathlib import Path

from docx import Document


SRC = Path(r"D:/SHARE/OneDrive/Desktop/毕业论文 初稿7.1 检测.docx")
OUT = Path(r"D:/TheEnd/qnyproj-main/recommendation-system/论文相关/毕业论文 初稿7.1 引用重排版.docx")


UPDATED_PARAGRAPHS = {
    52: (
        "信息过载主要是指用户在有限的时间内，很难快速找到自己真正需要或感兴趣的内容，大量相关性不高的信息也会分散用户的注意力。"
        "用户在一定时间里能看的内容数量本来就有限，平台需要从大量候选内容中挑出更适合展示的，再按某种规则排好序。"
        "传统的时间流大多按发布时间倒序展示，默认越新的内容越值得被优先看到，这其实隐含了一种假设——用户对不同发布者和不同内容的关注程度差别不大。"
        "但在实际的社交网络里，人们的兴趣方向、互动频率和内容偏好往往各不相同，光看时间顺序很容易跟用户的真实需求错位；"
        "推荐排序可以结合浏览、点赞、评论、关注这些行为去分析和判断用户兴趣，这样用户刷新页面时看到的内容会更贴近自己的偏好，"
        "阅读体验会变好，平台的使用黏性也能跟着增强，这个思路与基于用户浏览行为的个性化推荐研究以及深度学习推荐系统所追求的目标基本一致[1][2]。"
    ),
    55: (
        "推荐排序技术为解决上述问题提供了一种比较有效的方式。系统可以采集用户的浏览、点赞、评论、转发、搜索、关注、点踩、快速滑过等多种行为数据，"
        "并根据这些数据分析用户的兴趣方向和互动习惯。同时，系统还可以结合内容标签、文本相似度、热门话题、社交关系和时间衰减等因素，对信息流内容进行个性化排序。"
        "从用户画像构建、显隐式反馈融合以及协同过滤推荐等研究可以看出，用户行为、内容特征和交互关系能够共同支撑推荐排序模型[3][4][5]。"
        "与单纯按照时间倒序展示相比，个性化推荐在保留内容新鲜度的基础上，也能提高内容与用户兴趣的匹配程度，使用户能够更快看到符合自身需求的信息。"
    ),
    62: (
        "工业界推荐系统已从早期协同过滤与矩阵分解，逐步演进为融合深度模型、多任务学习与海量日志训练的排序架构。"
        "YouTube 推荐系统相关研究表明，深度神经网络可以综合用户历史行为、上下文和候选内容特征，从而服务于大规模视频推荐场景[6]。"
        "X 平台开源的推荐算法也展示了候选召回、过滤与重排序等工程流水线思想，为本课题“关注源 + 全站源”的双源候选组织提供了现实参照[7]。"
        "Amazon 的物品到物品协同过滤研究则说明，在电商和内容推荐场景中，基于历史交互关系计算相似内容仍然具有较强应用价值[8]。"
    ),
    63: (
        "在国内学术语境下，信息流推荐、用户画像和算法治理同样是研究热点。已有研究从科技传播、用户浏览行为、用户画像构建和推荐去偏等角度指出，"
        "推荐系统既要提升内容匹配效率，也要兼顾用户兴趣漂移、信息茧房和结果可解释性等问题[9][1][3][10]。"
        "因此，本文在实现混合推荐排序时，不只追求点击与互动分数，还加入作者多样性、负反馈衰减、冷启动探索和评分拆分展示，使算法逻辑能够对应系统实际功能与可视化分析要求。"
    ),
    64: (
        "近年来，深度学习推荐模型和 CTR 预估模型不断强化对用户兴趣序列、交叉特征和上下文信息的建模能力，DeepFM、DIEN 等方法说明了深度模型在点击率预测与兴趣演化建模中的价值[11][12]。"
        "基于深度神经网络的个性化推荐研究和深度学习推荐系统相关著作，也为理解召回、排序、多任务建模和推荐解释提供了方法基础[13][2]。"
        "本文中的系统没有直接训练大规模深度推荐模型，而是在可解释混合打分基础上增加 AI 推荐策略，用于候选理解、排序解释和降级对比，从而在毕业设计规模的数据条件下体现智能推荐思想。"
    ),
    65: (
        "矩阵分解和深度推荐模型通常会把用户和内容转化为低维向量表示，再根据已有的历史互动数据学习用户的潜在兴趣，这类方法能够从较大规模的行为记录中发现一些人工标签不容易覆盖的兴趣关系[14][11][12]。"
        "不过，这类方法一般需要持续训练，也需要较多的数据作为支撑。相比之下，本文当前实现的系统更重视可运行、可解释和便于调节参数，因此采用了多因子启发式排序方法。"
        "系统通过协同过滤加成、文本相似度加成和动态权重调节等方式，构成了比较适合本系统数据规模的推荐排序方案。"
    ),
    66: (
        "在国内常见的工程开发环境中，Spring Boot、Vue 和 MySQL 的组合学习成本相对较低，相关资料也比较丰富，比较适合完成任务书中前后端分离信息系统的开发要求[15][16][17]。"
        "在数据库设计上，本文使用关系型数据模型来保存系统运行所需的主要数据，包括用户信息、帖子内容、互动记录、消息通知、广告数据以及推荐日志等。"
        "为了降低复杂查询对系统运行效率的影响，系统会在常用查询字段上设置索引，并按照不同业务模块对数据表进行划分，这样数据存储和后续查询就会更清晰一些；"
        "数据密集型应用系统设计相关研究也说明，数据模型、索引、缓存和可靠性设计会直接影响系统后续扩展能力[18]。"
        "在推荐算法方面，矩阵分解、物品协同过滤以及深度学习推荐系统等相关研究，为本文分析用户兴趣、处理浏览和互动等隐式反馈、判断推荐排序效果提供了参考依据[14][8][2]。"
    ),
    67: (
        "在接口设计上，系统主要按照用户、帖子、互动、关注、消息、通知、广告、用户画像和推荐策略这几个功能模块来划分 RESTful API，前端页面可以通过相对统一的接口路径访问后端服务，这样能减少接口调用上的混乱。"
        "对于实时私信和通知这类对及时性要求较高的功能，系统会通过消息通道实现近实时推送，这种设计既符合任务书中关于 Spring Boot RESTful API、Vue 单页应用以及前后端分离 B/S 架构的要求，也方便后续进行接口测试、权限控制和功能扩展。"
        "推荐策略、AI 降级和广告投放等模块采用统一接口和可替换实现时，也体现了面向对象设计模式中封装变化、面向接口编程的思想[19]，并与数据密集型系统强调的可维护性和可演进性相一致[18]。"
    ),
    69: (
        "实时信息流排序与分发系统主要应用于社交媒体、新闻资讯、内容社区和广告投放等场景。本文系统以类 X/Twitter 平台为原型，将内容生产、用户互动、实时通知、私信通信、用户画像、推荐排序和广告分发组织在同一业务闭环中，"
        "使推荐算法不再停留于离线实验，而能够直接服务于用户刷新信息流、管理员切换策略和答辩演示中的效果对比。"
        "在推荐列表已经生成后，再根据个性化目标进行重排序，也是推荐系统中常见的优化思路[20]。"
    ),
    70: (
        "在具体应用中，推荐流用于根据用户行为和内容特征生成个性化帖子列表，关注流用于保留社交关系带来的确定性内容来源，广告分发模块则利用画像标签和广告统计实现较为基础的智能投放。"
        "个性化广告推荐研究表明，广告排序通常需要综合用户画像、广告内容、点击率估计和投放策略等因素[21]。"
        "该应用形态能够较好体现任务书中“实时信息流排序与分发”“AI 辅助功能”“算法效果可视化”和“广告智能分发”等要求。"
    ),
}


REFERENCES = [
    "刘华真, 王巍, 谷壬倩, 等. 基于用户浏览行为的个性化推荐研究综述[J]. Application Research of Computers/Jisuanji Yingyong Yanjiu, 2021, 38(8).",
    "王树森, 黎崎. 深度学习推荐系统[M]. 北京: 电子工业出版社, 2020.",
    "高广尚. 用户画像构建方法研究综述[J]. 数据分析与知识发现, 2019, 3(3): 25-35.",
    "欧朝荣, 胡军. 融合显隐式反馈的协同过滤推荐模型[J]. 控制与决策, 2024, 39(3): 1048-1056.",
    "付峻宇, 朱小栋, 陈晨. 基于图卷积的双通道协同过滤推荐算法[J]. 计算机应用研究, 2023, 40(1): 129-135.",
    "Covington P, Adams J, Sargin E. Deep neural networks for YouTube recommendations[C].In: Proceedings of the 10th ACM Conference on Recommender Systems. New York: ACM, 2016: 191-198.",
    "Twitter Inc. Twitter's Recommendation Algorithm[EB/OL]. (2023-03-31) [2026-03-13]. https://github.com/twitter/the-algorithm.",
    "Linden G, Smith B, York J. Amazon.com recommendations: Item-to-item collaborative filtering[J]. IEEE Internet Computing, 2003, 7(1): 76-80.",
    "夏丽云, 徐敏赟, 丁懿楠, 等. 智能推荐算法下的科技期刊国际传播策略研究[J]. 中国科技期刊研究, 2023, 34(11): 1486-1493.",
    "刘文贤, 朱海威, 武浩. 基于流行度和质量偏好建模的去偏推荐系统[J]. 云南大学学报(自然科学版), 2025, 47(3): 523-532.",
    "Guo H, Tang R, Ye Y, et al. DeepFM: A factorization-machine based neural network for CTR prediction[J/OL]. arXiv, 2017 [2026-03-13]. https://arxiv.org/abs/1703.04247.",
    "Zhou G, Mou N, Fan Y, et al. Deep Interest Evolution Network for click-through rate prediction[C].In: Proceedings of the AAAI Conference on Artificial Intelligence. Honolulu: AAAI Press, 2019: 5941-5948.",
    "张敏军, 华庆一, 贾伟, 等. 基于深度神经网络的个性化推荐系统研究[J]. 西南大学学报 (自然科学版), 2019, 41(11): 104-109.",
    "Koren Y, Bell R, Volinsky C. Matrix factorization techniques for recommender systems[J]. Computer, 2009, 42(8): 30-37.",
    "吴昌政. 基于前后端分离技术的 web 开发框架设计[D]. 南京: 南京邮电大学, 2020.",
    "Horstmann C S. Java核心技术·卷I：开发基础[M]. 林信良, 译. 第12版. 北京: 机械工业出版社, 2022.",
    "杨海民. Vue.js 3.0企业级管理后台开发实战：基于Element Plus[M]. 北京: 电子工业出版社, 2022.",
    "KLEPPMANN M. 数据密集型应用系统设计[M]. 赵健博等译. 北京: 中国电力出版社, 2018.",
    "Freeman E, Robson E. Head First设计模式[M]. 张晓菲, 等译. 第2版. 北京: 中国电力出版社, 2022.",
    "Pei C, Zhang Y, Zhang Y, et al. Personalized re-ranking for recommendation[C]//Proceedings of the 13th ACM conference on recommender systems. 2019: 3-11.",
    "张玉洁, 董政, 孟祥武. 个性化广告推荐系统及其应用研究[J]. 计算机学报, 2021, 44(3): 531-563.",
]


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def remove_citations(text: str) -> str:
    text = re.sub(r"(?:\[\d+\])+", "", text)
    text = re.sub(r"\s+([，。；、])", r"\1", text)
    text = re.sub(r" {2,}", " ", text)
    return text


def main() -> None:
    doc = Document(SRC)
    ref_heading_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")

    for idx, text in UPDATED_PARAGRAPHS.items():
        set_paragraph_text(doc.paragraphs[idx], text)

    for idx, paragraph in enumerate(doc.paragraphs[:ref_heading_idx]):
        if idx in UPDATED_PARAGRAPHS:
            continue
        if "[" in paragraph.text and "]" in paragraph.text:
            cleaned = remove_citations(paragraph.text)
            if cleaned != paragraph.text:
                set_paragraph_text(paragraph, cleaned)

    ref_paragraph_idxs = [
        i
        for i in range(ref_heading_idx + 1, len(doc.paragraphs))
        if doc.paragraphs[i].text.strip()
    ][: len(REFERENCES)]

    if len(ref_paragraph_idxs) != len(REFERENCES):
        raise RuntimeError(f"Expected {len(REFERENCES)} reference paragraphs, found {len(ref_paragraph_idxs)}")

    for num, (idx, ref) in enumerate(zip(ref_paragraph_idxs, REFERENCES), start=1):
        set_paragraph_text(doc.paragraphs[idx], f"[{num}] {ref}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
