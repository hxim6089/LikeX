from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


INPUT = Path(r"C:\Users\甄\Downloads\YouTube推荐系统论文翻译_1.docx")
OUTPUT = Path(r"C:\Users\甄\Downloads\YouTube推荐系统论文翻译_论文格式.docx")

BODY_EA = "SimSun"
HEADING_EA = "SimHei"
LATIN = "Times New Roman"

FIGURES = {
    1: {
        "anchor": "图 1 展示了 YouTube 移动端首页的推荐效果",
        "after": "same",
        "height_cm": 4.2,
        "caption": "图 1 YouTube 移动应用首页显示的推荐内容",
    },
    2: {
        "anchor": "线上 A/B 结果并不总与离线实验结果相吻合",
        "after": "same",
        "height_cm": 4.0,
        "caption": "图 2 推荐系统架构：候选视频经由候选生成和排序后，仅向用户展示少量视频",
    },
    3: {
        "anchor": "图 3 展示了包含下述非视频观看特征的通用网络架构",
        "after": "same",
        "height_cm": 6.2,
        "caption": "图 3 深度候选生成模型架构",
    },
    4: {
        "anchor": "图 4 展示了该方案在一个随机选定视频上的有效性",
        "after": "same",
        "height_cm": 5.5,
        "caption": "图 4 加入样本年龄特征后，模型能够表示视频上传时间与随时间变化的流行度",
    },
    5: {
        "anchor": "图 5b",
        "after": "same",
        "height_cm": 4.2,
        "caption": "图 5 标签与输入上下文选择示意：预测未来观看在 A/B 测试中表现更好",
    },
    6: {
        "anchor": "深度 4：2048 ReLU",
        "after": "same",
        "height_cm": 5.2,
        "caption": "图 6 视频嵌入之外的特征和网络深度对留出集 MAP 的提升效果",
    },
    7: {
        "anchor": "通过逻辑回归为每个视频曝光独立打分（图 7）",
        "after": "same",
        "height_cm": 5.8,
        "caption": "图 7 深度排序网络架构",
    },
}


def clean_text(text: str) -> str:
    replacements = {
        "201c": "“",
        "201d": "”",
        "2018": "‘",
        "2019": "’",
        "留存集": "留出集",
        "留存数据": "留出数据",
        "留存曝光": "留出曝光",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def set_run_font(run, size=12, east_asia=BODY_EA, latin=LATIN, bold=None, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def format_paragraph_runs(paragraph, size=12, east_asia=BODY_EA, bold=None, color=None):
    if not paragraph.runs:
        run = paragraph.add_run("")
        set_run_font(run, size=size, east_asia=east_asia, bold=bold, color=color)
    for run in paragraph.runs:
        set_run_font(run, size=size, east_asia=east_asia, bold=bold, color=color)


def set_text(paragraph, text):
    paragraph.text = clean_text(text)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=10.5)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFBFBF"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_cm=15.0):
    width_dxa = int(width_cm / 2.54 * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    jc = tbl_pr.first_child_found_in("w:jc")
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def insert_paragraph_after(anchor, text="", style=None):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = anchor._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    new_p.addnext(paragraph._p)
    paragraph._p.getparent().remove(new_p)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_figure_placeholder(doc, anchor, fig_no, caption, height_cm):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_width(table, 15.0)
    set_table_borders(table, "A6A6A6")
    row = table.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F7F7F7")
    set_cell_margins(cell, top=160, bottom=160, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"此处插入图 {fig_no}")
    set_run_font(run, size=12, east_asia=HEADING_EA, bold=True, color=RGBColor(80, 80, 80))
    run = p.add_run("\n请按原 PDF 中对应图片插入，并保留下方图注")
    set_run_font(run, size=10.5, color=RGBColor(100, 100, 100))

    anchor._p.addnext(table._tbl)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(6)
    cap.add_run(caption)
    format_paragraph_runs(cap, size=10.5, east_asia=BODY_EA)
    table._tbl.addnext(cap._p)


def is_numbered_heading(text):
    return bool(re.match(r"^\d+(\.\d+)?\s*[\u4e00-\u9fffA-Za-z]", text.strip()))


def is_level1(text):
    return bool(re.match(r"^\d+\.\s+", text.strip()))


def is_level2(text):
    return bool(re.match(r"^\d+\.\d+\s+", text.strip()))


def apply_document_format(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

    footer = section.footer
    footer.paragraphs[0].text = ""
    add_page_number(footer.paragraphs[0])

    normal = None
    for style_name in ("Normal", "正文"):
        try:
            normal = doc.styles[style_name]
            break
        except KeyError:
            pass
    if normal is not None:
        normal.font.name = LATIN
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_EA)
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

    references = False
    for idx, p in enumerate(doc.paragraphs):
        set_text(p, p.text)
        text = p.text.strip()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(0.74)
        pf.left_indent = None
        pf.right_indent = None

        if idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.space_after = Pt(8)
            format_paragraph_runs(p, size=16, east_asia=HEADING_EA, bold=True)
        elif idx in (1, 2, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.line_spacing = 1.15
            format_paragraph_runs(p, size=10.5, east_asia=BODY_EA)
        elif text == "摘要":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.space_before = Pt(10)
            pf.space_after = Pt(4)
            format_paragraph_runs(p, size=12, east_asia=HEADING_EA, bold=True)
        elif text.startswith("关键词"):
            pf.first_line_indent = None
            pf.space_after = Pt(8)
            format_paragraph_runs(p, size=12, east_asia=BODY_EA)
            if p.runs:
                p.runs[0].bold = True
        elif is_level1(text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.space_before = Pt(10)
            pf.space_after = Pt(4)
            format_paragraph_runs(p, size=14, east_asia=HEADING_EA, bold=True)
            if text.startswith("7."):
                references = True
        elif is_level2(text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.space_before = Pt(8)
            pf.space_after = Pt(3)
            format_paragraph_runs(p, size=12, east_asia=HEADING_EA, bold=True)
        elif text in {"特征工程", "类别特征嵌入", "连续特征归一化"}:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.space_before = Pt(6)
            pf.space_after = Pt(3)
            format_paragraph_runs(p, size=12, east_asia=HEADING_EA, bold=True)
        elif text.startswith("["):
            pf.first_line_indent = Cm(-0.74)
            pf.left_indent = Cm(0.74)
            pf.line_spacing = 1.15
            format_paragraph_runs(p, size=10.5, east_asia=BODY_EA)
        elif text.startswith("表 "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.space_before = Pt(4)
            pf.space_after = Pt(6)
            format_paragraph_runs(p, size=10.5, east_asia=BODY_EA)
        elif not text:
            pf.first_line_indent = None
            pf.line_spacing = 1.0
            pf.space_after = Pt(0)
        else:
            format_paragraph_runs(p, size=12, east_asia=BODY_EA)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        set_table_width(table, 14.5)
        set_table_borders(table, "BFBFBF")
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.first_line_indent = None
                    format_paragraph_runs(p, size=10.5, east_asia=BODY_EA, bold=(row_idx == 0))
                if row_idx == 0:
                    set_cell_shading(cell, "F2F2F2")


def main():
    doc = Document(INPUT)
    apply_document_format(doc)

    existing_text = "\n".join(p.text for p in doc.paragraphs)
    if "此处插入图" not in existing_text:
        for fig_no in sorted(FIGURES.keys(), reverse=True):
            spec = FIGURES[fig_no]
            anchor = None
            for p in doc.paragraphs:
                if spec["anchor"] in p.text:
                    anchor = p
            if anchor is None:
                raise RuntimeError(f"未找到图 {fig_no} 的插入锚点：{spec['anchor']}")
            insert_figure_placeholder(doc, anchor, fig_no, spec["caption"], spec["height_cm"])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
