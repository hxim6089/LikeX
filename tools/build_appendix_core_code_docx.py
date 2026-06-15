# -*- coding: utf-8 -*-
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "论文相关" / "附录_程序源代码核心代码.docx"


@dataclass
class Snippet:
    title: str
    source: str
    code: str
    start_line: int | None = None
    end_line: int | None = None


@dataclass
class SectionSpec:
    title: str
    note: str
    snippets: list[Snippet]


def read_source(rel_path: str) -> str:
    return (ROOT / rel_path).read_text(encoding="utf-8")


def line_number_at(text: str, index: int) -> int:
    return text.count("\n", 0, index) + 1


def find_line_start(text: str, needle: str, occurrence: int = 1) -> tuple[int, int]:
    count = 0
    offset = 0
    while True:
        idx = text.find(needle, offset)
        if idx < 0:
            raise ValueError(f"Cannot find marker: {needle}")
        count += 1
        if count == occurrence:
            line_start = text.rfind("\n", 0, idx) + 1
            return line_start, line_number_at(text, line_start)
        offset = idx + len(needle)


def include_annotations_above(text: str, start_idx: int) -> int:
    lines_before = text[:start_idx].splitlines(keepends=True)
    if not lines_before:
        return start_idx

    idx = start_idx
    pos = start_idx
    for line in reversed(lines_before):
        pos -= len(line)
        stripped = line.strip()
        if stripped.startswith("@"):
            idx = pos
            continue
        break
    return idx


def balanced_java_block(text: str, start_idx: int) -> tuple[str, int, int]:
    start_idx = include_annotations_above(text, start_idx)
    open_idx = text.find("{", start_idx)
    if open_idx < 0:
        raise ValueError("No opening brace after marker")

    depth = 0
    in_string = False
    in_char = False
    in_line_comment = False
    in_block_comment = False
    escaped = False

    i = open_idx
    while i < len(text):
        ch = text[i]
        nxt = text[i + 1] if i + 1 < len(text) else ""

        if in_line_comment:
            if ch == "\n":
                in_line_comment = False
            i += 1
            continue

        if in_block_comment:
            if ch == "*" and nxt == "/":
                in_block_comment = False
                i += 2
            else:
                i += 1
            continue

        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            i += 1
            continue

        if in_char:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == "'":
                in_char = False
            i += 1
            continue

        if ch == "/" and nxt == "/":
            in_line_comment = True
            i += 2
            continue
        if ch == "/" and nxt == "*":
            in_block_comment = True
            i += 2
            continue
        if ch == '"':
            in_string = True
            i += 1
            continue
        if ch == "'":
            in_char = True
            i += 1
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end_idx = i + 1
                while end_idx < len(text) and text[end_idx] in "\r\n":
                    end_idx += 1
                code = text[start_idx:end_idx].strip("\r\n")
                return code, line_number_at(text, start_idx), line_number_at(text, i)
        i += 1

    raise ValueError("Unbalanced Java block")


def extract_block(rel_path: str, marker: str, occurrence: int = 1) -> Snippet:
    text = read_source(rel_path)
    start_idx, _ = find_line_start(text, marker, occurrence)
    code, start_line, end_line = balanced_java_block(text, start_idx)
    return Snippet(marker.strip(), rel_path, normalize_code(code), start_line, end_line)


def extract_between(rel_path: str, start_marker: str, end_marker: str, title: str) -> Snippet:
    text = read_source(rel_path)
    start_idx, start_line = find_line_start(text, start_marker)
    end_idx = text.find(end_marker, start_idx)
    if end_idx < 0:
        raise ValueError(f"Cannot find end marker: {end_marker}")
    code = text[start_idx:end_idx].strip("\r\n")
    end_line = line_number_at(text, end_idx)
    return Snippet(title, rel_path, normalize_code(code), start_line, end_line)


def extract_field_summary(rel_path: str, title: str) -> Snippet:
    text = read_source(rel_path)
    pattern = re.compile(r"(?ms)^public class ScoreBreakdown \{\s*(.*?)\n\s*public ScoreBreakdown\(\)")
    match = pattern.search(text)
    if not match:
        raise ValueError("Cannot extract ScoreBreakdown fields")
    header_start = text.find("public class ScoreBreakdown")
    code = text[header_start:match.end(1)].rstrip() + "\n}"
    return Snippet(title, rel_path, normalize_code(code), line_number_at(text, header_start), line_number_at(text, match.end(1)))


def normalize_code(code: str) -> str:
    code = code.replace("\t", "    ")
    code = textwrap.dedent(code).strip("\n")
    return code


def wrap_code_lines(code: str, width: int = 116) -> list[str]:
    wrapped: list[str] = []
    for raw_line in code.splitlines():
        line = raw_line.rstrip()
        if len(line) <= width:
            wrapped.append(line)
            continue
        indent = len(line) - len(line.lstrip(" "))
        continuation = " " * min(indent + 8, 24)
        parts = textwrap.wrap(
            line,
            width=width,
            subsequent_indent=continuation,
            break_long_words=False,
            break_on_hyphens=False,
            replace_whitespace=False,
        )
        wrapped.extend(parts if parts else [""])
    return wrapped


def set_run_font(run, size_pt: float, bold: bool = False, color: RGBColor | None = None, font_name: str = "宋体") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman" if font_name == "宋体" else font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman" if font_name == "宋体" else font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_paragraph_spacing(paragraph, line_spacing: float = 1.5, before: float = 0, after: float = 0) -> None:
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def add_paragraph(doc: Document, text: str, size: float = 12, bold: bool = False, align=None, after: float = 0) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, 1.5, after=after)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(p, 1.5, before=6 if level == 1 else 3, after=3)
    run = p.add_run(text)
    set_run_font(run, 14 if level == 1 else 12, bold=True, font_name="黑体")


def add_source_line(doc: Document, snippet: Snippet) -> None:
    line_range = ""
    if snippet.start_line is not None and snippet.end_line is not None:
        line_range = f"（第 {snippet.start_line}-{snippet.end_line} 行）"
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 1.5, after=2)
    run = p.add_run(f"代码来源：{snippet.source}{line_range}")
    set_run_font(run, 10.5, color=RGBColor(89, 89, 89))


def add_code_block(doc: Document, code: str) -> None:
    for line in wrap_code_lines(code):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        shade_paragraph(p, "F2F2F2")
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8.5)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 1.0, after=4)


def build_sections() -> list[SectionSpec]:
    hybrid = "backend/src/main/java/com/example/rec/service/HybridRecommendationStrategy.java"
    profile = "backend/src/main/java/com/example/rec/service/UserBehaviorProfileService.java"
    ai = "backend/src/main/java/com/example/rec/service/AiRecommendationStrategy.java"
    ad = "backend/src/main/java/com/example/rec/service/AdService.java"
    ad_controller = "backend/src/main/java/com/example/rec/controller/AdController.java"
    behavior = "backend/src/main/java/com/example/rec/service/BehaviorService.java"
    behavior_controller = "backend/src/main/java/com/example/rec/controller/BehaviorController.java"
    score_breakdown = "backend/src/main/java/com/example/rec/dto/ScoreBreakdown.java"

    return [
        SectionSpec(
            "A.1 推荐排序模块核心代码",
            "该部分对应论文中个性化信息流排序与推荐策略设计，重点展示候选内容如何经过多维度评分、行为惩罚、作者多样性处理和加权随机重排后形成最终推荐列表。",
            [
                extract_between(hybrid, "public class HybridRecommendationStrategy", "@Override", "策略依赖与权重参数"),
                extract_block(hybrid, "public List<Content> recommend(Long userId, List<Content> candidates)"),
                extract_block(hybrid, "public List<ContentWithScore> recommendWithScore(Long userId, List<Content> candidates, Map<String, Double> manualWeights)"),
                extract_block(hybrid, "private double calculateScore(Content content, List<String> userInterests,"),
                extract_block(hybrid, "private ScoredContentWithDetails calculateScoreWithDetails(Content content,"),
                extract_block(hybrid, "private void applyAuthorDiversityPenaltyWithDetails(List<ScoredContentWithDetails> scoredList)"),
                extract_block(hybrid, "private List<Content> weightedShuffle(List<ScoredContent> sortedList, double explorationFactor)"),
                extract_block(hybrid, "private double applyInteractionPenalty(Long contentId, double score,"),
                extract_block(hybrid, "private Map<String, Double> dynamicWeightsToMap(DynamicWeights dw)"),
                extract_field_summary(score_breakdown, "ScoreBreakdown 评分明细对象关键字段"),
            ],
        ),
        SectionSpec(
            "A.2 用户画像构建模块核心代码",
            "该部分对应论文中用户画像建模与动态权重调整，展示系统如何根据历史行为计算话题偏好、作者偏好、互动风格、内容深度偏好、新鲜度偏好与探索度。",
            [
                extract_block(profile, "public enum UserStage"),
                extract_block(profile, "public static class BehaviorProfile"),
                extract_block(profile, "public BehaviorProfile buildProfile(Long userId)"),
                extract_block(profile, "private Map<String, Double> computeTopicPreferences(List<Behavior> behaviors,"),
                extract_block(profile, "private Map<Long, Double> computeAuthorPreferences(List<Behavior> behaviors,"),
                extract_block(profile, "private String computeEngagementStyle(List<Behavior> behaviors)"),
                extract_block(profile, "private String computeDepthPreference(List<Behavior> behaviors, Map<Long, Content> contentMap)"),
                extract_block(profile, "private double computeFreshnessPreference(List<Behavior> behaviors,"),
                extract_block(profile, "private double computeExplorationRate(Map<String, Double> topicPreferences)"),
                extract_block(profile, "public static class DynamicWeights"),
                extract_block(profile, "public DynamicWeights computeDynamicWeights(Long userId)"),
                extract_block(profile, "private void adjustByEngagementStyle(DynamicWeights w, String style)"),
                extract_block(profile, "private void adjustByDepthPreference(DynamicWeights w, String depthPref)"),
                extract_block(profile, "private void adjustByFreshnessPreference(DynamicWeights w, double freshness)"),
            ],
        ),
        SectionSpec(
            "A.3 AI 推荐与降级模块核心代码",
            "该部分对应论文中 AI 辅助推荐与系统降级策略，展示系统如何构造提示词、调用本地模型、解析排序结果，并在模型不可用时回退到基础热度排序。",
            [
                extract_between(ai, "public class AiRecommendationStrategy", "@Override", "AI 推荐策略依赖与参数"),
                extract_block(ai, "public List<Content> recommend(Long userId, List<Content> candidates)"),
                extract_block(ai, "public List<ContentWithScore> recommendWithScore(Long userId, List<Content> candidates)"),
                extract_block(ai, "private List<Content> presortByEngagement(List<Content> candidates, int limit)"),
                extract_block(ai, "private String buildPrompt(BehaviorProfile profile, List<Content> candidates)"),
                extract_block(ai, "private AiRankingResult callAiRankingFull(BehaviorProfile profile, List<Content> candidates)"),
                extract_block(ai, "private AiRankingResult parseAiResponse(String text, List<Content> candidates)"),
                extract_block(ai, "private List<Content> applyAiRanking(List<Content> allCandidates, List<Content> topCandidates,"),
                extract_block(ai, "private List<Content> fallbackSort(List<Content> candidates)"),
                extract_block(ai, "public boolean isOllamaAvailable()"),
            ],
        ),
        SectionSpec(
            "A.4 广告分发模块核心代码",
            "该部分对应论文中广告位与信息流结合的实现，展示广告如何依据用户画像标签进行相关性匹配，并记录展示、点击和统计指标。",
            [
                extract_between(ad, "public class AdService", "public AdService", "广告服务依赖与投放限制"),
                extract_block(ad, "public List<Map<String, Object>> getRelevantAds(Long userId, int count)"),
                extract_block(ad, "public void recordImpression(Long adId, Long userId)"),
                extract_block(ad, "public void recordClick(Long adId)"),
                extract_block(ad, "public Map<String, Object> getAdStats()"),
                extract_block(ad_controller, "public ResponseEntity<?> getRelevantAds("),
                extract_block(ad_controller, "public ResponseEntity<?> recordImpression("),
                extract_block(ad_controller, "public ResponseEntity<?> recordClick(@PathVariable Long id)"),
            ],
        ),
        SectionSpec(
            "A.5 用户行为记录模块核心代码",
            "该部分对应论文中行为数据采集与推荐反馈闭环，展示点赞、点踩、浏览和跳过等行为如何写入数据库，并作为后续用户画像与推荐排序的输入。",
            [
                extract_between(behavior, "public class BehaviorService", "public BehaviorService", "行为服务依赖关系"),
                extract_block(behavior, "public void likeContent(Long userId, Long contentId)"),
                extract_block(behavior, "public void recordView(Long userId, Long contentId, Integer duration)"),
                extract_block(behavior, "public void dislikeContent(Long userId, Long contentId)"),
                extract_block(behavior, "public void recordSkip(Long userId, Long contentId)"),
                extract_between(behavior_controller, "public class BehaviorController", "public BehaviorController", "行为接口控制器依赖"),
                extract_block(behavior_controller, "public String like(@RequestBody Map<String, Long> payload)"),
                extract_block(behavior_controller, "public String view(@RequestBody Map<String, Object> payload)"),
                extract_block(behavior_controller, "public String dislike(@RequestBody Map<String, Long> payload)"),
                extract_block(behavior_controller, "public String skip(@RequestBody Map<String, Long> payload)"),
            ],
        ),
    ]


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)

    heading1 = styles["Heading 1"]
    heading1.font.name = "黑体"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading1._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading1._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading1.font.size = Pt(14)
    heading1.font.bold = True
    heading1.paragraph_format.line_spacing = 1.5
    heading1.paragraph_format.space_before = Pt(6)
    heading1.paragraph_format.space_after = Pt(3)

    heading2 = styles["Heading 2"]
    heading2.font.name = "黑体"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading2._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading2._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.paragraph_format.line_spacing = 1.5
    heading2.paragraph_format.space_before = Pt(3)
    heading2.paragraph_format.space_after = Pt(3)


def build_docx() -> Path:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, 1.5, after=8)
    run = title.add_run("附录  程序源代码")
    set_run_font(run, 16, bold=True, font_name="黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, 1.5, after=10)
    run = subtitle.add_run("模块的核心代码")
    set_run_font(run, 14, bold=True, font_name="黑体")

    add_paragraph(
        doc,
        "本文档依据当前项目源码整理，选取与论文系统设计和实现关系最紧密的核心模块代码。为便于毕业论文附录排版，文档保留关键方法、参数、注释和主要业务逻辑，省略实体类、Repository 接口、Getter/Setter 等基础性代码。",
        size=12,
        after=6,
    )
    add_paragraph(
        doc,
        "代码模块包括推荐排序、用户画像、AI 推荐与降级、广告分发以及用户行为记录五个部分，可作为论文“附录 程序源代码”中的核心代码材料。",
        size=12,
        after=8,
    )

    for section_index, spec in enumerate(build_sections(), start=1):
        if section_index > 1:
            doc.add_page_break()
        add_heading(doc, spec.title, level=1)
        add_paragraph(doc, spec.note, size=12, after=4)

        for snippet in spec.snippets:
            add_heading(doc, snippet.title, level=2)
            add_source_line(doc, snippet)
            add_code_block(doc, snippet.code)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_docx()
    print(path)
