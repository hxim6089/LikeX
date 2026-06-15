from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
SOURCE_PDF = Path(r"D:\下载\s11257-023-09379-6 (1).pdf")
OUTPUT_TXT = ROOT / "论文相关" / "Hybrid_session-aware_recommendation_英文正文及参考文献_编号版.txt"
OUTPUT_DOCX = ROOT / "论文相关" / "Hybrid_session-aware_recommendation_英文正文及参考文献_编号版.docx"
OUTPUT_REFERENCES_TXT = ROOT / "论文相关" / "Hybrid_session-aware_recommendation_仅参考文献_编号版.txt"

TITLE = "Hybrid session-aware recommendation with feature-based models"
AUTHORS = "Josef Bauer · Dietmar Jannach"

HEADER_PATTERNS = (
    re.compile(r"^\d{3}\s+J\.\s+Bauer,\s+D\.\s+Jannach$"),
    re.compile(r"^Hybrid session-aware recommendation with feature-based models\s+\d{3}$"),
)

SECTION_PATTERN = re.compile(
    r"^(?:"
    r"\d+(?:\.\d+){0,3}\s+\S.*|"
    r"Abstract|Keywords\b.*|References|"
    r"Appendix\s+A:.*|"
    r"Table\s+\d+.*|Fig\.\s+\d+.*|Algorithm\s+\d+.*"
    r")$"
)

REFERENCE_START_PATTERN = re.compile(r"^[A-Z][A-Za-zÀ-ÖØ-öø-ÿ’' .-]+,\s.*:")


def normalize_text(text: str) -> str:
    replacements = {
        "\ufb00": "ff",
        "\ufb01": "fi",
        "\ufb02": "fl",
        "\ufb03": "ffi",
        "\ufb04": "ffl",
        "\u00ad": "",
        "\xa0": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    artifact_replacements = {
        "ofcollaborative": "of collaborative",
        "V ery": "Very",
        "R e n": "Ren",
        "Y uan": "Yuan",
        "Y u": "Yu",
        "Y ang": "Yang",
        "S A S R e c": "SASRec",
        "w ei n t r o - duce": "we introduce",
        "H i d a s ie ta l .": "Hidasi et al.",
        "T h e": "The",
        "theactual": "the actual",
        "N o wt h et a s ki st oe n r i c h": "Now the task is to enrich",
        "a r e": "are",
        "U k i le ta l .": "Ukil et al.",
        "theSHAP": "the SHAP",
        "sessionbased": "session-based",
        "longterm": "long-term",
        "timedependent": "time-dependent",
        "add-tocart": "add-to-cart",
        "o rSASRec": "or SASRec",
        "w ei n t r o duce": "we introduce",
        "Sect. 2,w ei n t r o - duce": "Sect. 2, we introduce",
        "Sect. 2,w ei n t r o duce": "Sect. 2, we introduce",
        "Sect. 2,we introduce": "Sect. 2, we introduce",
        "Table2.": "Table 2.",
        "Table6": "Table 6",
        "Table7": "Table 7",
        "Sect.2": "Sect. 2",
        "Sect.4.1": "Sect. 4.1",
        "Sect. 4.3.N o t e": "Sect. 4.3. Note",
    }
    for old, new in artifact_replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\b([A-Z]) \.", r"\1.", text)
    text = re.sub(r"(?<=[A-Za-z])- (?=[a-z])", "-", text)
    return text


def extract_clean_lines(pdf_path: Path) -> list[str]:
    reader = PdfReader(str(pdf_path))
    lines: list[str] = []
    for page in reader.pages:
        page_text = normalize_text(page.extract_text() or "")
        for raw in page_text.splitlines():
            line = re.sub(r"\s+", " ", raw).strip()
            if not line or line == "123":
                continue
            if any(pattern.match(line) for pattern in HEADER_PATTERNS):
                continue
            lines.append(line)

    joined = "\n".join(lines)
    joined = re.sub(
        r"Hybrid session-aware recommendation with feature-based\nmodels",
        TITLE,
        joined,
        count=1,
    )

    title_pos = joined.find(TITLE)
    if title_pos == -1:
        raise RuntimeError("Could not locate the article title in the extracted PDF text.")
    joined = joined[title_pos:]

    publisher_pos = joined.find("Publisher’s Note")
    if publisher_pos != -1:
        joined = joined[:publisher_pos]

    joined = re.sub(
        r"B Josef Bauer.*?1 University of Klagenfurt, Klagenfurt, Austria",
        "",
        joined,
        flags=re.DOTALL,
    )
    joined = re.sub(
        r"Received:.*?© The Author\(s\) 2023",
        "",
        joined,
        flags=re.DOTALL,
    )
    joined = re.sub(
        r"Author Contributions.*?(?=Appendix A: List of engineered features)",
        "",
        joined,
        flags=re.DOTALL,
    )

    # Repair words split by PDF line wrapping, e.g. "recommen-\ndation".
    joined = re.sub(r"(?<=[A-Za-z])-\n(?=[a-z])", "", joined)
    # Some artifacts only become visible after the PDF's wrapped lines are joined.
    joined = normalize_text(joined)
    raw_cleaned = []
    for raw in joined.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line or line == "123":
            continue
        if any(pattern.match(line) for pattern in HEADER_PATTERNS):
            continue
        if re.match(r"^Josef Bauer\s+\d+\s+·\s+Dietmar Jannach\s+\d+$", line):
            line = AUTHORS
        raw_cleaned.append(line)

    cleaned: list[str] = []
    index = 0
    while index < len(raw_cleaned):
        line = raw_cleaned[index]
        if line.startswith("Keywords ") and index + 1 < len(raw_cleaned):
            next_line = raw_cleaned[index + 1]
            if not is_heading(next_line):
                line = f"{line} {next_line}"
                index += 1
        cleaned.append(line)
        index += 1
    return cleaned


def is_heading(line: str) -> bool:
    return bool(SECTION_PATTERN.match(line))


def should_end_paragraph(line: str) -> bool:
    if re.search(r"\b(?:et al|e\.g|i\.e|cf)\.$", line):
        return False
    return bool(re.search(r"[.!?)](?:\s*\d+)?$", line))


def build_blocks(lines: list[str]) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    body_buffer: list[str] = []
    bullet_buffer: list[str] = []
    reference_buffer: list[str] = []
    in_references = False

    def flush_body() -> None:
        if body_buffer:
            blocks.append(("body", normalize_text(" ".join(body_buffer))))
            body_buffer.clear()

    def flush_reference() -> None:
        if reference_buffer:
            blocks.append(("reference", normalize_text(" ".join(reference_buffer))))
            reference_buffer.clear()

    def flush_bullet() -> None:
        if bullet_buffer:
            blocks.append(("bullet", normalize_text(" ".join(bullet_buffer))))
            bullet_buffer.clear()

    for index, line in enumerate(lines):
        if index == 0 and line == TITLE:
            blocks.append(("title", line))
            continue
        if line == AUTHORS:
            blocks.append(("authors", line))
            continue
        if line == "References":
            flush_body()
            flush_bullet()
            flush_reference()
            in_references = True
            blocks.append(("heading1", line))
            continue

        if in_references:
            if REFERENCE_START_PATTERN.match(line):
                flush_reference()
                reference_buffer.append(line)
            else:
                reference_buffer.append(line)
            continue

        if is_heading(line):
            flush_body()
            flush_bullet()
            if re.match(r"^\d+\s", line) or line.startswith("Appendix A"):
                blocks.append(("heading1", line))
            elif re.match(r"^\d+\.\d+\s", line):
                blocks.append(("heading2", line))
            elif line in {"Abstract"} or line.startswith("Keywords"):
                blocks.append(("heading2", line))
            elif line.startswith(("Table ", "Fig. ", "Algorithm ")):
                blocks.append(("caption", line))
            else:
                blocks.append(("heading2", line))
            continue

        if line.startswith(("–", "- ")):
            flush_body()
            flush_bullet()
            bullet_buffer.append(line.lstrip("–- ").strip())
            continue

        if bullet_buffer:
            bullet_buffer.append(line)
            if should_end_paragraph(line):
                flush_bullet()
            continue

        body_buffer.append(line)
        if should_end_paragraph(line):
            flush_body()

    flush_body()
    flush_bullet()
    flush_reference()
    return blocks


def write_text(blocks: list[tuple[str, str]], path: Path) -> None:
    rendered = []
    reference_number = 0
    for block_type, content in blocks:
        if block_type == "reference":
            reference_number += 1
            rendered.append(f"[{reference_number}] {content}")
        else:
            rendered.append(f"– {content}" if block_type == "bullet" else content)
    text = "\n\n".join(rendered).strip() + "\n"
    path.write_text(text, encoding="utf-8")


def write_references_text(blocks: list[tuple[str, str]], path: Path) -> None:
    references = [
        f"[{number}] {content}"
        for number, (_, content) in enumerate(
            (block for block in blocks if block[0] == "reference"),
            start=1,
        )
    ]
    path.write_text("\n\n".join(references).strip() + "\n", encoding="utf-8")


def set_run_font(run, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, 10)


def write_docx(blocks: list[tuple[str, str]], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)

    reference_number = 0
    for block_type, content in blocks:
        p = doc.add_paragraph()
        if block_type == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            set_run_font(p.add_run(content), 16, bold=True)
        elif block_type == "authors":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            set_run_font(p.add_run(content), 11)
        elif block_type == "heading1":
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            set_run_font(p.add_run(content), 13, bold=True)
        elif block_type == "heading2":
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            set_run_font(p.add_run(content), 11.5, bold=True)
        elif block_type == "caption":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(4)
            set_run_font(p.add_run(content), 10, italic=True)
        elif block_type == "bullet":
            p.style = doc.styles["List Bullet"]
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(content), 11)
        elif block_type == "reference":
            reference_number += 1
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.space_after = Pt(3)
            set_run_font(p.add_run(f"[{reference_number}] {content}"), 10)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.15
            set_run_font(p.add_run(content), 11)

    doc.save(path)


def main() -> None:
    if not SOURCE_PDF.exists():
        raise FileNotFoundError(SOURCE_PDF)
    lines = extract_clean_lines(SOURCE_PDF)
    blocks = build_blocks(lines)
    write_text(blocks, OUTPUT_TXT)
    write_references_text(blocks, OUTPUT_REFERENCES_TXT)
    write_docx(blocks, OUTPUT_DOCX)

    text = OUTPUT_TXT.read_text(encoding="utf-8")
    print(f"pages={len(PdfReader(str(SOURCE_PDF)).pages)}")
    print(f"blocks={len(blocks)}")
    print(f"chars={len(text)}")
    print(f"references={sum(1 for block_type, _ in blocks if block_type == 'reference')}")
    print(f"has_references_heading={'References' in text}")
    print(f"txt={OUTPUT_TXT}")
    print(f"references_txt={OUTPUT_REFERENCES_TXT}")
    print(f"docx={OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
