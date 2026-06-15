from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "2959100.2959190.zh.md"
OUTPUT = ROOT / "2959100.2959190.zh.docx"

ACCENT = RGBColor(0, 0, 0)
MUTED = RGBColor(96, 96, 96)


def set_run_font(run, east_asia="SimSun", ascii_font="Times New Roman", size=None, bold=None, color=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, east_asia="SimSun", ascii_font="Times New Roman", size=10.5, color=None):
    for run in paragraph.runs:
        set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EA")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    jc = tbl_pr.first_child_found_in("w:jc")
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "left")


def add_paragraph_box(paragraph, color="BFBFBF", fill="F8F8F8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        node = p_bdr.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            p_bdr.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), color)
    shd = p_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_figure_placeholder(doc, caption_line):
    match = re.match(r"\*\*图\s*(\d+)：\*\*\s*(.*)", caption_line)
    if not match:
        return
    fig_no, caption = match.groups()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.right_indent = Cm(0.25)
    add_paragraph_box(p)
    run = p.add_run(f"【此处插入图 {fig_no}】")
    set_run_font(run, size=10.5, bold=True, color=RGBColor(80, 80, 80))
    run = p.add_run(f"\n{caption[:72]}{'...' if len(caption) > 72 else ''}")
    set_run_font(run, size=9.5, color=MUTED)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_run_font(run, size=9, color=MUTED)


def add_border_bottom(paragraph, color="D9E2EA", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_markdown_runs(paragraph, text, size=11, color=None):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        set_run_font(run, size=size, bold=bold, color=color)


def add_hyperlink_like_run(paragraph, text, size=10):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=ACCENT)
    run.underline = True
    return run


def add_cover_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("用于 YouTube 推荐的深度神经网络")
    set_run_font(run, size=24, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Deep Neural Networks for YouTube Recommendations")
    set_run_font(run, size=13, color=MUTED)

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    set_table_width(meta, width_dxa=7600)
    set_table_borders(meta)
    rows = [
        ("作者", "Paul Covington, Jay Adams, Emre Sargin"),
        ("机构", "Google, Mountain View, CA"),
        ("会议", "RecSys 2016, Boston, MA, USA"),
        ("DOI", "10.1145/2959100.2959190"),
        ("文档类型", "中文译文整理版"),
    ]
    for r_idx, (label, value) in enumerate(rows):
        cells = meta.rows[r_idx].cells
        for c in cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c, top=110, bottom=110, start=150, end=150)
        set_cell_shading(cells[0], "F3F6F9")
        cells[0].text = ""
        cells[1].text = ""
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p0.add_run(label)
        set_run_font(run, size=10, bold=True, color=RGBColor(60, 60, 60))
        p1 = cells[1].paragraphs[0]
        run = p1.add_run(value)
        set_run_font(run, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("说明：本文档依据用户提供的 PDF 翻译整理，保留原论文结构、公式、图表说明与参考文献编号。")
    set_run_font(run, size=10, color=MUTED)

    doc.add_page_break()


def collect_toc_entries(lines):
    entries = []
    for line in lines:
        if line.startswith("## "):
            title = line[3:].strip()
            if title != "摘要":
                entries.append((1, title))
            else:
                entries.append((1, title))
        elif line.startswith("### "):
            entries.append((2, line[4:].strip()))
    return entries


def add_manual_toc(doc, entries):
    p = doc.add_paragraph(style="Heading 1")
    add_markdown_runs(p, "目录", size=16, color=ACCENT)

    for level, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.0 if level == 1 else 0.55)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        set_run_font(run, size=11 if level == 1 else 10.5, bold=(level == 1), color=RGBColor(50, 50, 50))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("注：页码可在 Word 中通过“引用 -> 目录”更新为自动目录；本文档标题样式已配置为可识别层级。")
    set_run_font(run, size=9, color=MUTED)
    doc.add_page_break()


def is_table_block(lines, idx):
    return idx + 1 < len(lines) and lines[idx].startswith("|") and re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", lines[idx + 1])


def parse_table(lines, idx):
    rows = []
    while idx < len(lines) and lines[idx].startswith("|"):
        row = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(row)
        idx += 1
    return rows[0], rows[2:], idx


def configure_styles(doc):
    styles = doc.styles
    for style_name, size in (("Normal", 11), ("Body Text", 11)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(3)

    for style_name, size, color in (
        ("Title", 16, ACCENT),
        ("Subtitle", 10.5, MUTED),
        ("Heading 1", 12, ACCENT),
        ("Heading 2", 11, ACCENT),
        ("Heading 3", 10.5, RGBColor(40, 40, 40)),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(size)
        style.font.bold = True if style_name != "Subtitle" else False
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6 if style_name != "Subtitle" else 12)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def build_docx():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    run = header_p.add_run("用于 YouTube 推荐的深度神经网络：中文译文")
    set_run_font(run, size=9, color=MUTED)
    add_border_bottom(header_p)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    in_code = False
    code_lines = []
    title_done = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                for j, code in enumerate(code_lines):
                    if j:
                        p.add_run().add_break()
                    run = p.add_run(code)
                    set_run_font(run, east_asia="DengXian", ascii_font="Consolas", size=10, color=RGBColor(64, 64, 64))
                p.paragraph_format.left_indent = Cm(0.35)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if is_table_block(lines, i):
            headers, rows, i = parse_table(lines, i)
            table = doc.add_table(rows=1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = "Table Grid"
            set_table_width(table)
            set_table_borders(table)
            for col_idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                add_markdown_runs(p, header_text, size=10, color=RGBColor(30, 30, 30))
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "F3F6F9")
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for row in rows:
                cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    cells[col_idx].text = ""
                    p = cells[col_idx].paragraphs[0]
                    add_markdown_runs(p, value, size=10)
                    if col_idx == len(row) - 1:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    set_cell_margins(cells[col_idx])
                    cells[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_markdown_runs(p, line[2:], size=16, color=ACCENT)
            title_done = True
        elif title_done and re.match(r"^(Paul|Google|\{)", line):
            p = doc.add_paragraph(style="Subtitle")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_markdown_runs(p, line.replace("  ", ""), size=10.5, color=MUTED)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_markdown_runs(p, line[3:], size=12, color=ACCENT)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_markdown_runs(p, line[4:], size=11, color=ACCENT)
        elif line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_markdown_runs(p, line[5:], size=10.5, color=RGBColor(40, 40, 40))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(p, line[2:], size=10.5)
        elif line.startswith("**图") or line.startswith("**表"):
            if line.startswith("**图"):
                add_figure_placeholder(doc, line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.35)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_markdown_runs(p, line, size=9.5, color=RGBColor(50, 50, 50))
            for run in p.runs:
                run.italic = True
        elif re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(3)
            add_markdown_runs(p, line, size=9)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            add_markdown_runs(p, line, size=10.5)
        i += 1

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_docx())
