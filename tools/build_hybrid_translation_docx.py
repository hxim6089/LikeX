from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "论文相关" / "Hybrid_session-aware_recommendation_中文正文15000.md"
OUTPUT = ROOT / "论文相关" / "Hybrid_session-aware_recommendation_中文正文15000字.docx"

BODY_EA = "SimSun"
HEAD_EA = "SimHei"
LATIN = "Times New Roman"


def set_run_font(run, size=12, east_asia=BODY_EA, latin=LATIN, bold=None, italic=None, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_runs(paragraph, text, size=12, east_asia=BODY_EA, bold=False, italic=False, color=None):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        part_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if part_bold else part
        run = paragraph.add_run(content)
        set_run_font(run, size=size, east_asia=east_asia, bold=(bold or part_bold), italic=italic, color=color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=10.5)


def set_cell_margins(cell, top=160, start=160, bottom=160, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="A6A6A6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_cm=15.0):
    width_dxa = int(width_cm / 2.54 * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    jc = tbl_pr.first_child_found_in("w:jc")
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def add_figure_placeholder(doc, caption_line):
    match = re.match(r"\*\*图\s*(\d+)：\*\*\s*(.*)", caption_line)
    if not match:
        return
    fig_no, caption = match.groups()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table)
    row = table.rows[0]
    row.height = Cm(6.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F7F7F7")
    set_cell_margins(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"此处插入图 {fig_no}")
    set_run_font(run, size=12, east_asia=HEAD_EA, bold=True, color=RGBColor(80, 80, 80))
    run = p.add_run("\n请按原 PDF 中对应图片插入，并保留下方图注")
    set_run_font(run, size=10.5, color=RGBColor(100, 100, 100))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(6)
    add_runs(cap, f"图 {fig_no} {caption}", size=10.5)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)
    add_page_number(section.footer.paragraphs[0])

    for style_name in ("Normal", "Body Text"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_EA)
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)


def format_body_paragraph(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    title_seen = False

    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            add_runs(p, line[2:], size=16, east_asia=HEAD_EA, bold=True)
            title_seen = True
        elif title_seen and re.match(r"^(Josef|User Modeling|DOI)", line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            add_runs(p, line.replace("  ", ""), size=10.5)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if line.startswith("## 摘要") else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, line[3:], size=12 if line.startswith("## 摘要") else 14, east_asia=HEAD_EA, bold=True)
        elif line.startswith("**图"):
            add_figure_placeholder(doc, line)
        else:
            p = doc.add_paragraph()
            format_body_paragraph(p)
            add_runs(p, line, size=12)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
