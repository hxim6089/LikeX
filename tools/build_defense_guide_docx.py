import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = "1F4E79"
DARK_BLUE = "17365D"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MUTED = "64748B"
BLACK = "111827"
TABLE_WIDTH_DXA = 10080
TABLE_INDENT_DXA = 0


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=35, start=65, bottom=35, end=65):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=None,
                 bold=None, italic=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    heading_specs = {
        "Heading 1": ("黑体", 14, BLUE, 6, 3),
        "Heading 2": ("黑体", 12, BLUE, 5, 2),
        "Heading 3": ("黑体", 10.5, DARK_BLUE, 4, 2),
    }
    for name, (font, size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Cm(0.65)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.05


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("毕业答辩常见问题与参考回答 · 打印速查版")
    set_run_font(run, east_asia="宋体", size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("第 ")
    set_run_font(r, east_asia="宋体", size=9, color=MUTED)
    add_field(fp, "PAGE")
    r = fp.add_run(" 页 / 共 ")
    set_run_font(r, east_asia="宋体", size=9, color=MUTED)
    add_field(fp, "NUMPAGES")
    r = fp.add_run(" 页")
    set_run_font(r, east_asia="宋体", size=9, color=MUTED)


def add_inline(paragraph, text, base_size=9.5, base_color=BLACK):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="微软雅黑", latin="Consolas",
                         size=max(base_size - 0.5, 9), color=DARK_BLUE)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEF3F8")
            run._element.get_or_add_rPr().append(shading)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", latin="Arial", size=18, bold=True, color=DARK_BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    run = sub.add_run("答辩现场打印速查版")
    set_run_font(run, east_asia="宋体", size=10, color=MUTED)


def add_heading(doc, level, text, major_index):
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    add_inline(p, text, base_size={1: 14, 2: 12, 3: 10.5}[min(level, 3)], base_color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    add_inline(p, text, base_size=9.2, base_color=DARK_BLUE)


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F6F9")
    p_pr.append(shd)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, east_asia="微软雅黑", latin="Consolas", size=8.5, color=BLACK)
        if idx < len(lines) - 1:
            run.add_break()


def table_widths(headers):
    cols = len(headers)
    if cols == 2:
        return [2500, 7580]
    if cols == 3:
        return [2800, 4800, 2480]
    if cols == 4:
        return [1700, 2300, 3600, 2480]
    return [TABLE_WIDTH_DXA // cols] * cols


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = table_widths(headers)

    header_cells = table.rows[0].cells
    for idx, value in enumerate(headers):
        header_cells[idx].text = ""
        p = header_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, value, base_size=8.8, base_color=DARK_BLUE)
        for run in p.runs:
            run.bold = True
        set_cell_shading(header_cells[idx], LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])

    for row_idx, values in enumerate(rows[1:], start=1):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            cells[col_idx].text = ""
            p = cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, base_size=8.5, base_color=BLACK)
            if row_idx % 2 == 0:
                set_cell_shading(cells[col_idx], "F8FAFC")
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        values = [value.strip() for value in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        idx += 1
    return rows, idx


def build_docx(source, output):
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    idx = 0
    major_index = 0
    in_code = False
    code_lines = []
    first_title_seen = False

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue
        if not stripped or stripped == "---":
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            if rows:
                add_table(doc, rows)
            continue
        if stripped.startswith("# "):
            if not first_title_seen:
                add_title(doc, stripped[2:].strip())
                first_title_seen = True
            idx += 1
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            hashes, text = heading_match.groups()
            level = len(hashes) - 1
            if level == 1:
                add_heading(doc, 1, text, major_index)
                major_index += 1
            else:
                add_heading(doc, level, text, major_index)
            idx += 1
            continue
        if stripped.startswith("> "):
            quote_parts = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_parts.append(lines[idx].strip().lstrip(">").strip())
                idx += 1
            add_quote(doc, " ".join(quote_parts))
            continue
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet_match or number_match:
            p = doc.add_paragraph(style="List Bullet" if bullet_match else "List Number")
            add_inline(p, (bullet_match or number_match).group(1))
            idx += 1
            continue

        parts = [stripped]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if (not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|")
                    or nxt.startswith(">") or nxt.startswith("```")
                    or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt)):
                break
            parts.append(nxt)
            idx += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(parts))

    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
        if p.text == "参考回答：":
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(BLUE)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_defense_guide_docx.py SOURCE.md OUTPUT.docx")
    build_docx(Path(sys.argv[1]), Path(sys.argv[2]))
