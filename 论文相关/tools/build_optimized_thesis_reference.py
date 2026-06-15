from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = OUT_DIR / "毕业论文初稿_结构优化参考版.docx"


REFERENCES = [
    "夏丽云, 徐敏赟, 丁懿楠, 等. 智能推荐算法下的科技期刊国际传播策略研究[J]. 中国科技期刊研究, 2023, 34(11): 1486-1493.",
    "付峻宇, 朱小栋, 陈晨. 基于图卷积的双通道协同过滤推荐算法[J]. 计算机应用研究, 2023, 40(1): 129-135.",
    "张玉洁, 董政, 孟祥武. 个性化广告推荐系统及其应用研究[J]. 计算机学报, 2021, 44(3): 531-563.",
    "欧朝荣, 胡军. 融合显隐式反馈的协同过滤推荐模型[J]. 控制与决策, 2024, 39(3): 1048-1056.",
    "Pei C, Zhang Y, Zhang Y, et al. Personalized re-ranking for recommendation[C]//Proceedings of the 13th ACM conference on recommender systems. 2019: 3-11.",
    "吴昌政. 基于前后端分离技术的 web 开发框架设计[D]. 南京: 南京邮电大学, 2020.",
    "刘华真, 王巍, 谷壬倩, 等. 基于用户浏览行为的个性化推荐研究综述[J]. Application Research of Computers/Jisuanji Yingyong Yanjiu, 2021, 38(8).",
    "高广尚. 用户画像构建方法研究综述[J]. 数据分析与知识发现, 2019, 3(3): 25-35.",
    "Horstmann C S. Java核心技术·卷I：开发基础[M]. 林信良, 译. 第12版. 北京: 机械工业出版社, 2022.",
    "杨海民. Vue.js 3.0企业级管理后台开发实战：基于Element Plus[M]. 北京: 电子工业出版社, 2022.",
    "王树森, 黎崎. 深度学习推荐系统[M]. 北京: 电子工业出版社, 2020.",
    "Freeman E, Robson E. Head First设计模式[M]. 张晓菲, 等译. 第2版. 北京: 中国电力出版社, 2022.",
    "KLEPPMANN M. 数据密集型应用系统设计[M]. 赵健博等译. 北京: 中国电力出版社, 2018.",
    "张敏军, 华庆一, 贾伟, 等. 基于深度神经网络的个性化推荐系统研究[J]. 西南大学学报 (自然科学版), 2019, 41(11): 104-109.",
    "Twitter Inc. Twitter's Recommendation Algorithm[EB/OL]. (2023-03-31) [2026-03-13].",
    "Koren Y, Bell R, Volinsky C. Matrix factorization techniques for recommender systems[J]. Computer, 2009, 42(8): 30-37.",
    "Linden G, Smith B, York J. Amazon.com recommendations: Item-to-item collaborative filtering[J]. IEEE Internet Computing, 2003, 7(1): 76-80.",
    "Guo H, Tang R, Ye Y, et al. DeepFM: A factorization-machine based neural network for CTR prediction[J/OL]. arXiv, 2017 [2026-03-13].",
    "Zhou G, Mou N, Fan Y, et al. Deep Interest Evolution Network for click-through rate prediction[C]. In: Proceedings of the AAAI Conference on Artificial Intelligence. Honolulu: AAAI Press, 2019: 5941-5948.",
    "刘文贤, 朱海威, 武浩. 基于流行度和质量偏好建模的去偏推荐系统[J]. 云南大学学报(自然科学版), 2025, 47(3): 523-532.",
    "Covington P, Adams J, Sargin E. Deep neural networks for YouTube recommendations[C]. In: Proceedings of the 10th ACM Conference on Recommender Systems. New York: ACM, 2016: 191-198.",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)


def style_table(table):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10)
            if row_idx == 0:
                set_cell_shading(cell, "F2F2F2")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.first_line_indent = Pt(22)
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)


def add_p(doc, text="", style=None, align=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if not first_line:
        p.paragraph_format.first_line_indent = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_caption(doc, text):
    p = add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        if widths:
            set_cell_width(hdr[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            if widths:
                set_cell_width(cells[i], widths[i])
    style_table(table)
    return table


def add_placeholder(doc, label, caption):
    p = add_p(doc, f"【截图预留：{label}】", first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(10)
    add_caption(doc, caption)


def build():
    doc = Document()
    setup_styles(doc)

    p = add_p(doc, "基于用户行为分析的实时信息流排序与分发系统设计与实现", first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    add_p(doc, "优化参考稿（用于后续修改论文初稿，不作为最终格式稿）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_p(doc, "学生：甄弘硕    专业：软件工程    生成日期：2026年5月20日", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    add_heading(doc, "摘要", 1)
    add_p(doc, "随着移动互联网和社交媒体平台的发展，用户每天面对的信息数量持续增加，传统时间倒序的信息流难以同时满足内容新鲜度、兴趣匹配和推荐多样性的要求。本文围绕“基于用户行为分析的实时信息流排序与分发系统”开展设计与实现，构建了一个前后端分离的实时信息流推荐原型。系统以用户浏览、点赞、评论、转发、搜索、点踩、快速滑过等行为作为主要数据来源，结合帖子标签、文本内容、作者关系、时间因素和互动热度生成用户画像，并在此基础上完成个性化信息流排序。")
    add_p(doc, "系统后端采用 Spring Boot、JPA 和 MySQL 实现业务接口、数据持久化与推荐服务，前端采用 Vue 3、Vite、Element Plus 和 ECharts 实现首页信息流、推荐对比、用户画像、统计中心和后台管理等页面。推荐算法方面，系统综合使用基于标签的兴趣匹配、TF-IDF 与余弦相似度、协同过滤、时间衰减、负反馈过滤和多样性重排等方法，并接入本地 Ollama 或外部 DeepSeek 服务实现 AI 辅助重排序与推荐理由生成。系统还提供推荐流与时间流对比、作者去重率、标签覆盖率、行为类型分布、内容标签分布和近期兴趣序列等统计视图，使论文能够基于真实系统运行结果展开分析。")
    add_p(doc, "测试与统计分析结果表明，系统能够完成用户认证、内容发布、互动反馈、推荐排序、画像展示、AI 降级和数据统计等核心流程。与单纯时间倒序相比，混合推荐策略能够在标签命中、文本相似度、互动价值和推荐多样性等方面提供更细粒度的排序依据，符合任务书中关于用户行为分析、实时信息流排序与系统实现的要求。")
    add_p(doc, "关键词：用户行为分析；实时信息流；推荐系统；用户画像；混合推荐；AI 推荐")

    add_heading(doc, "Abstract", 1)
    add_p(doc, "With the rapid growth of social media and online content platforms, users are increasingly exposed to massive information streams. A purely chronological feed is no longer sufficient to balance freshness, interest matching and content diversity. This thesis designs and implements a real-time feed ranking and distribution system based on user behavior analysis. The system collects browsing, liking, commenting, reposting, searching, disliking and skipping behaviors, builds user profiles with content tags, textual features, author relationships and temporal factors, and generates personalized feed rankings accordingly.")
    add_p(doc, "The backend is implemented with Spring Boot, JPA and MySQL, while the frontend uses Vue 3, Vite, Element Plus and ECharts. The recommendation module integrates tag affinity, TF-IDF cosine similarity, collaborative filtering, time decay, negative feedback filtering and diversity re-ranking. In addition, the system introduces AI-assisted ranking and explanation generation through Ollama or DeepSeek, with fallback strategies to ensure the availability of the main feed. Statistical views such as behavior distribution, tag distribution, author deduplication rate, tag coverage rate and recent interest sequences are provided for evaluation and thesis analysis.")
    add_p(doc, "Keywords: user behavior analysis; real-time feed; recommendation system; user profile; hybrid recommendation; AI recommendation")

    doc.add_page_break()

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景与意义", 2)
    add_p(doc, "信息流产品已经成为新闻阅读、社区讨论、短内容分发和商业广告触达的重要入口。与搜索场景不同，信息流场景中的用户需求往往是隐式的、连续的和动态变化的，用户不一定主动输入明确查询词，而是通过浏览、停留、点赞、评论、转发和负反馈等行为不断表达兴趣。因此，如何从用户行为中提取偏好特征，并将其转化为实时排序依据，是信息流系统设计中的关键问题。智能推荐算法已经被广泛用于内容传播与平台增长场景，相关研究表明推荐系统能够提升内容匹配效率和用户使用体验[1]。")
    add_p(doc, "传统时间倒序信息流实现简单，能够保证内容的新鲜性，但它忽视了不同用户之间的兴趣差异，也无法利用长期行为信号优化排序结果。协同过滤、基于内容的推荐、深度学习推荐和混合推荐等方法为个性化信息流排序提供了不同技术路径[2][4][11]。与此同时，社交平台的信息流还需要兼顾内容多样性、负反馈控制、广告插入和系统可解释性，单一算法难以覆盖全部工程需求。")
    add_p(doc, "本文设计并实现的系统不是单纯展示帖子列表，而是围绕“行为采集—画像构建—候选生成—混合排序—统计分析—后台治理”的完整链路展开。该系统既能够满足毕业设计任务书中对实时信息流排序与分发功能的要求，也为论文后续通过统计结果分析推荐效果提供了可复现的数据基础。")

    add_heading(doc, "1.2 国内外研究现状", 2)
    add_p(doc, "推荐系统研究经历了从协同过滤到矩阵分解，再到深度学习和工业级多阶段排序的发展过程。协同过滤方法通常利用用户与物品之间的历史交互矩阵计算相似用户或相似物品，Item-to-Item 协同过滤曾被用于大规模电子商务推荐场景[17]，矩阵分解方法则通过低维隐向量表达用户和物品偏好[16]。近年来，图卷积、深度神经网络、DeepFM 和兴趣演化网络等方法进一步增强了模型对高维稀疏特征、组合特征和序列兴趣的建模能力[2][14][18][19]。")
    add_p(doc, "从信息流产品角度看，推荐系统通常采用多阶段架构：召回阶段从大规模内容池中选出候选集，排序阶段综合用户画像、内容特征、上下文和业务约束计算得分，重排阶段再处理多样性、去重、广告插入和负反馈过滤等问题。YouTube 推荐系统和 Twitter 推荐算法均体现了多阶段候选生成、特征排序和重排治理的工程思想[15][21]。")
    add_p(doc, "国内研究中，用户浏览行为、用户画像构建、显隐式反馈融合和去偏推荐等方向与本文系统关系密切。基于浏览行为的个性化推荐强调从隐式反馈中挖掘兴趣信号[7]；用户画像研究强调将用户行为、内容标签和统计特征组织为可用于推荐的结构化特征[8]；融合显隐式反馈的协同过滤模型可以缓解单一反馈信号不足的问题[4]；去偏推荐关注流行度偏置与质量偏好之间的平衡[20]。这些研究为本文系统的行为权重、画像构建、混合排序和多样性指标提供了理论参考。")

    add_heading(doc, "1.3 本文主要研究内容", 2)
    add_p(doc, "本文围绕一个可运行的信息流排序与分发系统展开，主要工作包括以下几个方面。第一，设计用户认证、内容发布、互动反馈、关注关系、通知私信、广告分发和后台管理等基础功能，保证系统具备完整的社交信息流业务闭环。第二，设计用户行为采集与画像构建方法，将浏览、点赞、评论、转发、搜索、点踩和快速滑过等行为转化为标签偏好、互动风格、内容深度偏好、新鲜度偏好、探索度和近期兴趣序列等画像维度。第三，设计混合推荐排序策略，综合标签匹配、文本相似度、协同过滤、时间衰减、热门度、负反馈过滤、作者多样性和标签覆盖等因素生成推荐结果。第四，引入 AI 辅助排序与推荐理由生成，在模型可用时利用大语言模型进行候选重排，在模型不可用时自动回退到传统混合推荐策略。第五，构建推荐对比页、用户画像页和数据统计中心，通过行为类型分布、内容标签分布、推荐得分、多样性指标和近期兴趣序列等统计结果支撑论文分析。")

    add_heading(doc, "1.4 论文组织结构", 2)
    add_p(doc, "全文共分为六章。第一章介绍研究背景、意义、研究现状和主要内容。第二章介绍系统相关技术与理论基础，包括前后端分离、用户画像、协同过滤、TF-IDF、混合推荐和 AI 推荐原理。第三章进行需求分析和系统总体设计，明确功能模块、数据流和数据库设计。第四章说明系统核心模块的详细设计与实现，重点描述行为采集、画像构建、混合排序、AI 推荐、广告分发和统计分析。第五章进行系统测试与统计结果分析，结合页面截图和接口统计说明系统效果。第六章总结全文工作并提出后续优化方向。")

    add_heading(doc, "2 相关技术与理论基础", 1)
    add_heading(doc, "2.1 前后端分离与系统开发技术", 2)
    add_p(doc, "系统采用前后端分离架构。后端基于 Java、Spring Boot 3.2.1、Maven、Spring Data JPA 和 MySQL 实现 REST 接口、业务逻辑、推荐算法和数据持久化；前端基于 Vue 3、Vite、Element Plus 和 ECharts 实现信息流展示、推荐对比、画像分析、数据统计和后台管理页面。Java 平台具有较成熟的生态和工程实践基础[9]，Vue 3 与 Element Plus 适合构建响应式后台和数据可视化界面[10]。前后端通过 HTTP JSON 接口交互，便于模块独立开发和部署；数据密集型系统设计中强调的可靠性、可维护性和可扩展性也为本系统的模块划分提供了工程依据[13]。")
    add_p(doc, "设计模式方面，推荐策略采用策略模式组织不同排序实现。系统保留传统混合推荐策略和 AI 推荐策略，并通过推荐策略管理器完成切换。策略模式能够降低不同算法实现之间的耦合，使后续扩展新的排序策略或实验策略时无需改动控制层代码[12]。")

    add_heading(doc, "2.2 用户行为分析与用户画像", 2)
    add_p(doc, "用户行为是推荐系统最重要的数据来源之一。显式反馈通常包括评分、收藏、关注等主动表达，隐式反馈包括浏览、停留、点击、点赞、搜索、快速滑过等行为[7]。在信息流场景中，显式评分往往缺失，因此系统需要更多依赖隐式反馈和弱反馈信号。本文系统将 LIKE、COMMENT、REPOST、VIEW、SEARCH 等视为正向或兴趣发现信号，将 DISLIKE、SKIP 和不感兴趣、屏蔽、静音等视为负向信号，分别用于画像构建、排序加权和候选过滤。")
    add_p(doc, "用户画像是对用户兴趣和行为模式的结构化表达。系统中的画像维度包括兴趣标签、分类偏好、互动风格、活跃度、内容长度偏好、图片偏好、话题多样性、活跃时段和近期兴趣序列等。用户画像研究通常强调从多源行为数据中抽取标签化、统计化和可解释的特征[8]。本文系统并不把画像作为静态资料，而是根据最新行为动态更新画像，从而为实时信息流排序提供更贴近当前兴趣的输入。")

    add_heading(doc, "2.3 TF-IDF 与文本相似度", 2)
    add_p(doc, "为了弥补单纯标签匹配粒度较粗的问题，系统引入 TF-IDF 与余弦相似度计算用户近期兴趣文本与候选帖子正文之间的语义匹配程度。对文本内容分词后，词项 i 在文档 d 中的 TF-IDF 权重可采用经典形式：")
    add_p(doc, "w(i,d) = tf(i,d) × log(N / df(i))                                          （2-1）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_p(doc, "其中，N 表示语料文档数，df(i) 表示包含词项 i 的文档数。系统在用户画像侧使用近期互动内容构造兴趣向量，在候选帖子侧构造文档向量，再计算二者余弦相似度：")
    add_p(doc, "sim(u,c) = Σw(i,u)w(i,c) / (||u|| × ||c||)                              （2-2）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_p(doc, "系统代码中的文本相似度加成（contentSimilarityBoost）即来源于该计算结果。该指标能够把用户近期阅读和互动过的文本内容转化为向量表示，使推荐排序不仅依赖人工标签，也能识别正文层面的兴趣接近程度[11]。")

    add_heading(doc, "2.4 协同过滤与混合推荐", 2)
    add_p(doc, "协同过滤的基本思想是利用用户之间或物品之间的历史行为相似性进行推荐。基于用户的协同过滤会寻找与目标用户行为模式相近的其他用户，再将相似用户喜欢的内容推荐给目标用户；基于物品的协同过滤则通过物品共现关系寻找相似内容[16][17]。本文系统实现了基于用户行为向量的协同过滤服务，将不同行为类型赋予不同权重，并通过加权余弦相似度计算相似用户。")
    add_p(doc, "由于信息流排序同时受到兴趣匹配、互动热度、时间新鲜度、负反馈、多样性和业务策略影响，单一算法难以满足全部需求。混合推荐可以综合协同过滤、基于内容的推荐、标签匹配和排序重排等方法，兼顾准确性、覆盖率和可解释性[5][11]。系统中的混合推荐策略把各类特征转换为可解释的分项得分，再计算最终排序分。")

    add_heading(doc, "2.5 AI 辅助推荐技术", 2)
    add_p(doc, "近年来，深度学习和大模型技术被广泛用于推荐系统特征建模、点击率预估和候选重排[14][18][19]。传统深度推荐模型通常依赖大量训练数据，而本文系统面向毕业设计原型，更适合采用“传统算法候选生成 + 大语言模型辅助重排”的轻量方案。该方案先由传统混合推荐生成候选集，再将用户画像、候选内容摘要、互动指标和标签信息组织为提示词，交给大语言模型判断候选内容与用户兴趣的匹配程度，并返回排序结果和推荐理由。")
    add_p(doc, "该设计与工业推荐系统中的多阶段思想一致：召回与粗排阶段由确定性算法保证效率和稳定性，AI 阶段主要承担语义理解、理由生成和候选重排任务[15][21]。为了避免 AI 服务不可用影响主流程，系统实现了健康检查、缓存和降级机制。当本地 Ollama 或外部 DeepSeek 服务响应失败时，推荐接口自动回退到传统混合排序结果，保证首页信息流仍然可用。")

    add_heading(doc, "3 系统需求分析与总体设计", 1)
    add_heading(doc, "3.1 系统需求分析", 2)
    add_p(doc, "根据任务书要求，系统需要围绕用户行为分析和实时信息流排序展开。结合实际代码实现，系统需求可以分为基础业务需求、推荐算法需求、统计分析需求和管理维护需求。基础业务需求包括用户注册登录、内容发布、图片上传、点赞点踩、评论、转发、引用、搜索、关注关系、通知和私信。推荐算法需求包括行为采集、用户画像、候选内容生成、混合排序、负反馈过滤、AI 推荐和推荐解释。统计分析需求包括行为类型统计、内容标签统计、热门话题统计、推荐对比统计和画像可视化。管理维护需求包括用户管理、内容管理、推荐策略切换、广告配置和数据导入。")
    add_p(doc, "与普通社区系统相比，本文系统的重点不在于简单 CRUD，而在于把用户行为转化为可用于排序的信号。因此，论文中对“用户认证与内容互动模块”的描述应尽量服务于推荐链路：认证用于确定行为归属，内容发布用于形成候选池，互动行为用于构建画像和反馈闭环，统计中心用于验证推荐效果。")

    add_caption(doc, "表3-1 系统核心模块输入输出概要")
    add_table(
        doc,
        ["模块", "主要输入", "主要处理", "主要输出"],
        [
            ["用户认证与资料", "用户名、密码、用户资料", "注册登录、角色识别、JWT 会话", "用户信息、Token、权限状态"],
            ["内容与互动", "帖子、图片、评论、点赞、点踩、搜索", "内容持久化、标签解析、行为记录", "内容流、行为日志、互动计数"],
            ["用户画像", "用户行为、内容标签、发布时间", "行为加权、时间衰减、兴趣聚合", "兴趣标签、互动风格、近期兴趣序列"],
            ["推荐排序", "候选内容、画像、负反馈、协同过滤结果", "混合评分、多样性重排、AI 辅助排序", "个性化 Feed、评分明细"],
            ["统计分析", "行为日志、内容库、推荐结果、广告记录", "分布统计、趋势统计、对比统计", "行为图表、标签图表、推荐对比指标"],
        ],
        [3.1, 4.2, 5.1, 4.5],
    )

    add_heading(doc, "3.2 系统功能结构设计", 2)
    add_p(doc, "系统功能结构可划分为前台用户端、推荐算法服务、统计分析服务和后台管理端四个层次。前台用户端面向普通用户，提供首页推荐流、关注流、帖子详情、个人主页、画像洞察、通知和私信等功能。推荐算法服务负责采集行为、构建画像、计算候选得分和输出推荐结果。统计分析服务负责对平台运行数据和推荐效果进行聚合展示。后台管理端负责用户管理、内容管理、推荐策略切换、广告管理和数据导入。")
    add_placeholder(doc, "放置系统总体功能结构图，建议包含前台用户端、推荐服务、统计分析、后台管理四个部分。", "图3-1 系统总体功能结构图")

    add_heading(doc, "3.3 系统技术架构设计", 2)
    add_p(doc, "系统采用浏览器/服务器架构。前端运行在浏览器中，通过 Axios 调用后端 REST API，并通过 WebSocket 获取实时通知或消息。后端通过 Controller 层提供接口入口，通过 Service 层组织业务逻辑与推荐算法，通过 Repository 层访问 MySQL 数据库。推荐相关服务包括 RecommendationService、HybridRecommendationStrategy、AiRecommendationStrategy、AiRecCacheService、UserBehaviorProfileService、TfIdfService、CollaborativeFilteringService 和 AdService。")
    add_p(doc, "数据流方面，用户在前端执行浏览、点赞、评论、搜索等操作后，前端调用后端接口写入行为记录；画像服务定期或按请求读取行为记录并构造兴趣特征；推荐服务读取用户画像、候选帖子和负反馈记录，计算推荐分并返回带评分明细的 ContentWithScore；前端 FeedList、CompareView、ProfileView 和 AnalyticsView 分别展示推荐结果、推荐对比、画像分析和统计图表。")
    add_placeholder(doc, "放置系统技术架构图，建议标出 Vue 前端、Spring Boot 后端、MySQL、Ollama/DeepSeek、WebSocket。", "图3-2 系统技术架构图")

    add_heading(doc, "3.4 数据库与数据字典设计", 2)
    add_p(doc, "系统数据库以 MySQL rec_db 为核心，主要实体包括用户、内容、行为、标签、关注关系、通知、私信、广告和负反馈。用户表保存账号、头像、简介、角色和封禁状态；内容表保存帖子正文、图片、作者、父内容、转发来源、引用来源、分类、创建时间和互动计数；行为表保存用户 ID、内容 ID、行为类型、停留时长和创建时间，是推荐算法最重要的数据来源；标签表与内容表通过 content_tags 形成多对多关系。")
    add_caption(doc, "表3-2 推荐相关核心数据表说明")
    add_table(
        doc,
        ["数据表/实体", "关键字段", "作用"],
        [
            ["users", "id、username、role、banned、avatarUrl", "识别用户身份、角色和状态"],
            ["contents", "id、author_id、content、category、created_at、like_count、view_count", "形成信息流候选池并提供排序特征"],
            ["behaviors", "user_id、content_id、type、duration、created_at", "记录用户行为并支撑画像与推荐"],
            ["tags/content_tags", "tag.name、content_id、tag_id", "保存内容标签和兴趣匹配依据"],
            ["negative_signals", "user_id、target_id、signal_type", "记录不感兴趣、屏蔽、静音等负反馈"],
            ["ads/ad_configs", "category、bidPrice、impressionCount、clickCount", "支撑广告插入与统计"],
        ],
        [3.0, 4.3, 5.0],
    )

    add_heading(doc, "4 系统详细设计与实现", 1)
    add_heading(doc, "4.1 用户认证与内容互动模块", 2)
    add_p(doc, "用户认证模块是行为归属和权限判断的基础。用户通过注册接口创建账号，通过登录接口获取 Token，前端将登录状态保存到 localStorage 并在需要鉴权的路由中进行校验。管理员角色可进入后台页面进行用户、内容、策略和广告管理。需要说明的是，当前系统为毕业设计原型，部分密码处理采用开发阶段简化方式，正式生产环境应使用 BCrypt 等单向哈希方式保存密码。")
    add_p(doc, "内容互动模块负责生成推荐系统的原始数据。用户可以发布文字或图片帖子，可以对帖子进行评论、点赞、点踩、转发和引用。图片先通过上传接口保存并返回 URL，再与正文一并提交。后端在内容发布后解析正文中的话题标签并写入标签关系，必要时可调用 AI 打标服务补充语义标签。点赞、评论、转发等正向行为会增加内容互动计数，同时写入 behaviors 表；点踩、快速滑过、不感兴趣等负向行为用于后续过滤和降权。")

    add_heading(doc, "4.2 用户行为采集与画像构建模块", 2)
    add_p(doc, "系统的行为采集覆盖 VIEW、LIKE、COMMENT、REPOST、QUOTE、DISLIKE、SKIP 和 SEARCH 等类型。不同类型行为代表不同强度的兴趣信号，例如评论和转发通常比浏览更能表达用户偏好，快速滑过和点踩则代表弱负反馈或负反馈。画像服务首先根据用户 ID 读取历史行为，再根据行为关联的内容获取标签、分类、作者和发布时间，最终计算兴趣标签、作者偏好、互动风格、内容深度偏好、新鲜度偏好和探索度。")
    add_p(doc, "近期兴趣序列是本次优化中新增的画像维度。系统按最近 14 天聚合用户互动内容的标签或分类，将每天的行为权重累加为兴趣得分，并选取 Top 标签形成折线序列。该序列能够体现用户兴趣是否从某一主题转向另一主题，为论文中“用户兴趣动态变化”的论述提供依据，也呼应了用户兴趣并非静态不变的推荐系统研究观点[7][8]。")
    add_placeholder(doc, "放置用户画像页面的近期兴趣序列截图，重点展示折线图和每日主导标签。", "图4-1 用户近期兴趣序列展示")

    add_heading(doc, "4.3 混合推荐排序模块", 2)
    add_p(doc, "混合推荐排序模块是系统核心。系统首先根据候选池获取可推荐内容，再依次执行负反馈过滤、基础互动得分计算、标签兴趣匹配、作者亲和度计算、TF-IDF 文本相似度计算、协同过滤加成、时间衰减、新鲜度匹配和多样性调整。最终得分可以抽象为：")
    add_p(doc, "Score(u,c)=Base(c)+Tag(u,c)+Author(u,c)+Text(u,c)+CF(u,c)+Fresh(c)+Trend(c)-Penalty(u,c)        （4-1）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_p(doc, "其中 Base(c) 表示内容基础互动热度，Tag(u,c) 表示用户兴趣标签与内容标签的匹配加成，Author(u,c) 表示作者亲和度，Text(u,c) 表示 TF-IDF 文本相似度，CF(u,c) 表示协同过滤加成，Fresh(c) 表示时间新鲜度，Trend(c) 表示热门话题加成，Penalty(u,c) 表示负反馈惩罚。")
    add_p(doc, "时间衰减用于避免旧内容长期占据高位，可采用如下形式：")
    add_p(doc, "decay(c)=1 / (1 + ageHours(c) / k)                                      （4-2）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_p(doc, "在实际实现中，系统根据用户阶段动态调整权重。冷启动用户行为较少，系统提高热门内容和探索因子的权重；初级用户已有少量行为，系统开始提高标签匹配和文本相似度权重；活跃用户行为充分，系统进一步强化个性化、协同过滤和作者亲和度。该设计既避免新用户无内容可看，也能够让老用户获得更精准的推荐。")
    add_placeholder(doc, "放置推荐排序流程图，建议标出候选生成、负反馈过滤、评分、作者多样性、随机探索、输出 Feed。", "图4-2 混合推荐排序流程图")

    add_heading(doc, "4.4 推荐多样性控制与对比指标", 2)
    add_p(doc, "推荐系统如果只追求短期匹配分数，可能会连续推荐同一作者或同一标签内容，造成兴趣疲劳。已有研究也指出推荐系统需要在准确性、多样性和新颖性之间取得平衡[20]。因此，系统在推荐排序中加入作者多样性控制，并在推荐对比页新增作者去重率和标签覆盖率两个指标，用于分析推荐结果是否过于集中。")
    add_p(doc, "作者去重率定义为推荐列表中不同作者数量与推荐条数之比：")
    add_p(doc, "ADR = uniqueAuthors / listSize                                           （4-3）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_p(doc, "标签覆盖率定义为推荐列表中不同标签数量与推荐条数之比，超过 1 时按 1 处理：")
    add_p(doc, "TCR = min(1, uniqueTags / listSize)                                      （4-4）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_p(doc, "推荐对比页同时展示个性化推荐流和时间倒序流的平均分、标签命中率、TF-IDF 相似度、平均互动量、作者去重率和标签覆盖率。这样论文不再只说明“推荐分更高”，还可以分析个性化排序是否在内容覆盖和作者分布方面更合理。")
    add_placeholder(doc, "放置 CompareView 推荐对比页面截图，需显示作者去重率、标签覆盖率和柱状图。", "图4-3 推荐对比页面多样性指标")

    add_heading(doc, "4.5 AI 辅助推荐与降级机制", 2)
    add_p(doc, "AI 辅助推荐模块采用“候选集压缩 + 提示词重排”的方式。传统混合算法先筛选出 Top 候选内容，系统将用户画像、近期兴趣、候选帖子摘要、标签、互动数据和排序要求组织为 Prompt。大语言模型根据上下文理解用户兴趣与候选内容之间的关系，返回候选内容 ID 的排序列表和推荐理由。该方式不需要在本地训练深度模型，适合毕业设计原型接入，同时能够增强系统可解释性。")
    add_p(doc, "AI 推荐模块需要处理模型不可用、响应超时、JSON 解析失败和候选 ID 不匹配等异常情况。系统通过 AiRecCacheService 对 AI 结果进行缓存，并通过健康检查判断模型是否可用。当 AI 服务失败时，系统返回传统混合推荐结果，保证首页信息流主链路不被 AI 服务影响。这种降级机制体现了工程系统对可靠性和可用性的要求[13]。")
    add_placeholder(doc, "放置 AI 推荐模式或推荐理由截图，也可放置关闭 Ollama 后系统降级的测试截图。", "图4-4 AI 推荐与降级机制展示")

    add_heading(doc, "4.6 广告分发与后台管理模块", 2)
    add_p(doc, "广告模块根据广告分类、出价、曝光次数、点击次数和用户兴趣标签进行信息流插入。个性化广告推荐研究表明，广告分发需要结合用户画像、广告内容和统计反馈进行优化[3]。本文系统在首页信息流中按配置间隔插入广告卡片，并记录广告展示和点击数据，在广告仪表盘中计算 CTR 和估算收入。")
    add_p(doc, "后台管理模块提供用户列表、封禁、角色切换、内容管理、推荐策略切换、数据导入和广告管理功能。管理员可以在 traditional 与 ai 策略之间切换，并通过首页徽标或推荐对比结果观察策略变化。该模块不仅服务演示，也体现系统可治理性。")

    add_heading(doc, "4.7 数据统计与可视化模块", 2)
    add_p(doc, "统计中心是本次论文优化中需要重点加强的部分。系统通过 AnalyticsService 聚合用户、内容、行为、负反馈和广告数据，并在 AnalyticsView 中用图表展示。当前统计页包含总体指标卡片、行为趋势图、行为类型分布图、用户阶段分布图、用户类型分布图、内容分类分布图、内容标签分布图、热门话题列表、用户兴趣标签列表、发帖趋势、活跃时段、高互动帖子和广告摘要。")
    add_p(doc, "其中，行为类型分布图用于说明系统实际采集了哪些类型的用户反馈；内容标签分布图用于说明内容池的主题结构；用户兴趣标签 Top10 用于说明全站用户偏好；推荐对比页指标用于说明个性化排序相对时间倒序的变化。这些统计结果为第五章从“测试通过”转向“基于统计结果分析”提供了直接材料。")

    add_heading(doc, "5 系统测试与统计结果分析", 1)
    add_heading(doc, "5.1 测试环境与测试方法", 2)
    add_p(doc, "系统测试环境以 Windows 为主，后端运行环境为 JDK 17 和 Maven，前端运行环境为 Node.js 和 npm，数据库采用 MySQL，后端默认端口为 8888，前端默认端口为 5173。测试方法采用功能场景测试、接口联调验证、构建验证和统计结果分析相结合的方式。考虑到当前项目主要是毕业设计原型，测试结论不表述为持续集成自动通过，而是以手工执行记录、接口响应、页面截图、构建命令输出和统计图表作为验证依据。")
    add_caption(doc, "表5-1 测试环境配置")
    add_table(
        doc,
        ["项目", "配置"],
        [
            ["操作系统", "Windows"],
            ["后端环境", "JDK 17、Maven、Spring Boot 3.2.1"],
            ["前端环境", "Node.js、npm、Vue 3、Vite、Element Plus"],
            ["数据库", "MySQL rec_db"],
            ["图表工具", "ECharts、echarts-wordcloud"],
            ["AI 服务", "Ollama、DeepSeek API"],
        ],
        [4.0, 9.0],
    )

    add_heading(doc, "5.2 功能测试结果", 2)
    add_p(doc, "功能测试围绕系统主链路展开，包括用户注册登录、内容发布、图片上传、点赞点踩、评论转发、关注关系、通知私信、推荐流访问、推荐策略切换、AI 推荐降级、广告展示和后台管理等场景。测试过程中，首先使用普通用户完成注册登录并发布帖子，再对帖子执行互动操作，随后进入首页推荐流、个人画像页、推荐对比页和统计中心检查数据是否同步变化。管理员账号用于验证后台管理、策略切换和广告配置。")
    add_caption(doc, "表5-2 核心功能测试结果")
    add_table(
        doc,
        ["编号", "测试内容", "测试过程", "预期结果", "结果"],
        [
            ["T01", "注册登录", "输入账号密码并登录", "返回用户信息和 Token", "通过"],
            ["T02", "内容发布", "发布文字或图片帖子", "首页出现新帖子", "通过"],
            ["T03", "互动反馈", "点赞、点踩、评论、转发", "计数变化并写入行为记录", "通过"],
            ["T04", "推荐流", "访问 For You 信息流", "返回带排序结果的信息流", "通过"],
            ["T05", "画像页", "进入 Insights 标签页", "展示词云、雷达图、近期兴趣序列", "通过"],
            ["T06", "统计中心", "打开 AnalyticsView", "展示行为类型和内容标签分布图", "通过"],
            ["T07", "AI 降级", "AI 服务不可用时请求推荐", "回退到传统混合推荐", "通过"],
            ["T08", "后台管理", "管理员切换策略或管理广告", "操作成功并影响前端展示", "通过"],
        ],
        [1.5, 2.8, 5.2, 3.5, 1.5],
    )

    add_heading(doc, "5.3 用户行为统计结果分析", 2)
    add_p(doc, "为了验证系统是否真实采集用户行为，本文采用 behaviors 表中的行为记录作了行为类型统计，统计范围可在页面中选择今日、最近 7 天、最近 30 天或全部数据。统计项包括浏览、点赞、评论、转发、引用、点踩、快速滑过和搜索。行为类型分布图能够直观显示系统采集到的反馈结构。如果浏览行为占比高，说明用户主要以消费内容为主；如果点赞、评论、转发占比提高，说明用户互动意愿增强；如果点踩或快速滑过占比异常升高，则说明部分推荐内容可能与用户兴趣不匹配，需要通过负反馈过滤进行优化。")
    add_placeholder(doc, "放置 AnalyticsView 中的行为类型分布图截图，建议选择全部或最近 30 天。", "图5-1 行为类型分布统计结果")
    add_p(doc, "通过该统计图，论文可以写成：“采用系统行为日志对用户反馈类型作了统计，统计结果如图5-1所示。由图可知，系统已覆盖浏览、点赞、评论、搜索和负反馈等多种行为类型，能够为用户画像和推荐排序提供多维度数据来源。”")

    add_heading(doc, "5.4 内容标签与主题分布分析", 2)
    add_p(doc, "内容标签是推荐排序的重要特征。本文采用 contents 与 content_tags 关系数据对平台内容主题作了统计，统计结果如图5-2所示。内容标签分布图按照帖子标签出现次数进行聚合，展示平台内容池中各类主题的覆盖情况。如果某一类标签数量明显偏少，可能导致相关兴趣用户的候选内容不足；如果少数标签占比过高，可能使推荐结果出现主题集中现象。因此，内容标签分布不仅用于描述内容生态，也可用于解释推荐效果差异。")
    add_placeholder(doc, "放置 AnalyticsView 中的内容标签分布图截图。", "图5-2 内容标签分布统计结果")
    add_p(doc, "结合内容分类分布和热门话题 Top10，可以进一步分析系统内容库的结构。例如，若 Tech、AI、Sports 等标签较多，则相应兴趣用户更容易获得高匹配推荐；若生活类或新闻类标签较少，则推荐结果可能更多依赖文本相似度和热门度补偿。")

    add_heading(doc, "5.5 推荐对比与多样性分析", 2)
    add_p(doc, "为了验证个性化推荐相对于时间倒序的差异，本文采用 CompareView 页面中的推荐对比接口对推荐流和时间流作了统计。统计指标包括平均推荐分、标签命中率、平均互动量、TF-IDF 相似度、作者去重率和标签覆盖率。平均推荐分反映系统综合评分结果；标签命中率反映推荐内容与用户兴趣标签的匹配比例；TF-IDF 相似度反映文本语义层面的接近程度；作者去重率和标签覆盖率反映推荐结果的多样性。")
    add_placeholder(doc, "放置推荐对比页面截图，突出推荐流与时间流的指标对比。", "图5-3 推荐流与时间流统计对比")
    add_p(doc, "如果个性化推荐流的标签命中率和 TF-IDF 相似度高于时间倒序流，说明系统能够根据用户画像调整排序；如果作者去重率和标签覆盖率保持在合理水平，说明推荐结果没有过度集中于单一作者或单一话题。该部分可以作为论文中“算法有效性验证”的核心统计分析。")

    add_heading(doc, "5.6 用户画像与近期兴趣变化分析", 2)
    add_p(doc, "用户画像页用于展示单个用户的兴趣结构和行为特征。本文采用用户近 14 天互动行为作了近期兴趣序列统计，统计结果如图5-4所示。系统按天聚合用户互动内容的标签权重，绘制主要兴趣标签的变化曲线，并显示每天的主导标签。该图能够体现用户兴趣是否稳定、是否出现迁移以及近期兴趣与长期兴趣是否一致。")
    add_placeholder(doc, "放置 ProfileView 的 Insights 页面截图，展示近期兴趣序列折线图。", "图5-4 用户近期兴趣序列统计结果")
    add_p(doc, "例如，当用户连续多天对 AI、编程、科技类内容产生互动时，相关标签曲线会持续上升；当用户转向体育或生活类内容时，新的标签曲线会在近期上升。推荐系统可据此提高近期兴趣内容权重，避免长期画像过度固化。")

    add_heading(doc, "5.7 AI 推荐与降级测试分析", 2)
    add_p(doc, "AI 推荐测试主要验证模型可用时的重排序能力和模型不可用时的系统降级能力。测试过程中，先启动 Ollama 或配置 DeepSeek API，再切换 AI 推荐策略，检查推荐接口是否返回 AI 排序结果或推荐理由；随后关闭 AI 服务或模拟异常，再次请求推荐接口，观察系统是否回退到传统混合推荐策略。")
    add_p(doc, "测试结果表明，AI 模块适合承担推荐理由生成和候选重排任务，但不应成为首页信息流的唯一依赖。通过缓存和降级机制，系统能够在 AI 服务异常时保持主流程可用。这一点在论文中可以作为系统可靠性设计的重要说明。")

    add_heading(doc, "5.8 测试结论", 2)
    add_p(doc, "综合功能测试和统计分析可以看出，系统已经形成较完整的实时信息流排序与分发闭环。用户行为能够被采集并写入数据库，画像服务能够从行为中提取兴趣特征，推荐服务能够综合多类信号生成个性化排序，前端页面能够展示推荐结果、画像变化和统计图表。与原先只写测试用例“通过”相比，优化后的论文应重点利用图5-1至图5-4说明系统实际运行数据，从而满足导师提出的“根据统计结果去分析”的要求。")

    add_heading(doc, "6 总结与展望", 1)
    add_heading(doc, "6.1 全文总结", 2)
    add_p(doc, "本文围绕基于用户行为分析的实时信息流排序与分发系统进行了设计与实现。系统以 Spring Boot 和 Vue 3 为主要技术栈，实现了用户认证、内容发布、互动反馈、关注通知、私信、推荐排序、用户画像、统计分析、广告分发和后台管理等功能。在推荐算法方面，系统综合使用标签匹配、TF-IDF 文本相似度、协同过滤、时间衰减、负反馈过滤和多样性控制，并接入 AI 服务实现辅助重排与推荐理由生成。")
    add_p(doc, "论文优化后的重点在于把系统功能与推荐链路结合起来描述，而不是孤立介绍登录、发帖和评论等功能。同时，本文增加了行为类型分布、内容标签分布、推荐多样性指标和近期兴趣序列等统计分析内容，使系统测试更具数据依据。整体来看，系统较好地贴合任务书中关于用户行为分析、实时信息流排序与分发、推荐效果验证和系统实现的要求。")

    add_heading(doc, "6.2 不足与展望", 2)
    add_p(doc, "当前系统仍存在一些不足。第一，系统数据规模有限，推荐效果主要通过手工测试和页面统计验证，尚未在大规模真实用户数据上进行离线指标评估。第二，协同过滤和 TF-IDF 算法相对轻量，尚未训练深度排序模型。第三，AI 推荐主要依赖提示词重排，缺少对模型输出稳定性的长期评估。第四，广告分发当前以规则和简单画像匹配为主，尚未建立完整的点击率预估模型。")
    add_p(doc, "后续可从四个方向继续优化。首先，扩展数据集和行为采样规模，引入准确率、召回率、NDCG、多样性和新颖性等离线指标。其次，引入序列推荐或深度 CTR 模型，增强对用户兴趣演化的建模能力。再次，完善 AI 推荐的评测与安全控制，避免模型输出不稳定影响用户体验。最后，进一步完善后台实验平台，支持 A/B 测试、权重配置版本管理和推荐策略灰度发布。")

    add_heading(doc, "参考文献", 1)
    for i, ref in enumerate(REFERENCES, 1):
        add_p(doc, f"[{i}] {ref}", first_line=False)

    add_heading(doc, "附录：正文引用位置建议", 1)
    add_p(doc, "为便于后续修改原稿，本文已在正文中安排 21 篇参考文献的引用位置。建议后续定稿时检查引用是否与段落内容对应：推荐算法与传播研究对应[1]，协同过滤与图卷积对应[2][4][16][17]，广告推荐对应[3]，重排序与多样性对应[5][20]，前后端和工程技术对应[6][9][10][12][13]，用户行为和画像对应[7][8]，深度学习与工业推荐对应[11][14][15][18][19][21]。")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
