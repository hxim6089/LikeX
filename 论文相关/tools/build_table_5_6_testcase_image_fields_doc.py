from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = r"D:\TheEnd\qnyproj-main\recommendation-system\论文相关\表5-6_核心测试用例补充表_图示字段版.docx"


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, east="楷体", size=10.5)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_border(cell, edge, sz="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    node = borders.find(qn("w:" + edge))
    if node is None:
        node = OxmlElement("w:" + edge)
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), sz)
    node.set(qn("w:space"), "0")
    node.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=65, start=65, bottom=65, end=65):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(int(w * 567) for w in widths_cm)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 567)))


def fill_cell(cell, text, center=False, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


headers = ["序号", "测试模块", "测试用例", "预期结果", "实际结果", "结论"]
rows = [
    ("1", "用户认证", "输入正确用户名和密码登录", "登录成功，跳转至首页", "与预期一致", "通过"),
    ("2", "权限控制", "普通用户访问后台入口或接口", "前端隐藏入口或后端限制访问", "与预期一致", "通过"),
    ("3", "内容发布", "发布文字帖子并上传图片", "帖子保存成功，图片正常展示", "与预期一致", "通过"),
    ("4", "内容互动", "对帖子评论、点赞、点踩和转发", "互动计数更新，行为记录写入", "与预期一致", "通过"),
    ("5", "社交关系", "关注用户并切换关注流", "关注状态更新，关注流正常展示", "与预期一致", "通过"),
    ("6", "通知私信", "触发点赞通知并发送私信", "生成通知或私信记录，未读状态更新", "与预期一致", "通过"),
    ("7", "行为采集", "执行浏览、互动和搜索操作", "系统记录 VIEW、LIKE、SEARCH 等行为", "与预期一致", "通过"),
    ("8", "推荐排序", "访问 For You 个性化推荐流", "返回结合行为、标签和时间因素的排序结果", "与预期一致", "通过"),
    ("9", "推荐解释", "打开算法对比页面查看评分拆分", "展示推荐分项和推荐流对比指标", "与预期一致", "通过"),
    ("10", "用户画像", "查看用户兴趣画像页面", "显示词云、雷达图、偏好统计和近期兴趣序列", "与预期一致", "通过"),
    ("11", "AI 模块", "执行 AI 打标并切换 AI 推荐", "模型可用时返回 AI 结果，异常时自动降级", "与预期一致", "通过"),
    ("12", "广告分发", "浏览信息流并点击广告", "广告按间隔插入，展示和点击统计更新", "与预期一致", "通过"),
    ("13", "后台管理", "管理员切换策略、管理用户和导入数据", "后台操作成功，相关数据同步更新", "与预期一致", "通过"),
    ("14", "数据统计", "查看行为类型分布和内容标签分布", "统计图表正常展示平台运行数据", "与预期一致", "通过"),
]

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.6)
section.right_margin = Cm(2.4)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.font.size = Pt(10.5)

add_caption(doc, "表 5-6 核心测试用例补充表")
table = doc.add_table(rows=1, cols=6)
set_table_width(table, [1.15, 2.25, 4.05, 4.75, 2.35, 1.65])
repeat_header(table.rows[0])

for idx, text in enumerate(headers):
    clear_cell_borders(table.rows[0].cells[idx])
    fill_cell(table.rows[0].cells[idx], text, center=True, bold=True)

for data in rows:
    row = table.add_row()
    for c in row.cells:
        clear_cell_borders(c)
    for idx, text in enumerate(data):
        fill_cell(row.cells[idx], text, center=(idx in (0, 1, 4, 5)))

for cell in table.rows[0].cells:
    set_cell_border(cell, "top", sz="12")
    set_cell_border(cell, "bottom", sz="8")
for cell in table.rows[-1].cells:
    set_cell_border(cell, "bottom", sz="12")

doc.save(OUT)
print(OUT)
