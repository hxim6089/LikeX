from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader


PDF = Path(r"D:\下载\s11257-023-09379-6.pdf")
OUT = Path(r"D:\TheEnd\qnyproj-main\recommendation-system\论文相关\s11257-023-09379-6_正文提取_Abstract_to_References.docx")


LIGATURES = {
    "\ufb00": "ff",
    "\ufb01": "fi",
    "\ufb02": "fl",
    "\ufb03": "ffi",
    "\ufb04": "ffl",
    "\ufb05": "st",
    "\ufb06": "st",
    "\u00ad": "",
}


def normalize_text(text: str) -> str:
    for src, dst in LIGATURES.items():
        text = text.replace(src, dst)
    fixes = {
        "V ery": "Very",
        "V asiloglou": "Vasiloglou",
        "V .": "V.",
        "P .": "P.",
        "Y u,": "Yu,",
        "Y uan": "Yuan",
        "Y ang": "Yang",
        "R e n": "Ren",
        "Latiﬁ": "Latifi",
        "M.: ": "M.: ",
    }
    for src, dst in fixes.items():
        text = text.replace(src, dst)
    text = text.replace(" \u201c", " \u201c").replace("\u201d)", "\u201d)")
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"\(\s+", "(", text)
    text = re.sub(r"\s+\)", ")", text)
    return text.strip()


def is_header_footer(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if s == "123":
        return True
    if re.match(r"^\d{3}\s+J\.\s*Bauer,\s*D\.\s*Jannach$", s):
        return True
    if re.match(r"^Hybrid session-aware recommendation with feature-based models\s+\d{3}$", s):
        return True
    if s in {
        "B Josef Bauer",
        "josef.b.bauer@gmail.com",
        "Dietmar Jannach",
        "dietmar.jannach@aau.at",
        "1 University of Klagenfurt, Klagenfurt, Austria",
    }:
        return True
    return False


def is_heading(line: str) -> bool:
    s = line.strip()
    if s in {"Abstract", "References", "Appendix", "Acknowledgements"}:
        return True
    if re.match(r"^\d+(\.\d+)*\s+[A-Z][A-Za-z0-9,:;()\-– ]+$", s):
        return True
    if re.match(r"^Table\s+\d+", s):
        return True
    if re.match(r"^Fig\.\s+\d+", s):
        return True
    return False


def ends_sentence(line: str) -> bool:
    s = line.rstrip()
    if not s:
        return False
    if re.search(r"(e\.g|i\.e|et al|Fig|Table|Sec)\.$", s):
        return False
    return s.endswith((".", "?", "!", ".)", ".)", "\u201d", "\u201d)"))


def should_break_paragraph(prev_raw: str, curr: str, in_refs: bool) -> bool:
    curr = curr.strip()
    if not curr:
        return True
    if is_heading(curr):
        return True
    if in_refs:
        if re.match(r"^[A-ZÀ-Ý][A-Za-zÀ-ÿ\-]+,\s+[A-Z]", curr):
            return True
        if re.match(r"^[A-ZÀ-Ý][A-Za-zÀ-ÿ\-]+\s+[A-ZÀ-Ý][A-Za-zÀ-ÿ\-]+,\s+[A-Z]", curr):
            return True
        return False
    if ends_sentence(prev_raw) and len(prev_raw) < 92 and re.match(r"^[A-Z0-9]", curr):
        return True
    return False


def collect_clean_lines() -> list[str]:
    reader = PdfReader(str(PDF))
    lines = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for raw in text.splitlines():
            line = normalize_text(raw)
            if is_header_footer(line):
                continue
            lines.append(line)
    joined = "\n".join(lines)
    start = joined.find("Abstract")
    if start == -1:
        raise RuntimeError("Cannot find Abstract in PDF text.")
    end = joined.find("Publisher\u2019s Note", start)
    if end == -1:
        end = joined.find("Publisher’s Note", start)
    if end == -1:
        end = joined.find("Josef Bauer studied", start)
    if end == -1:
        end = len(joined)
    body = joined[start:end].strip()
    return [normalize_text(x) for x in body.splitlines()]


def reconstruct_paragraphs(lines: list[str]) -> list[str]:
    paragraphs = []
    buf = ""
    prev_raw = ""
    in_refs = False

    def flush():
        nonlocal buf
        if buf.strip():
            paragraphs.append(normalize_text(buf))
        buf = ""

    for line in lines:
        if not line:
            flush()
            prev_raw = ""
            continue
        if line == "References":
            flush()
            paragraphs.append(line)
            in_refs = True
            prev_raw = line
            continue
        if is_heading(line):
            flush()
            paragraphs.append(line)
            prev_raw = line
            continue
        if not buf:
            buf = line
        else:
            if prev_raw.endswith("-") and re.match(r"^[a-z]", line):
                buf = buf[:-1] + line
            elif should_break_paragraph(prev_raw, line, in_refs):
                flush()
                buf = line
            else:
                buf += " " + line
        prev_raw = line
    flush()
    return paragraphs


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if text in {"Abstract", "References", "Appendix", "Acknowledgements"} or re.match(r"^\d+(\.\d+)*\s+", text):
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True)
        return

    if text.startswith("Keywords "):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=12, bold=False)
        return

    if re.match(r"^(Table|Fig\.)\s+\d+", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=11, bold=True)
        return

    if text.startswith("References"):
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True)
        return

    run = p.add_run(text)
    set_run_font(run, size=12, bold=False)


def build_doc(paragraphs: list[str]):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for para in paragraphs:
        add_para(doc, para)

    doc.save(OUT)


if __name__ == "__main__":
    clean_lines = collect_clean_lines()
    paragraphs = reconstruct_paragraphs(clean_lines)
    build_doc(paragraphs)
    print(OUT)
    print(f"paragraphs={len(paragraphs)}")
    print(f"first={paragraphs[0] if paragraphs else ''}")
    print(f"last={paragraphs[-1][:120] if paragraphs else ''}")
