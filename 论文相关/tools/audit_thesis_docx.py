import io
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document


DOC_PATH = Path(r"D:\SHARE\OneDrive\Desktop\毕业论文 初稿07.2.docx")


def para_texts(doc):
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def is_heading(text):
    return bool(
        re.match(r"^第[一二三四五六七八九十0-9]+章", text)
        or re.match(r"^\d+(\.\d+){0,3}\s*", text)
    )


def chapter_of(text):
    m = re.match(r"^第([一二三四五六七八九十0-9]+)章", text)
    if not m:
        return None
    return m.group(0)


def main():
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    doc = Document(str(DOC_PATH))
    texts = para_texts(doc)

    print("FILE", DOC_PATH)
    print("PARAGRAPHS", len(texts))
    print("TABLES", len(doc.tables))
    print("IMAGES", len(doc.inline_shapes))
    print()

    print("OUTLINE")
    headings = []
    for idx, text in enumerate(texts):
        if is_heading(text):
            headings.append((idx, text))
            print(f"{idx:04d}: {text}")
    print()

    print("CHAPTER PARAGRAPH COUNTS")
    current = "front"
    counts = Counter()
    for text in texts:
        ch = chapter_of(text)
        if ch:
            current = ch
        counts[current] += 1
    for key, value in counts.items():
        print(key, value)
    print()

    body_texts = []
    for t in texts:
        if t.strip() == "外文翻译":
            break
        body_texts.append(t)

    print("BODY_PARAGRAPHS_BEFORE_TRANSLATION", len(body_texts))
    print()

    print("FIGURE_CAPTIONS_BODY")
    fig_caps = [t for t in body_texts if re.match(r"^图\s*\d+[-－]\d+", t)]
    for t in fig_caps:
        print(t)
    print("FIG_COUNT", len(fig_caps))
    print()

    fig_nums = [re.match(r"^图\s*(\d+[-－]\d+)", t).group(1).replace("－", "-") for t in fig_caps]
    print("DUP_FIG_NUMS", [k for k, v in Counter(fig_nums).items() if v > 1])
    print()

    print("TABLE_CAPTIONS_BODY")
    table_caps = [t for t in body_texts if re.match(r"^表\s*\d+[-－]\d+", t)]
    for t in table_caps:
        print(t)
    print("TABLE_CAPTION_COUNT", len(table_caps))
    table_nums = [re.match(r"^表\s*(\d+[-－]\d+)", t).group(1).replace("－", "-") for t in table_caps]
    print("DUP_TABLE_NUMS", [k for k, v in Counter(table_nums).items() if v > 1])
    print()

    print("FIG_TABLE_REFERENCE_COUNTS")
    print("如图", sum(t.count("如图") for t in body_texts))
    print("如表", sum(t.count("如表") for t in body_texts))
    print()

    print("CITATION_NUMBERS")
    refs = []
    for t in body_texts:
        refs.extend(int(x) for x in re.findall(r"\[(\d{1,2})\]", t))
    c = Counter(refs)
    for i in range(1, max(c.keys(), default=0) + 1):
        print(f"[{i}] {c[i]}")
    print("TOTAL_CITATIONS", len(refs), "UNIQUE", len(c))
    print()

    print("REFERENCE_LIST_GUESSES")
    in_refs = False
    ref_lines = []
    for t in texts:
        if t.strip() == "参考文献":
            in_refs = True
            continue
        if in_refs and (t.strip() == "致谢" or re.match(r"^附录", t) or t.strip() == "外文翻译"):
            break
        if in_refs:
            ref_lines.append(t)
    for t in ref_lines[:40]:
        print(t)
    print("REF_LINES", len(ref_lines))
    print()

    print("STYLE_RED_FLAGS")
    red_flags = ["这头", "跑", "兜底", "拉起来", "一块用", "不太", "比较方便", "这边", "基本", "可以看到", "像", "啥", "靠"]
    for flag in red_flags:
        hits = [t for t in body_texts if flag in t]
        if hits:
            print(flag, len(hits))
            for h in hits[:3]:
                print("  ", h[:160])
    print()

    print("PLACEHOLDERS")
    for t in body_texts:
        if "预留" in t or "TODO" in t or "待" in t or "截图" in t:
            print(t)


if __name__ == "__main__":
    main()
