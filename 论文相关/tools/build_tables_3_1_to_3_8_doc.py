from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = r"D:\TheEnd\qnyproj-main\recommendation-system\论文相关\表3-1至表3-8_统一格式.docx"


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, east="楷体", size=10.5)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_border(cell, edge, sz="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    node = borders.find(qn("w:" + edge))
    if node is None:
        node = OxmlElement("w:" + edge)
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), sz)
    node.set(qn("w:space"), "0")
    node.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=62, start=70, bottom=62, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(int(w * 567) for w in widths_cm)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 567)))


def fill_cell(cell, text, center=False, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


def add_three_line_table(doc, caption, headers, rows, widths, center_cols):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    repeat_header(table.rows[0])
    for idx, text in enumerate(headers):
        clear_cell_borders(table.rows[0].cells[idx])
        fill_cell(table.rows[0].cells[idx], text, center=True, bold=True)
    for data in rows:
        row = table.add_row()
        for c in row.cells:
            clear_cell_borders(c)
        for idx, text in enumerate(data):
            fill_cell(row.cells[idx], text, center=(idx in center_cols))
    for cell in table.rows[0].cells:
        set_cell_border(cell, "top", sz="12")
        set_cell_border(cell, "bottom", sz="8")
    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", sz="12")
    return table


tables = [
    (
        "表 3-1 核心模块输入输出概要",
        ["模块", "主要输入", "主要输出", "典型异常"],
        [
            ("用户认证与资料", "用户名、密码、资料字段", "用户信息与 Token", "用户名重复、账号封禁"),
            ("内容与互动", "帖子、评论、转发参数", "Content 与计数更新", "上传失败、权限不足"),
            ("行为采集", "行为类型、contentId、时长", "Behavior 记录", "类型非法、未登录"),
            ("推荐排序", "userId、分页", "排序 Feed、评分明细", "候选为空、AI 降级"),
            ("用户画像", "userId", "标签分布、动态权重", "冷启动默认策略"),
            ("通知与私信", "事件、消息体", "通知记录、推送", "连接断开仅落库"),
            ("广告分发", "画像标签、广告配置", "排序广告、展示统计", "无启用广告"),
        ],
        [3.2, 4.2, 4.2, 4.6],
        {0},
    ),
    (
        "表 3-2 用户表设计",
        ["字段", "类型", "说明"],
        [
            ("id", "BIGINT", "用户 ID"),
            ("username", "VARCHAR", "用户名，唯一"),
            ("handle", "VARCHAR", "用户昵称标识，如 @user1"),
            ("avatarUrl", "VARCHAR", "头像地址"),
            ("bio", "VARCHAR", "个人简介"),
            ("password", "VARCHAR", "登录密码"),
            ("role", "VARCHAR", "USER 或 ADMIN"),
            ("banned", "BOOLEAN", "是否封禁"),
            ("customWeights", "TEXT", "用户自定义推荐权重 JSON"),
            ("createdAt", "DATETIME", "创建时间"),
        ],
        [4.4, 3.8, 8.0],
        {0, 1},
    ),
    (
        "表 3-3 内容表设计",
        ["字段", "类型", "说明"],
        [
            ("id", "BIGINT", "帖子 ID"),
            ("author", "BIGINT", "作者 ID"),
            ("parentContent", "BIGINT", "父评论 ID"),
            ("content", "TEXT", "正文内容"),
            ("imageUrl", "VARCHAR", "图片地址"),
            ("category", "VARCHAR", "内容分类"),
            ("viewCount", "INT", "浏览数"),
            ("likeCount", "INT", "点赞数"),
            ("commentCount", "INT", "评论数"),
            ("dislikeCount", "INT", "点踩数"),
            ("repostCount", "INT", "转发或引用数"),
            ("createdAt", "DATETIME", "发布时间"),
        ],
        [4.4, 3.8, 8.0],
        {0, 1},
    ),
    (
        "表 3-4 行为表设计",
        ["字段", "类型", "说明"],
        [
            ("id", "BIGINT", "行为 ID"),
            ("userId", "BIGINT", "用户 ID"),
            ("contentId", "BIGINT", "内容 ID，搜索等行为可为空"),
            ("type", "VARCHAR", "VIEW、LIKE、DISLIKE、SKIP、SEARCH 等"),
            ("duration", "INT", "浏览停留时长"),
            ("createdAt", "DATETIME", "行为发生时间"),
        ],
        [4.4, 3.8, 8.0],
        {0, 1},
    ),
    (
        "表 3-5 关注表设计概要",
        ["字段", "类型", "说明"],
        [
            ("id", "BIGINT", "主键"),
            ("followerId", "BIGINT", "关注者用户 ID，逻辑外键指向 users.id"),
            ("followeeId", "BIGINT", "被关注者用户 ID"),
            ("createdAt", "DATETIME", "关注时间，可用于推荐中新关注加权"),
        ],
        [4.4, 3.8, 8.0],
        {0, 1},
    ),
    (
        "表 3-6 标签实体与内容标签关联（逻辑结构）",
        ["对象", "说明"],
        [
            ("tags", "标签表：id、name（唯一）、可选统计字段。"),
            ("content_tags", "关联表：contentId、tagId，联合唯一约束防止重复绑定。"),
        ],
        [4.2, 12.0],
        {0},
    ),
    (
        "表 3-7 通知表设计概要",
        ["字段", "类型", "说明"],
        [
            ("id", "BIGINT", "主键"),
            ("recipientId", "BIGINT", "接收者"),
            ("actorId", "BIGINT", "触发者"),
            ("type", "VARCHAR", "LIKE、COMMENT、FOLLOW、REPOST、QUOTE 等"),
            ("entityId", "BIGINT", "关联帖子或内容 ID"),
            ("isRead", "BOOLEAN", "是否已读"),
            ("createdAt", "DATETIME", "通知时间"),
        ],
        [4.4, 3.8, 8.0],
        {0, 1},
    ),
    (
        "表 3-8 广告表设计概要",
        ["字段", "类型", "说明"],
        [
            ("id", "BIGINT", "主键"),
            ("title、description、imageUrl", "VARCHAR/TEXT", "展示文案与素材"),
            ("targetUrl", "VARCHAR", "点击跳转地址"),
            ("advertiser", "VARCHAR", "广告主标识"),
            ("targetTags、category", "VARCHAR", "定向标签与类别"),
            ("bidPrice", "DOUBLE", "出价，用于排序加权"),
            ("impressionCount、clickCount", "INT", "展示与点击计数"),
            ("active", "BOOLEAN", "是否启用投放"),
        ],
        [5.2, 3.5, 7.5],
        {0, 1},
    ),
]


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.6)
section.right_margin = Cm(2.4)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.font.size = Pt(10.5)

for idx, spec in enumerate(tables):
    add_three_line_table(doc, *spec)
    if idx != len(tables) - 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)

doc.save(OUT)
print(OUT)
