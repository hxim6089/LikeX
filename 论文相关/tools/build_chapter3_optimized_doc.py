from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\TheEnd\qnyproj-main\recommendation-system\论文相关")
OUT = ROOT / "第3章_系统总体设计_优化版_补充图表引用.docx"


IMAGES = {
    "arch": ROOT / "任务书技术方案图_优化版_大字加粗_高清.png",
    "func": ROOT / "图3-2_系统功能模块结构图_白底版.png",
    "er": ROOT / "图3-3_系统总体ER图_论文简略版_终版.png",
    "persona": ROOT / "图3-4_用户画像与动态权重流程图_无说明版.png",
    "rec": ROOT / "图3-5_推荐排序与信息流分发流程图.png",
}


def set_run_font(run, size=10.5, bold=False, italic=False, east="宋体", west="Times New Roman"):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_para_font(paragraph, size=10.5, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_run_font(run, 10.5)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(0)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, 15, bold=True, east="黑体")
    elif level == 1:
        run = p.add_run(text)
        set_run_font(run, 14, bold=True, east="黑体")
    else:
        run = p.add_run(text)
        set_run_font(run, 12, bold=True, east="黑体")
    return p


def add_caption(doc, text, is_table=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(3 if is_table else 0)
    p.paragraph_format.space_after = Pt(3 if is_table else 6)
    run = p.add_run(text)
    set_run_font(run, 10.5, east="楷体")
    return p


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 10.5, bold=bold)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn("w:{}".format(key)), str(value))


def apply_three_line_table(table):
    rows = table.rows
    for row in rows:
        for cell in row.cells:
            set_cell_border(cell,
                            top={"val": "nil"},
                            left={"val": "nil"},
                            bottom={"val": "nil"},
                            right={"val": "nil"})
    for cell in rows[0].cells:
        set_cell_border(cell,
                        top={"val": "single", "sz": "12", "color": "000000"},
                        bottom={"val": "single", "sz": "6", "color": "000000"},
                        left={"val": "nil"},
                        right={"val": "nil"})
    for cell in rows[-1].cells:
        set_cell_border(cell,
                        bottom={"val": "single", "sz": "12", "color": "000000"},
                        left={"val": "nil"},
                        right={"val": "nil"})


def add_table(doc, caption, headers, rows, widths=None):
    add_caption(doc, caption, is_table=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            hdr[idx].width = Cm(widths[idx])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, item in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], item, align=align)
            if widths:
                cells[idx].width = Cm(widths[idx])
    apply_three_line_table(table)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    return table


def add_image(doc, path, caption, width_cm=14.5):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, caption)
    else:
        add_para(doc, f"【图预留：{caption}】", first_line=False)
        add_caption(doc, caption)


def add_formula(doc, formula, number):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.3)
    table.columns[1].width = Cm(2.0)
    left, right = table.rows[0].cells
    set_cell_text(left, formula, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(right, number, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"},
                        bottom={"val": "nil"}, right={"val": "nil"})


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    add_heading(doc, "第3章  系统总体设计", 0)
    add_para(doc, "本章在需求分析和技术可行性分析的基础上，对基于用户行为分析的实时信息流排序与分发系统进行总体设计。为避免总体设计与第四章详细实现内容重复，本章重点说明系统的总体架构、功能结构、核心推荐算法、数据库结构以及支撑系统落地的关键技术。通过本章设计，可以明确系统各模块之间的协作关系、数据在前后端和数据库之间的流转方式，以及推荐结果从行为采集到排序分发的形成过程。")

    add_heading(doc, "3.1 系统体系架构设计", 1)
    add_para(doc, "系统采用 B/S 架构和前后端分离开发模式。浏览器端运行 Vue 3 单页应用，主要负责页面展示、用户操作、状态维护和图表渲染；后端采用 Spring Boot 提供 RESTful API 与 WebSocket 服务，负责用户认证、内容管理、行为记录、推荐排序、AI 调用、广告分发和后台管理等业务逻辑；数据库采用 MySQL 存储用户、内容、行为、标签、关注、通知、私信、负反馈和广告等业务数据。")
    add_para(doc, "从工程结构看，后端项目以 controller、service、repository、model、dto 和 config 等包组织代码。controller 层作为接口入口接收前端请求，service 层承载业务流程和算法计算，repository 层负责 JPA 数据访问，model 层定义实体映射，dto 层封装接口传输对象，config 层集中管理跨域、WebSocket、文件访问映射和 AI 服务参数等配置。前端项目则以 views、components 和 router 为主要结构，页面级组件负责业务视图，可复用组件负责信息流卡片、权重调节、画像展示和图表展示，路由层负责页面跳转与权限控制。")
    add_para(doc, "系统整体数据流可以概括为四个层次：数据层负责保存内容、行为和社交关系；特征层根据行为记录生成用户画像、标签偏好和文本向量；排序层根据传统混合推荐策略或 AI 推荐策略完成候选内容排序；展示层将普通内容、推荐解释、统计图表和广告卡片呈现给用户。该分层方式能够使页面交互、业务处理、算法计算和数据存储之间保持较清晰的边界。")
    add_para(doc, "系统总体架构如图3-1所示。")
    add_image(doc, IMAGES["arch"], "图3-1 系统总体架构图", width_cm=14.2)

    add_heading(doc, "3.2 系统功能结构设计", 1)
    add_para(doc, "系统功能按业务职责划分为用户端业务模块、推荐算法模块、AI 智能模块、广告模块、管理后台模块和数据导入模块等部分。用户端业务模块覆盖注册登录、内容发布、评论转发、点赞点踩、关注、通知和私信等基础社交功能；推荐算法模块负责行为采集、用户画像、候选内容召回、综合评分、重排和负反馈过滤；AI 智能模块用于内容打标、AI 推荐辅助排序和推荐理由生成；广告模块负责广告配置、画像匹配、信息流插入以及曝光点击统计；管理后台模块负责用户管理、策略切换、系统统计和内容治理；数据导入模块负责公开数据集或内置种子内容的导入与补齐。")
    add_para(doc, "这些模块并不是相互独立运行，而是通过 REST API、WebSocket 推送和共用数据库表协同工作。普通用户的浏览、点赞、评论和搜索等行为会写入行为表，并进一步影响画像生成和推荐排序；关注关系既用于关注流展示，也为 In-Network 候选内容召回提供来源；AI 打标和数据导入能够补充内容标签与候选内容池，降低冷启动阶段内容不足带来的影响；广告分发则复用用户画像标签完成广告匹配，但与普通内容推荐保持相对独立的排序目标。")
    add_para(doc, "系统功能模块结构如图3-2所示。")
    add_image(doc, IMAGES["func"], "图3-2 系统功能模块结构图", width_cm=14.0)
    add_para(doc, "为进一步说明各模块在输入、输出和设计职责上的差异，核心功能模块的输入输出概要如表3-1所示。")
    add_table(
        doc,
        "表3-1 核心功能模块输入输出概要",
        ["模块", "主要输入", "主要输出", "设计说明"],
        [
            ["用户端业务模块", "账号、内容、互动请求", "用户信息、帖子、通知、私信", "承担普通用户主要操作流程，是行为数据产生入口。"],
            ["推荐算法模块", "行为记录、画像、候选内容、负反馈", "推荐列表、评分明细、推荐解释", "负责个性化信息流排序，是系统核心模块。"],
            ["AI 智能模块", "画像摘要、候选内容摘要、待打标文本", "AI 排序、推荐理由、语义标签", "通过本地 Ollama 模型增强语义理解能力。"],
            ["广告模块", "广告配置、画像标签、曝光点击记录", "广告卡片、CTR、统计指标", "在信息流中按规则插入广告并记录效果。"],
            ["管理后台模块", "管理员操作、系统统计请求", "策略状态、用户列表、统计图表", "提供系统治理、策略切换和运行监控入口。"],
            ["数据导入模块", "公开数据集、内置种子内容", "内容库、标签数据、候选池补充", "用于扩展演示数据和冷启动内容来源。"],
        ],
        widths=[3.0, 4.0, 4.0, 5.0],
    )

    add_heading(doc, "3.3 系统核心算法设计", 1)
    add_para(doc, "推荐排序模块是系统实现个性化信息流分发的核心。系统在生成 For You 信息流时，不再单纯按照发布时间倒序展示内容，而是综合考虑用户行为、内容特征、社交关系、文本相似度、协同过滤、时间因素、负反馈和随机探索机制，对候选内容进行综合排序。整体流程包括候选召回、过滤处理、特征计算、综合评分、重排调整和结果返回。")

    add_heading(doc, "3.3.1 推荐策略总体设计", 2)
    add_para(doc, "为了便于推荐算法切换和扩展，后端采用策略模式组织推荐逻辑。系统定义统一的 RecommendationStrategy 接口，传统混合推荐策略 HybridRecommendationStrategy 和 AI 推荐策略 AiRecommendationStrategy 均实现该接口。RecommendationStrategyManager 维护当前启用的策略标识，并对外提供统一的推荐调用入口。RecommendationService 在完成候选内容池组装后，将候选集合交给策略管理器，由当前启用的策略完成排序计算。")
    add_para(doc, "这种设计使推荐业务流程与具体排序算法解耦。管理员可以在后台切换 traditional 与 ai 两种策略，系统无需修改信息流主流程即可观察不同策略下的推荐效果。若后续需要增加其他算法策略，只需实现统一接口并注册到策略管理器中，能够降低推荐模块扩展成本。")

    add_heading(doc, "3.3.2 候选内容召回与过滤", 2)
    add_para(doc, "候选内容召回采用双源候选机制，即将候选内容分为 In-Network 和 Out-of-Network 两类。In-Network 候选来自当前用户已关注作者发布的内容，能够保留社交关系带来的稳定信息来源；Out-of-Network 候选来自全站公开内容池，用于帮助用户发现尚未关注但可能感兴趣的作者和话题。两类候选合并后，系统根据内容 ID 去重，形成统一候选集合。")
    add_para(doc, "候选过滤阶段主要处理负反馈信号。系统通过 negative_signals 表记录不感兴趣、屏蔽作者和静音作者等行为。若用户屏蔽或静音某一作者，该作者发布的内容会在候选阶段被过滤；若用户对某条内容标记不感兴趣，该内容会被降低展示优先级或直接剔除。将负反馈过滤放在评分之前，可以减少无效内容参与 TF-IDF、协同过滤等后续计算，也能提高推荐结果与用户真实偏好的匹配程度。")

    add_heading(doc, "3.3.3 多因子评分模型", 2)
    add_para(doc, "传统混合推荐策略会对每条候选内容计算多个评分分项，并将这些分项组合为最终推荐得分。主要评分项包括多行为加权互动分、标签匹配分、文本相似度加成、协同过滤加成、热门话题加成、时间衰减因子、作者偏好加成、内容深度匹配加成和新鲜度偏好加成等。不同用户阶段对应不同动态权重，行为较少的冷启动用户更多依赖热门内容和探索因子，活跃用户则更多依赖个性化画像、协同过滤和作者偏好。")
    add_para(doc, "在文本相似度计算方面，系统采用 TF-IDF 方法表示内容文本特征。结合项目中 TfIdfService 的实现，词项 i 在文档 d 中的权重可表示为：")
    add_formula(doc, "wᵢ,ᵈ = tf(i,d) × (log(N / (df(i)+1)) + 1)", "（3-1）")
    add_para(doc, "式中，N 表示语料库文档总数，df(i) 表示包含词项 i 的文档数量，tf(i,d) 表示词项 i 在文档 d 中出现的频率。系统基于用户已点赞内容构造用户兴趣向量，基于候选帖子正文构造内容向量，再利用余弦相似度计算二者之间的语义接近程度：")
    add_formula(doc, "sim(u,c) = (Vᵤ · V꜀) / (|Vᵤ| × |V꜀|)", "（3-2）")
    add_para(doc, "其中，Vᵤ 表示用户兴趣向量，V꜀ 表示候选内容向量。该相似度作为文本相似度加成（tfidfSimilarityBoost）的主要来源，用于弥补单纯依靠标签匹配粒度较粗的问题。")
    add_para(doc, "协同过滤加成由 CollaborativeFilteringService 计算。系统基于用户—内容交互矩阵，将转发、评论、点赞和浏览等行为映射为不同权重，并结合行为发生时间进行衰减，再通过加权余弦相似度查找相似用户。如果候选内容出现在相似用户偏好集合中，则获得协同过滤加成（collaborativeBoost）。该方法能够在标签较少或新话题出现时，利用其他用户的行为记录补充兴趣判断。")
    add_para(doc, "综合上述因素后，候选内容最终推荐得分可概括为：")
    add_formula(doc, "finalScore = E/Dₜ + Btag + Btfidf + Bcf + Btrend + Bauthor + Bdepth + Bfresh + ε", "（3-3）")
    add_para(doc, "式中，E 表示多行为加权互动分，Dₜ 表示时间衰减因子，Btag 表示标签匹配加成，Btfidf 表示文本相似度加成，Bcf 表示协同过滤加成，Btrend 表示热门话题加成，Bauthor 表示作者偏好加成，Bdepth 表示内容深度匹配加成，Bfresh 表示新鲜度偏好加成，ε 表示随机扰动项。各评分项会结合 UserBehaviorProfileService 输出的动态权重集合 DynamicWeights 进行组合，使不同用户在排序时体现不同兴趣侧重点。")
    add_para(doc, "根据行为数量划分的用户阶段与推荐策略关系如表3-2所示。")
    add_table(
        doc,
        "表3-2 用户阶段与推荐策略",
        ["用户阶段", "判定条件", "推荐策略特点"],
        [
            ["冷启动用户", "行为数少于 10", "更多依赖热门内容、新鲜内容和探索因子。"],
            ["初级用户", "行为数 10 到 49", "热门内容与个性化标签匹配混合排序。"],
            ["活跃用户", "行为数不少于 50", "更多依赖个性化画像、协同过滤和作者偏好。"],
        ],
        widths=[3.5, 4.0, 8.0],
    )

    add_heading(doc, "3.3.4 重排与探索机制", 2)
    add_para(doc, "在完成综合评分后，系统还会对初步排序结果进行重排调整。首先，系统会对用户已经产生过明显反馈的内容进行重复曝光惩罚。例如，已点踩内容会受到较强惩罚，已点赞或已浏览内容则适当降低再次出现的优先级。这样既可以减少用户已经明确不感兴趣内容的重复展示，也能为新内容保留更多曝光机会。")
    add_para(doc, "其次，系统引入作者多样性约束。若同一作者在候选列表中连续出现较多内容，系统会降低后续同作者内容的排序位置，避免推荐列表被单一作者占据。最后，系统加入随机探索机制，使部分分数略低但仍具有潜在价值的内容获得展示机会。加权随机采样的思想可表示为：")
    add_formula(doc, "P(cᵢ) = score(cᵢ) / Σ score(cⱼ)", "（3-4）")
    add_para(doc, "其中，P(cᵢ) 表示候选内容 cᵢ 被选中的概率，score(cᵢ) 表示该候选内容的推荐得分。通过这种方式，系统能够在利用高分内容和探索新内容之间保持一定平衡。")

    add_heading(doc, "3.3.5 AI 推荐与降级机制", 2)
    add_para(doc, "除传统混合推荐策略外，系统还设计了 AI 推荐作为智能增强功能。本文中的 AI 推荐并不是重新训练大规模推荐模型，而是在候选召回、负反馈过滤和基础排序完成后，借助本地大语言模型对部分候选内容进行语义理解、辅助重排序并生成推荐理由。它不替代传统推荐链路，而是作为候选内容语义判断和推荐解释生成的补充。")
    add_para(doc, "AI 推荐功能主要依赖本地部署的 Ollama 服务。后端在配置文件中设置模型服务地址、推荐模型名称、温度参数和输出长度，并由 AiRecommendationStrategy 统一完成 Prompt 构造、HTTP 调用、JSON 解析和异常处理。进入 AI 推荐流程后，系统会整理用户画像摘要和候选内容摘要，其中画像摘要包括用户阶段、兴趣话题、互动风格、内容深度偏好和新鲜度偏好，候选内容摘要包括内容 id、正文片段、标签以及点赞、评论、转发等互动统计。随后系统要求模型按固定 JSON 格式返回 ranking 和 reasons 字段，ranking 表示候选内容排序，reasons 表示推荐理由。")
    add_para(doc, "考虑到本地模型推理耗时较长且输出格式可能不稳定，系统在 AI 推荐中加入缓存、异步计算和降级处理。AI 模式下，如果当前用户存在可用缓存结果，系统直接返回缓存内容；如果缓存未命中，系统优先返回传统混合推荐结果，同时在后台触发 AI 推荐计算并更新缓存。当 Ollama 服务不可用、请求超时、返回内容为空或 JSON 格式不符合要求时，系统自动回退到传统混合推荐策略；若传统排序所需数据也不足，则按基础互动量进行兜底排序。该机制保证 AI 模块异常时，信息流主流程仍可正常运行。")

    add_heading(doc, "3.3.6 推荐流水线与复杂度分析", 2)
    add_para(doc, "一次 For You 请求的推荐流水线可以概括为身份解析、候选召回、负反馈过滤、画像读取、特征计算、排序重排、随机探索和分页返回几个阶段。设候选内容数量为 N，内容平均标签数为 T，文本向量维度为 V，则候选过滤和标签匹配通常为 O(N×T)，TF-IDF 相似度计算约为 O(N×V)，最终排序为 O(N log N)。在实际系统中，候选池规模会通过分页、预筛选和 AI 输入数量限制进行控制，因此该复杂度能够满足毕业设计演示和中小规模数据运行需要。")
    add_para(doc, "从系统设计角度看，推荐流水线的关键不只是排序公式本身，还包括候选池规模控制、负反馈提前过滤、画像计算与推荐读取分离、AI 模型调用降级等工程措施。这些措施能够减少实时推荐接口的计算压力，也提高了信息流请求在异常情况下的稳定性。")
    add_para(doc, "推荐排序与信息流分发的整体流程如图3-3所示。")
    add_image(doc, IMAGES["rec"], "图3-3 推荐排序与信息流分发流程图", width_cm=14.2)

    add_heading(doc, "3.4 系统数据库设计", 1)
    add_para(doc, "数据库设计围绕用户、内容、行为、社交关系、推荐辅助数据和广告数据展开。系统核心实体包括用户、内容、行为、标签、关注、通知、私信、负反馈、广告和广告配置等。用户与内容之间是一对多关系，一个用户可以发布多条内容；用户与行为之间也是一对多关系，用户在浏览、点赞、评论、搜索等过程中会产生多条行为记录；用户之间通过关注表形成有向社交关系，为关注流和 In-Network 候选召回提供依据。")
    add_para(doc, "内容实体是信息流展示和推荐排序的核心数据对象。内容可以关联多个标签，也可以触发多条通知记录；在评论、转发和引用等场景下，内容还可以作为其他内容的目标对象。行为表记录用户与内容或搜索意图之间的交互，为画像构建和推荐排序提供基础信号。负反馈表记录不感兴趣、屏蔽和静音等显式反馈，用于候选过滤和排序惩罚。广告相关表则保存广告素材、定向标签、展示次数、点击次数和投放配置，用于完成信息流广告插入和统计分析。")
    add_para(doc, "系统主要实体之间的关系如图3-4所示。")
    add_image(doc, IMAGES["er"], "图3-4 系统总体 E-R 图", width_cm=14.4)
    add_para(doc, "系统主要数据表及其作用如表3-3所示。")
    add_table(
        doc,
        "表3-3 主要数据表设计概要",
        ["数据表", "关键字段", "主要作用"],
        [
            ["user", "id、username、password、role、banned、avatarUrl、customWeights", "保存用户基础信息、角色权限和自定义推荐权重。"],
            ["content", "id、author_id、content、imageUrl、createdAt、viewCount、likeCount、commentCount、repostCount", "保存帖子、评论、转发和引用内容，是信息流排序的核心对象。"],
            ["behavior", "id、user_id、content_id、type、duration、createdAt", "记录浏览、点赞、评论、转发、搜索、点踩和跳过等行为。"],
            ["tag / content_tags", "tag.id、tag.name、content_id、tag_id", "维护内容与标签之间的多对多关系，支持画像和标签匹配。"],
            ["follow", "id、follower_id、followee_id、createdAt", "维护用户之间的有向关注关系，支持关注流和社交召回。"],
            ["notification", "id、recipient_id、actor_id、type、entity_id、isRead", "保存点赞、评论、关注等事件触发的通知记录。"],
            ["message", "id、sender_id、recipient_id、content、isRead、createdAt", "保存用户之间的私信消息和未读状态。"],
            ["negative_signals", "id、user_id、target_type、target_id、signal_type、created_at", "保存不感兴趣、屏蔽和静音等显式负反馈。"],
            ["ad / ad_config", "id、title、targetTags、impressions、clicks、enabled、frequency", "保存广告素材、投放配置以及展示点击统计。"],
        ],
        widths=[3.2, 6.3, 6.0],
    )

    add_heading(doc, "3.5 实现系统的关键技术", 1)
    add_para(doc, "在完成系统架构、功能模块、核心算法和数据库设计后，本节进一步说明系统实现过程中采用的关键技术及其在项目代码中的落地方式。本节不再展开单个功能模块的界面和流程，而是从工程实现角度归纳前后端交互、行为建模、推荐计算、实时通信、AI 服务接入和数据可视化等支撑技术。")

    add_heading(doc, "3.5.1 前后端分离与 REST 接口技术", 2)
    add_para(doc, "系统采用前后端分离方式实现。前端通过 Axios 向后端 REST 接口发起请求，后端 controller 层按照业务边界提供接口，如认证接口、内容接口、行为接口、推荐对比接口、用户画像接口和后台管理接口。接口返回数据不直接暴露 JPA 实体，而是通过 ContentDTO、ScoreBreakdown、PersonaResponse 等 DTO 封装输出字段，使页面展示字段和数据库实体字段保持分离。")
    add_para(doc, "这种接口设计能够降低前端对数据库表结构的耦合。以推荐流为例，前端只需要请求信息流接口并接收内容列表、推荐得分、评分拆分和策略标识；后端则在 service 层完成候选召回、行为统计、画像读取和策略排序。若后续数据库字段或推荐评分项发生变化，只需调整 DTO 组装逻辑，前端页面结构不必大范围修改。")
    add_para(doc, "系统主要 REST 接口约定如表3-4所示。")
    add_table(
        doc,
        "表3-4 主要 REST 接口一览（节选）",
        ["路径", "方法", "说明"],
        [
            ["/api/auth/register", "POST", "注册并签发 Token。"],
            ["/api/auth/login", "POST", "登录校验与签发 Token。"],
            ["/api/content/feed", "GET", "返回个性化或时间序信息流。"],
            ["/api/content/publish", "POST", "发布文字或图片内容。"],
            ["/api/behavior/view", "POST", "记录浏览行为和停留时长。"],
            ["/api/search", "GET", "综合搜索，并可记录 SEARCH 行为。"],
            ["/api/compare/feed", "GET", "返回推荐流与时间流对比数据。"],
            ["/api/admin/rec-strategy", "GET/PUT", "查询或切换推荐策略。"],
            ["/api/user/{id}/persona", "GET", "返回用户画像信息。"],
        ],
        widths=[5.2, 2.5, 8.0],
    )

    add_heading(doc, "3.5.2 用户行为采集与画像建模技术", 2)
    add_para(doc, "用户画像模块由行为数据驱动生成。系统读取用户历史行为记录后，查询相关内容和标签，并根据行为类型赋予不同权重。转发和评论代表较强兴趣，点赞代表正向兴趣，浏览代表弱兴趣，点踩、跳过和屏蔽代表负向兴趣；同时结合行为发生时间进行衰减，使近期行为对画像影响更大。搜索行为用于刻画用户主动检索意图，浏览行为附带的停留时长则用于推断浅阅读或深度阅读倾向。")
    add_para(doc, "画像结果包含兴趣标签分布、作者偏好、互动风格、内容深度偏好、新鲜度偏好和用户阶段等信息，并进一步映射为 DynamicWeights 动态权重集合。推荐策略在计算最终得分时，会根据这些权重调整标签匹配、协同过滤、文本相似度、新鲜度和探索因子的占比，从而使排序结果随用户行为积累逐步从通用推荐过渡到个性化推荐。")
    add_para(doc, "用户画像与动态权重生成流程如图3-5所示。")
    add_image(doc, IMAGES["persona"], "图3-5 用户画像与动态权重生成流程图", width_cm=14.2)

    add_heading(doc, "3.5.3 TF-IDF、协同过滤与混合排序技术", 2)
    add_para(doc, "文本相似度和协同过滤是系统推荐排序中的两类重要技术。TF-IDF 服务负责把用户已点赞内容和候选内容转换为向量，并通过余弦相似度计算文本匹配程度；协同过滤服务负责从用户行为矩阵中发现相似用户，并根据相似用户偏好的内容给候选帖子增加协同过滤分。二者分别从内容语义和用户群体行为两个角度补充标签匹配的不足。")
    add_para(doc, "混合排序策略并不依赖单一算法，而是把互动质量、标签匹配、文本相似度、协同过滤、热门话题、作者偏好、内容深度、新鲜度、时间衰减和随机探索等因素组合起来。为了增强推荐解释性，系统将各评分项写入 ScoreBreakdown 对象，前端可在推荐详情或算法对比页面中展示评分来源，帮助观察不同推荐因子对最终排序的影响。")

    add_heading(doc, "3.5.4 WebSocket 实时通信技术", 2)
    add_para(doc, "通知和私信功能通过 WebSocket 实现实时推送。后端启用 WebSocket 服务后，前端连接 /ws 地址并订阅个人消息通道。点赞、评论、关注、转发和私信等事件发生时，系统先将通知或消息写入数据库，再通过 SimpMessagingTemplate 向在线用户推送消息。若用户不在线，数据库记录仍然保留，用户再次进入系统后可以通过通知列表或私信列表补偿查看。")
    add_para(doc, "这种设计兼顾了实时性和可靠性。WebSocket 用于减少轮询请求，提高在线用户的消息到达速度；数据库持久化用于保证离线场景下的消息不丢失。对于毕业设计系统而言，该方案实现复杂度适中，能够满足通知提醒、未读状态和私信会话等基础实时通信需求。")

    add_heading(doc, "3.5.5 AI 服务接入、缓存与降级技术", 2)
    add_para(doc, "AI 服务接入采用后端统一封装方式，前端不直接调用模型服务。后端根据用户画像摘要和候选内容摘要生成 Prompt，通过 Ollama 本地接口请求模型生成结构化排序结果和推荐理由。为避免模型返回不可解析文本，Prompt 中要求模型输出固定 JSON 结构，后端再对 ranking 和 reasons 字段进行校验、过滤和补齐。")
    add_para(doc, "由于本地大模型推理存在响应慢和格式不稳定的问题，系统在策略管理器中加入缓存与降级处理。AI 缓存命中时直接返回已有结果，缓存未命中时优先返回传统混合推荐并异步触发 AI 计算。当模型异常、超时或返回格式错误时，系统回退到传统推荐策略，保证信息流主链路不被 AI 服务阻塞。")

    add_heading(doc, "3.5.6 数据统计与可视化技术", 2)
    add_para(doc, "系统统计与可视化主要由后端统计接口和前端 ECharts 图表共同完成。后端 AnalyticsService 聚合用户数、帖子数、行为类型分布、内容标签分布、热门话题、广告曝光点击和负反馈等指标，并以结构化数据返回前端；前端在数据统计中心、算法对比页和用户画像页中使用折线图、柱状图、环形图、雷达图和词云等形式展示结果。")
    add_para(doc, "算法对比页面展示推荐流与时间倒序流的平均推荐分、标签命中率、TF-IDF 相似度、作者去重率和标签覆盖率等指标；用户画像页展示兴趣标签、行为分布和近期兴趣序列。通过这些图表，论文后续测试章节可以结合统计结果分析推荐效果，而不是只依赖主观页面观察。")

    add_para(doc, "综上，本章按照模板要求对系统进行了总体设计，明确了系统架构、功能结构、核心算法、数据库结构和关键实现技术。下一章将在此基础上，从用户界面设计和程序处理流程两个角度，对各核心功能模块的详细设计与实现过程进行说明。")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
