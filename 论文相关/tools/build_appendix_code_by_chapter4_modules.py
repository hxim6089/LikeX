# -*- coding: utf-8 -*-
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\TheEnd\qnyproj-main\recommendation-system")
OUT = ROOT / "论文相关" / "附录_按第4章模块整理_核心代码.docx"


@dataclass
class Snippet:
    title: str
    source: str
    code: str
    start_line: int
    end_line: int


@dataclass
class Module:
    title: str
    note: str
    snippets: list[Snippet]


def read_source(rel: str) -> str:
    return (ROOT / rel).read_text(encoding="utf-8")


def line_no(text: str, idx: int) -> int:
    return text.count("\n", 0, idx) + 1


def normalize_code(code: str) -> str:
    return textwrap.dedent(code.replace("\t", "    ")).strip("\n")


def include_annotations(text: str, start_idx: int) -> int:
    pos = start_idx
    best = start_idx
    for line in reversed(text[:start_idx].splitlines(keepends=True)):
        pos -= len(line)
        stripped = line.strip()
        if stripped.startswith("@"):
            best = pos
            continue
        if stripped == "":
            continue
        break
    return best


def find_marker(text: str, marker: str, occurrence: int = 1) -> int:
    offset = 0
    count = 0
    while True:
        idx = text.find(marker, offset)
        if idx < 0:
            raise ValueError(f"Cannot find marker: {marker}")
        count += 1
        if count == occurrence:
            return text.rfind("\n", 0, idx) + 1
        offset = idx + len(marker)


def java_block(rel: str, marker: str, title: str | None = None, occurrence: int = 1) -> Snippet:
    text = read_source(rel)
    start = include_annotations(text, find_marker(text, marker, occurrence))
    open_idx = text.find("{", start)
    if open_idx < 0:
        raise ValueError(f"No opening brace after {marker}")

    depth = 0
    in_str = in_char = in_line_comment = in_block_comment = False
    escaped = False
    i = open_idx
    while i < len(text):
        ch = text[i]
        nxt = text[i + 1] if i + 1 < len(text) else ""
        if in_line_comment:
            if ch == "\n":
                in_line_comment = False
            i += 1
            continue
        if in_block_comment:
            if ch == "*" and nxt == "/":
                in_block_comment = False
                i += 2
            else:
                i += 1
            continue
        if in_str:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_str = False
            i += 1
            continue
        if in_char:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == "'":
                in_char = False
            i += 1
            continue
        if ch == "/" and nxt == "/":
            in_line_comment = True
            i += 2
            continue
        if ch == "/" and nxt == "*":
            in_block_comment = True
            i += 2
            continue
        if ch == '"':
            in_str = True
            i += 1
            continue
        if ch == "'":
            in_char = True
            i += 1
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                code = normalize_code(text[start:end])
                return Snippet(title or marker.strip(), rel, code, line_no(text, start), line_no(text, i))
        i += 1
    raise ValueError(f"Unbalanced block: {marker}")


def range_block(rel: str, start_line: int, end_line: int, title: str) -> Snippet:
    lines = read_source(rel).splitlines()
    code = "\n".join(lines[start_line - 1:end_line])
    return Snippet(title, rel, normalize_code(code), start_line, end_line)


def wrap_line(line: str, width: int = 112) -> list[str]:
    if len(line) <= width:
        return [line]
    indent = len(line) - len(line.lstrip(" "))
    subsequent = " " * min(indent + 8, 24)
    return textwrap.wrap(
        line,
        width=width,
        subsequent_indent=subsequent,
        break_long_words=False,
        break_on_hyphens=False,
        replace_whitespace=False,
    ) or [""]


def set_font(run, name: str, size: float, bold: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_bottom_border(paragraph, color="BFBFBF", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_para(doc: Document, text: str, size=10.5, font="宋体", bold=False, color=None, first_line=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, font, size, bold, color)
    return p


def add_code(doc: Document, snippet: Snippet):
    header = add_para(
        doc,
        f"代码来源：{snippet.source}（第 {snippet.start_line}-{snippet.end_line} 行）",
        size=9,
        font="宋体",
        color=RGBColor(80, 80, 80),
    )
    add_bottom_border(header, "D9D9D9", "6")
    for raw in snippet.code.splitlines():
        for line in wrap_line(raw.rstrip()):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line if line else " ")
            set_font(run, "Consolas", 7.5, color=RGBColor(30, 30, 30))
    add_para(doc, "", size=6)


def build_modules() -> list[Module]:
    return [
        Module(
            "A.1 对应 4.1 用户认证与内容互动模块的核心代码",
            "本模块保留用户注册登录、内容发布、互动行为写入等核心代码，用于说明用户侧基础业务闭环。",
            [
                java_block("backend/src/main/java/com/example/rec/service/AuthService.java", "public User register(User user)", "用户注册逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/AuthService.java", "public User login(String username, String password)", "用户登录与封禁校验逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/ContentService.java", "public Content publish(Long authorId, String text, String imageUrl)", "内容发布与话题标签解析逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/BehaviorService.java", "public void likeContent(Long userId, Long contentId)", "点赞行为写入与计数更新逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/BehaviorService.java", "public void recordView(Long userId, Long contentId, Integer duration)", "浏览行为采集逻辑"),
            ],
        ),
        Module(
            "A.2 对应 4.2 推荐排序与用户画像模块的核心代码",
            "本模块保留推荐候选池构建、混合推荐评分、近期兴趣序列和对比指标计算等代码。",
            [
                java_block("backend/src/main/java/com/example/rec/service/RecommendationService.java", "public List<com.example.rec.dto.ContentWithScore> getRecommendedFeedWithScore(Long userId)", "推荐流入口与策略调用逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/RecommendationService.java", "private List<Content> buildCandidatePoolWithStats(Long userId, PipelineStats stats)", "候选内容池构建逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/HybridRecommendationStrategy.java", "public List<ContentWithScore> recommendWithScore(Long userId, List<Content> candidates, Map<String, Double> manualWeights)", "混合推荐排序主流程"),
                java_block("backend/src/main/java/com/example/rec/service/HybridRecommendationStrategy.java", "private ScoredContentWithDetails calculateScoreWithDetails(Content content,", "推荐评分明细计算逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/PersonaService.java", "private Map<String, Object> computeRecentInterestSequence(List<Behavior> allBehaviors, LocalDateTime now)", "近期兴趣序列生成逻辑"),
                java_block("backend/src/main/java/com/example/rec/controller/CompareController.java", "private Map<String, Object> calculateDiversityStats(List<ContentWithScore> list)", "作者去重率与标签覆盖率计算逻辑"),
            ],
        ),
        Module(
            "A.3 对应 4.3 AI推荐与后台管理模块的核心代码",
            "本模块保留推荐策略切换、AI 推荐降级、AI 打标和数据导入入口等核心实现。",
            [
                java_block("backend/src/main/java/com/example/rec/service/RecommendationStrategyManager.java", "public void switchStrategy(String type)", "推荐策略切换逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/RecommendationStrategyManager.java", "public List<ContentWithScore> recommendWithScore(Long userId, List<Content> candidates)", "traditional/ai 统一推荐入口"),
                java_block("backend/src/main/java/com/example/rec/service/AiRecommendationStrategy.java", "public List<ContentWithScore> recommendWithScore(Long userId, List<Content> candidates)", "AI 推荐排序与异常降级逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/AiTaggingService.java", "public List<String> tagContent(Content content)", "AI 内容打标逻辑"),
                java_block("backend/src/main/java/com/example/rec/controller/AdminController.java", "public Map<String, Object> switchRecStrategy(@RequestBody Map<String, String> payload)", "后台推荐策略切换接口"),
                java_block("backend/src/main/java/com/example/rec/controller/AdminController.java", "public Map<String, Object> importKaggleBatch(@RequestBody(required = false) Map<String, Object> payload)", "后台批量导入数据接口"),
            ],
        ),
        Module(
            "A.4 对应 4.4 广告分发模块的核心代码",
            "本模块保留广告画像匹配、广告曝光点击记录和前端广告事件上报代码。",
            [
                java_block("backend/src/main/java/com/example/rec/service/AdService.java", "public List<Map<String, Object>> getRelevantAds(Long userId, int count)", "基于画像标签的广告匹配逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/AdService.java", "public void recordImpression(Long adId, Long userId)", "广告曝光记录逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/AdService.java", "public void recordClick(Long adId)", "广告点击记录逻辑"),
                range_block("frontend/src/components/AdCard.vue", 50, 69, "前端广告曝光与点击上报逻辑"),
            ],
        ),
        Module(
            "A.5 对应 4.5 数据统计与画像可视化模块的核心代码",
            "本模块保留行为类型分布、内容标签分布、近期兴趣序列展示和统计图表渲染相关代码。",
            [
                java_block("backend/src/main/java/com/example/rec/service/AnalyticsService.java", "private Map<String, Object> buildBehaviorStats(List<Behavior> behaviors, int trendDays)", "行为类型分布统计逻辑"),
                java_block("backend/src/main/java/com/example/rec/service/AnalyticsService.java", "private Map<String, Object> buildContentStats(List<Content> contents, int trendDays, int trendingHours)", "内容分类与标签分布统计逻辑"),
                range_block("frontend/src/views/AnalyticsView.vue", 400, 480, "统计中心行为分布与标签分布图表渲染逻辑"),
                range_block("frontend/src/components/PersonaDetailCard.vue", 55, 80, "近期兴趣序列页面展示模板"),
                range_block("frontend/src/components/PersonaDetailCard.vue", 190, 220, "近期兴趣序列折线图数据构造逻辑"),
            ],
        ),
    ]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("附录  程序源代码（按第4章模块整理）")
    set_font(run, "黑体", 16, bold=True)

    add_para(
        doc,
        "以下代码依据项目实际源码整理，按照第4章“系统详细设计及实现”的模块顺序选取核心方法，省略实体类、Repository接口、Getter/Setter和样式代码，仅保留与业务流程、推荐算法、AI处理、广告分发和数据统计直接相关的关键代码。",
        size=10.5,
        font="宋体",
        first_line=True,
    )

    for idx, module in enumerate(build_modules()):
        if idx > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(module.title)
        set_font(run, "黑体", 13, bold=True)
        add_para(doc, module.note, size=10.5, font="宋体", first_line=True)
        for sn_idx, snippet in enumerate(module.snippets, 1):
            sub = doc.add_paragraph()
            sub.paragraph_format.space_before = Pt(6)
            sub.paragraph_format.space_after = Pt(3)
            run = sub.add_run(f"（{sn_idx}）{snippet.title}")
            set_font(run, "黑体", 11, bold=True)
            add_code(doc, snippet)

    doc.save(OUT)
    total_snippets = sum(len(m.snippets) for m in build_modules())
    total_lines = sum(sn.code.count("\n") + 1 for m in build_modules() for sn in m.snippets)
    print(OUT)
    print(f"modules={len(build_modules())}, snippets={total_snippets}, code_lines={total_lines}")


if __name__ == "__main__":
    build_doc()
