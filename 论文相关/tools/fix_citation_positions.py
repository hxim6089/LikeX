import sys
from pathlib import Path

from docx import Document


SRC = Path(sys.argv[1])
OUT = Path("毕业论文初稿_正文参考文献引用位置校正版.docx")


REPLACEMENTS = {
    "在国内的工程语境中，Spring Boot、Vue 与 MySQL 的组合具有较低学习成本": (
        "在国内的工程语境中，前后端分离 Web 系统具有较低学习成本和较好的工程实践基础[6]。"
        "本系统后端以 Java 语言和 Spring Boot 框架组织 RESTful API，相关编码规范与面向对象设计可参考 Java 开发基础文献[9]；"
        "前端采用 Vue 3 与 Element Plus 构建单页应用，其组件化开发方式与管理后台页面组织可参考 Vue 3 工程实践文献[10]。"
        "数据库层面，本文以关系型模型承载用户、帖子、互动、消息、通知、广告和推荐日志等核心对象，并通过索引与业务分层控制查询复杂度[13]。"
        "算法层面，推荐系统基础理论、矩阵分解和基于物品的协同过滤为本文理解用户兴趣建模、隐式反馈和排序评估提供了方法参照[11][16][17]。"
    ),
    "接口设计方面，系统围绕用户、帖子、互动、关注、消息、通知、广告、画像和推荐策略等资源划分 RESTful API": (
        "接口设计方面，系统围绕用户、帖子、互动、关注、消息、通知、广告、画像和推荐策略等资源划分 RESTful API，"
        "使前端页面能够通过统一路径访问后端服务。该设计与前后端分离 Web 框架设计思想一致[6]，"
        "也符合数据密集型应用中强调接口边界、数据一致性和运行维护协同的设计原则[13]。"
        "实时私信和通知则通过 WebSocket 通道实现近实时推送，便于后续进行接口测试、权限控制和功能扩展。"
    ),
    "推荐排序模块采用策略模式设计，统一定义推荐策略接口": (
        "推荐排序模块采用策略模式设计，统一定义推荐策略接口（RecommendationStrategy），传统混合推荐策略（HybridRecommendationStrategy）和 AI 推荐策略（AiRecommendationStrategy）均实现该接口。"
        "推荐策略管理器（RecommendationStrategyManager）持有两种策略实例并维护全局策略标识（如 traditional、ai），推荐服务（RecommendationService）在组装候选池后通过 getActiveStrategy() 获取当前策略执行排序，从而避免在业务代码中硬编码单一算法实现。"
        "该设计对应设计模式中“对变化点进行封装”的思想[12]。同时，系统保留个性化重排序扩展空间，能够在候选列表生成后继续接入新的重排序策略，这与个性化重排序研究中将多源特征用于排序优化的思路一致[5]。"
    ),
    "（4）TF-IDF 与余弦相似度。对文本内容分词后": (
        "（4）TF-IDF 与余弦相似度。文本特征是推荐系统理解内容语义的重要输入，推荐系统理论中通常会将文本、标签或类别特征转化为可计算向量[11]。"
        "本文对文本内容分词后，词项 i 在文档 d 中的 TF-IDF 权重采用式（3-1）计算，其中 N 为语料文档数，df_i 为包含词 i 的文档数，tf_i,d 为词项 i 在文档 d 中的词频。"
    ),
    "（9）AI 推荐与降级。AI 推荐策略": (
        "（9）AI 推荐与降级。AI 推荐策略（AiRecommendationStrategy）将画像摘要与候选摘要（含 id、正文片段、标签、互动统计）封装为提示词（Prompt），约束模型仅输出 JSON 结构，其中排序字段（ranking）表示候选顺序，理由字段（reasons）表示推荐解释。"
        "解析成功则重排列表；若 Ollama 不可用、超时或 JSON 非法，则降级为按基础互动分排序并在前端标识降级模式，保障系统可用性。"
        "该设计并非替代传统推荐模型，而是在混合排序基础上增加语义理解、排序解释与对比实验入口；其思想与深度神经网络个性化推荐、DeepFM 特征交叉建模以及 DIEN 兴趣演化建模等研究中强调“用户兴趣表征 + 排序优化”的方向一致[14][18][19]，同时仍以推荐系统基础理论作为整体方法参照[11]。"
    ),
    "广告业务实现补充。典型路径为": (
        "广告业务实现补充。典型路径为：管理员在广告报表页（AdDashboard）中执行创建（POST）与编辑（PUT），并通过更新启用状态字段（active）控制广告启停；列表页（GET）展示广告明细与统计数据。"
        "首页展示前再 GET /api/ads/relevant 或等价接口拉取当前用户最相关广告，与帖子流在客户端按广告间隔（interval）合并。"
        "展示、点击分别调用统计接口回写展示次数（impressionCount）和点击次数（clickCount）。"
        "本系统以用户画像、广告标签、广告出价和 CTR 统计为基础实现规则排序，属于个性化广告推荐系统在毕业设计规模下的工程化实现[3]。"
    ),
}


def replace_para_text(paragraph, new_text):
    # Preserve the paragraph object and basic position; formatting is not the priority here.
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def main():
    doc = Document(str(SRC))
    hits = {}
    for p in doc.paragraphs:
        text = p.text.strip()
        for key, value in REPLACEMENTS.items():
            if key in text:
                replace_para_text(p, value)
                hits[key] = hits.get(key, 0) + 1
                break

    missing = [k for k in REPLACEMENTS if hits.get(k, 0) == 0]
    if missing:
        raise RuntimeError("未找到待替换段落: " + " | ".join(missing))

    doc.save(str(OUT))
    print(str(OUT))
    print(hits)


if __name__ == "__main__":
    main()
