from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = "论文待优化章节_优化稿.docx"


def set_run_font(run, font="宋体", size=12, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_para(paragraph, first_line=True, align=None, before=0, after=0, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    if align is not None:
        paragraph.alignment = align


def add_para(doc, text="", first_line=True, align=None, size=12, bold=False, before=0, after=0):
    p = doc.add_paragraph()
    set_para(p, first_line=first_line, align=align, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(12 if level == 1 else 8)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right}
    for edge, val in specs.items():
        tag = "w:" + edge
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        if val is None:
            elem.set(qn("w:val"), "nil")
        else:
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), str(val))
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "000000")


def three_line_table(doc, caption, headers, rows, widths=None):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(caption)
    set_run_font(r, size=10.5, bold=False)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(val)) > 12 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], str(val), align=align)

    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell, left=None, right=None, top=None, bottom=None)
            if ridx == 0:
                set_cell_border(cell, top=12, bottom=6, left=None, right=None)
            if ridx == len(table.rows) - 1:
                set_cell_border(cell, bottom=12, left=None, right=None)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    add_para(doc, "", first_line=False, after=0)
    return table


def add_formula(doc, formula, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(formula)
    set_run_font(r, font="Times New Roman", size=12)
    r2 = p.add_run(" " * 12 + number)
    set_run_font(r2, font="Times New Roman", size=12)


def fig_placeholder(doc, text):
    add_para(doc, f"【截图/图示预留：{text}】", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, before=6, after=3)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.6)

styles = doc.styles
styles["Normal"].font.name = "宋体"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
r = title.add_run("毕业论文待优化章节修订稿")
set_run_font(r, size=18, bold=True)
add_para(doc, "说明：本稿依据任务书要求、当前项目代码功能和初稿中存在的薄弱环节整理，主要用于替换或补充论文第 2 章、第 3 章、第 4 章、第 5 章相关内容。统计分析部分预留了实际运行数据与系统截图位置，定稿时应根据本机接口返回值和页面截图补齐。", first_line=True)

add_heading(doc, "第2章 系统需求分析（优化稿）", 1)
add_heading(doc, "2.1 功能需求", 2)
add_para(doc, "根据任务书要求，系统面向类 X/Twitter 的社交媒体信息流场景，核心目标是在普通时间倒序信息流基础上，引入用户行为采集、用户画像、混合多因子推荐排序、AI 辅助能力、广告分发和算法效果可视化验证。系统采用 B/S 架构和前后端分离模式，前端使用 Vue 3、Vite 与 Element Plus 构建单页应用，后端使用 Spring Boot 提供 RESTful API 与 WebSocket 通信能力，数据库使用 MySQL 保存用户、内容、行为、社交关系、标签、通知、私信、广告和推荐相关数据。")
add_para(doc, "从用户角色看，系统主要包括普通用户和管理员两类对象。普通用户可完成注册登录、浏览 For You 推荐流、浏览 Following 关注流、发布帖子、上传图片、评论、点赞、点踩、转发、引用、关注、搜索、查看通知、发送私信、查看个人画像等操作；管理员可查看系统统计信息、切换推荐策略、导入外部数据、执行 AI 批量打标、维护广告配置并查看广告展示与点击统计。")

three_line_table(
    doc,
    "表 2-1 任务书要求与系统功能对应关系",
    ["任务书要求", "系统功能设计", "项目实现依据"],
    [
        ["基础社交平台", "注册登录、角色区分、发帖、评论、点赞、转发、引用、关注、搜索、通知、私信", "AuthController、ContentController、BehaviorController、RelationController、MessageController"],
        ["行为采集与画像", "记录 VIEW、LIKE、COMMENT、REPOST、SEARCH、DISLIKE、SKIP 等行为，生成兴趣标签、作者偏好、互动风格和动态权重", "BehaviorService、SearchController、UserBehaviorProfileService、PersonaService"],
        ["混合推荐排序", "候选池构建、负反馈过滤、TF-IDF 相似度、协同过滤、时间衰减、热门话题、作者多样性、随机探索", "RecommendationService、HybridRecommendationStrategy、TfIdfService、CollaborativeFilteringService"],
        ["AI 智能模块", "AI 问答、AI 自动打标、AI 推荐排序与降级", "AiService、AiTaggingService、AiRecommendationStrategy、AiRecCacheService"],
        ["广告智能分发", "按间隔插入信息流广告，依据标签匹配、出价和 CTR 计算广告得分", "AdService、AdController、AdDashboard"],
        ["可视化对比与调参", "推荐流与时间流对比、统计指标展示、权重调节、管道漏斗图", "CompareController、CompareView、WeightTuner、PipelineFunnel"],
    ],
    widths=[3.8, 6.4, 5.4],
)

add_heading(doc, "2.2 系统业务流程分析", 2)
add_para(doc, "系统的主业务流程可分为用户访问流程、内容生产流程、行为反馈流程、推荐排序流程和后台治理流程。用户首先通过注册或登录获取访问凭证，前端在本地保存 Token 与用户信息，并在需要身份校验的请求中携带会话信息。登录后，用户可在首页切换 For You 推荐流和 Following 关注流；前者由后端推荐服务根据候选池与推荐策略排序生成，后者主要根据关注关系筛选已关注作者的内容。")
add_para(doc, "内容生产流程以发帖为入口。用户输入正文并可选择上传图片，图片先通过上传接口保存并返回可访问地址，随后正文、作者信息和图片地址一起提交至内容发布接口。服务端保存帖子后解析正文中的话题标签，并可触发 AI 自动打标，为后续搜索、热门话题和推荐排序提供内容特征。")
add_para(doc, "行为反馈流程贯穿用户浏览和互动全过程。系统将浏览、点赞、评论、转发、搜索、点踩、快速滑过等事件写入行为表，并同步维护内容的互动计数。行为数据既用于用户画像构建，也用于协同过滤、热门内容统计、推荐得分拆分和后台统计分析。")
fig_placeholder(doc, "绘制系统业务流程图。建议图题为“图 2-1 系统核心业务流程图”，内容包括登录认证、发帖、行为采集、画像更新、推荐排序、广告插入和通知私信推送。")
add_para(doc, "图 2-1 系统核心业务流程图", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)

add_heading(doc, "2.3 数据流分析", 2)
add_para(doc, "系统数据流以用户行为为中心展开。前端页面产生的用户操作被封装为 API 请求进入后端控制层，控制层调用业务服务完成数据校验、状态更新和统计计算，最终通过 Repository 写入 MySQL。推荐服务读取内容表、行为表、关注表、标签表和负反馈表，构建候选池并计算推荐结果；前端再将推荐列表、评分拆分、广告项和统计图表组合展示给用户。")
add_para(doc, "在推荐场景中，数据流可概括为“行为采集—画像构建—候选召回—负反馈过滤—多因子打分—广告插入—页面展示”。其中，画像构建依赖历史行为、内容标签和作者信息；候选召回同时考虑关注来源和全站公开内容；负反馈过滤用于排除屏蔽作者和不感兴趣内容；多因子打分输出最终排序及 ScoreBreakdown 评分明细；广告模块再根据用户画像和广告配置选择合适广告插入信息流。")
fig_placeholder(doc, "绘制系统数据流图。建议图题为“图 2-2 系统数据流图”，重点体现用户行为数据如何进入画像、推荐、广告和统计模块。")
add_para(doc, "图 2-2 系统数据流图", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
three_line_table(
    doc,
    "表 2-2 核心数据流说明",
    ["数据流名称", "输入来源", "处理模块", "输出结果"],
    [
        ["认证数据流", "登录/注册表单", "AuthController、AuthService、JwtUtil", "用户信息与 Token"],
        ["内容数据流", "发帖、图片上传、评论、转发", "ContentController、FileUploadController、ContentService", "帖子、评论、图片地址、话题标签"],
        ["行为数据流", "浏览、点赞、评论、转发、搜索、点踩、跳过", "BehaviorController、SearchController、BehaviorService", "行为记录与内容计数"],
        ["推荐数据流", "用户画像、候选内容、标签、负反馈、协同结果", "RecommendationService、HybridRecommendationStrategy", "排序列表与评分拆分"],
        ["AI 数据流", "画像摘要、候选摘要、帖子正文", "AiRecommendationStrategy、AiTaggingService、AiService", "AI 排序、推荐理由、语义标签、问答回复"],
        ["广告数据流", "用户兴趣标签、广告标签、出价、CTR", "AdService、AdController", "相关广告、展示点击统计"],
    ],
    widths=[3.0, 4.5, 4.8, 3.5],
)

add_heading(doc, "2.4 数据字典", 2)
add_para(doc, "为保证论文设计描述与项目数据库模型一致，本系统的数据字典围绕用户、内容、行为、标签、关注、通知、私信、负反馈和广告等实体展开。数据字典不只是字段罗列，还应说明字段在业务流程和推荐排序中的作用。例如，Behavior 表的 type 字段决定行为权重，Content 表的 tags 和 category 字段参与标签匹配与热门话题统计，NegativeSignal 表用于候选过滤，Ad 表的 bidPrice、impressionCount 与 clickCount 用于广告排序和 CTR 统计。")
three_line_table(
    doc,
    "表 2-3 核心数据字典（节选）",
    ["数据对象", "主要字段", "含义与用途"],
    [
        ["User", "id、username、password、role、customWeights", "保存用户身份、角色和自定义推荐权重，是认证、权限控制和个性化推荐的基础。"],
        ["Content", "id、author、content、imageUrl、tags、parentContent、repostOf、quoteOf、likeCount、viewCount", "保存帖子、评论、转发和引用内容，互动计数用于推荐打分和后台统计。"],
        ["Behavior", "userId、contentId、type、duration、createdAt", "记录用户行为事件，支撑用户画像、协同过滤、推荐评分和测试统计。"],
        ["Tag", "id、name", "保存话题标签，与 Content 形成多对多关系，用于搜索、热门话题和兴趣匹配。"],
        ["Follow", "followerId、followeeId", "保存关注关系，是 Following 信息流和 In-Network 候选来源的依据。"],
        ["NegativeSignal", "userId、targetType、targetId、signalType", "保存屏蔽作者、不感兴趣等负反馈，用于候选过滤和降权。"],
        ["Ad", "targetTags、bidPrice、impressionCount、clickCount、active", "保存广告投放配置和统计字段，用于相关广告筛选和 CTR 计算。"],
    ],
    widths=[2.8, 5.8, 7.0],
)

add_heading(doc, "第3章 推荐算法与 AI 关键技术（优化稿）", 1)
add_heading(doc, "3.4 推荐算法设计补充", 2)
add_para(doc, "系统推荐排序采用传统混合推荐策略与 AI 推荐策略并存的设计。传统策略由 HybridRecommendationStrategy 实现，强调可解释性、响应速度和可调参；AI 策略由 AiRecommendationStrategy 实现，强调语义理解和推荐理由生成。RecommendationStrategyManager 统一管理当前生效策略，使管理员能够在后台切换 traditional 与 ai，并在 AI 服务不可用时回退到稳定的传统排序链路。")
add_para(doc, "传统推荐策略先由 RecommendationService 构建候选池。候选池由 In-Network 与 Out-of-Network 两部分组成，前者来自当前用户关注对象发布的内容，后者来自全站公开内容。候选合并后执行负反馈过滤，再交由推荐策略进行多因子打分。该设计与任务书中“双源候选机制、负反馈过滤、混合多因子排序和算法可视化”的要求一致。")
add_para(doc, "系统对文本内容进行分词并过滤停用词后，词项 i 在文档 d 中的 TF-IDF 权重采用式（3-1）计算：")
add_formula(doc, "w_i,d = tf_i,d × log(N / df_i)", "（3-1）")
add_para(doc, "其中，N 表示语料文档总数，df_i 表示包含词项 i 的文档数量，tf_i,d 表示词项 i 在文档 d 中的词频。项目实现中，TfIdfService 先基于全量帖子构建全局 IDF，再使用用户点赞内容的 TF-IDF 向量均值作为用户兴趣向量。候选帖子向量与用户兴趣向量通过余弦相似度计算，公式如下：")
add_formula(doc, "sim(u,c) = Σ_i w_i,u w_i,c / (||u|| · ||c||)", "（3-2）")
add_para(doc, "该相似度作为内容相似度加成（tfidfSimilarityBoost）的主要来源，在代码中对应 ScoreBreakdown 的 contentSimilarityBoost 字段，用于弥补单纯标签匹配粒度较粗的问题。")
add_para(doc, "综合评分由互动分、互动率修正、关注来源加成、话题亲和度、作者亲和度、标签匹配、文本相似度、热门话题、协同过滤、内容深度、新鲜度和随机扰动共同构成，可概括为式（3-3）：")
add_formula(doc, "finalScore = E / D_t + B_topic + B_author + B_personal + B_tfidf + B_trend + B_cf + B_depth + B_fresh + ε", "（3-3）")
add_para(doc, "其中，E 为基础互动得分，D_t 为时间衰减因子，B_topic 表示话题亲和度加成，B_author 表示作者亲和度加成，B_personal 表示标签匹配加成，B_tfidf 表示 TF-IDF 内容相似度加成，B_trend 表示热门话题加成，B_cf 表示协同过滤加成，B_depth 表示内容深度匹配加成，B_fresh 表示新鲜度偏好加成，ε 表示探索随机扰动项。")
add_para(doc, "推荐结果不是简单截取 Top-N，而是在排序后引入加权随机采样机制。设第 i 条候选的采样权重为 w_i，前 k 项权重前缀和为 S_k，则抽样过程可表示为式（3-4）：")
add_formula(doc, "S_k = Σ_{i=1}^{k} w_i，r ∈ (0, S_M]", "（3-4）")
add_para(doc, "当随机数 r 落入某一权重区间时选中对应候选。该机制使高分内容仍具有更大曝光概率，同时保留低分但非零分内容的探索机会，能够缓解信息流结果固化和同质化问题。")

add_heading(doc, "3.5 AI 底层技术与系统集成原理", 2)
add_para(doc, "AI 模块并非替代传统推荐算法，而是作为语义理解、辅助排序、内容打标和问答交互的增强层。系统使用本地 Ollama 服务调用大语言模型完成 AI 推荐与 AI 自动打标，并通过 DeepSeek API 提供类 Grok 的智能问答助手。大语言模型的底层通常基于 Transformer 架构，通过自注意力机制建模输入序列中不同词项之间的关联关系，从而在较长上下文中理解用户画像、候选帖子语义和输出格式要求。")
add_para(doc, "在 AI 推荐场景中，系统先按基础互动量对候选帖子进行预排序，并限制输入 AI 的候选数量。AiRecommendationStrategy 将用户阶段、兴趣话题、互动风格、内容深度偏好、新鲜度偏好和候选帖子摘要拼接为 Prompt，要求模型只返回包含 ranking 和 reasons 的 JSON 结构。ranking 字段表示 AI 认为的候选排序，reasons 字段表示推荐理由。后端解析 JSON 后按 AI 返回顺序重排候选，同时将推荐理由写入 ScoreBreakdown，供前端展示。")
add_para(doc, "由于大语言模型推理耗时明显高于传统公式打分，系统在工程实现上采用异步预计算、缓存和降级策略。AiRecCacheService 使用内存缓存保存用户的 AI 推荐结果，并设置 TTL；缓存未命中时，系统先返回传统推荐结果，同时在后台触发 AI 计算；缓存命中时直接返回已计算的 AI 结果。若 Ollama 服务不可用、调用超时或返回 JSON 解析失败，系统自动降级为按基础互动量排序，保证信息流主链路不会因 AI 服务异常而中断。")
add_para(doc, "AI 自动打标模块由 AiTaggingService 实现。系统将帖子正文和候选标签列表发送给模型，要求模型仅输出 1 至 3 个标签名称，不输出解释文本。后端再清洗模型返回值，将标签写入 Tag 表和 content_tags 关联关系。该设计将大语言模型的语义理解能力转化为结构化标签数据，为搜索、热门话题、用户画像和推荐排序提供补充特征。")

add_heading(doc, "第4章 关键模块实现（优化稿）", 1)
add_heading(doc, "4.5 AI 模块与管理后台实现", 2)
add_para(doc, "AI 模块包括 AI 问答、AI 自动打标和 AI 推荐三类功能。AI 问答由 AiService 调用外部对话模型接口，系统在 System Prompt 中注入平台功能背景，使回答更贴近本系统的推荐、画像、广告和可视化场景。AI 自动打标由 AiTaggingService 调用 Ollama 的 /api/generate 接口，输入帖子正文和候选标签，输出结构化标签名称。AI 推荐由 AiRecommendationStrategy 负责，将用户画像摘要和候选内容摘要封装为 Prompt，并解析模型返回的 JSON 排序结果。")
add_para(doc, "管理后台通过 AdminController 和 AdminView 提供推荐策略切换、统计概览、Kaggle 数据导入和 AI 批量打标等功能。管理员可在 traditional 与 ai 两种推荐策略之间切换，系统通过 RecommendationStrategyManager 保存当前策略。为避免 AI 推理服务影响主链路，策略管理接口同时返回 AI 可用状态和缓存统计信息，便于管理员判断当前是否适合启用 AI 推荐。")
fig_placeholder(doc, "放置 AI 推荐策略切换或 AI 助手页面截图。建议图题为“图 4-x AI 推荐策略与智能助手界面”。")
add_para(doc, "图 4-x AI 推荐策略与智能助手界面", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)

add_heading(doc, "4.8 统计分析与可视化实现", 2)
add_para(doc, "为满足任务书中“算法效果可视化对比与参数调优”的要求，系统在后端提供 CompareController 和 AnalyticsController 两类统计接口。CompareController 面向推荐算法验证，返回个性化推荐流、时间倒序流、评分拆分对象、推荐管道统计和对比指标；AnalyticsController 面向后台看板，统计用户数量、内容数量、行为类型分布、用户阶段分布、内容分类、热门话题、负反馈和广告投放数据。")
add_para(doc, "CompareView 使用 ECharts 将个性化推荐与时间倒序的统计差异以柱状图呈现，指标包括平均推荐分、兴趣标签命中率、TF-IDF 相似度均值、平均互动量和个性化提升倍率。PipelineFunnel 以漏斗图展示全量候选、关注/其他作者候选、负反馈过滤、评分截断和最终分发等阶段的数量变化。WeightTuner 提供点赞、评论、转发、个性化、热门话题和 TF-IDF 相似度等权重滑块，并提供默认、互动优先、个性化优先和热门优先等预设方案。")
fig_placeholder(doc, "放置 CompareView 中推荐流与时间流对比、统计柱状图、WeightTuner 或 PipelineFunnel 截图。建议图题为“图 4-x 推荐效果对比与参数调节界面”。")
add_para(doc, "图 4-x 推荐效果对比与参数调节界面", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)

add_heading(doc, "第5章 系统测试与统计分析（优化稿）", 1)
add_heading(doc, "5.1.2 测试过程及结果优化说明", 2)
add_para(doc, "系统测试采用功能场景测试、接口联调验证、推荐对比验证和统计结果分析相结合的方式进行。考虑到当前项目尚未建立独立的后端单元测试目录和前端组件测试脚本，本文不将测试结论表述为持续集成自动通过，而是以手工执行记录、接口响应结果、页面运行截图和推荐对比统计结果作为验证依据。测试过程围绕用户认证、内容发布、互动反馈、社交关系、实时通信、推荐排序、AI 模块、广告分发和后台管理等核心链路展开。")
add_para(doc, "与单纯功能测试相比，推荐系统更需要通过统计结果说明算法是否产生作用。因此，本文在推荐排序测试中采用 /api/compare/feed 接口对同一用户的个性化推荐流和时间倒序流作对比统计，并结合 CompareView 页面截图分析平均推荐分、兴趣标签命中率、TF-IDF 相似度均值、平均互动量和个性化提升倍率等指标。")

add_heading(doc, "5.1.3 推荐效果与行为数据统计分析", 2)
add_para(doc, "为验证推荐排序模块的有效性，本文采用 CompareController 提供的 /api/compare/feed 接口，对用户在同一候选集合下的个性化推荐结果和时间倒序结果作对比统计。统计范围限定为前 20 条内容，避免列表尾部低曝光内容对展示结果造成干扰。统计指标包括平均推荐分、兴趣标签命中率、TF-IDF 相似度均值、平均互动量和个性化提升倍率。")
three_line_table(
    doc,
    "表 5-6 推荐流与时间流对比统计表",
    ["统计指标", "个性化推荐流", "时间倒序流", "分析说明"],
    [
        ["平均推荐分", "【运行后填入】", "【运行后填入】", "反映多因子综合评分后的整体匹配程度。"],
        ["兴趣标签命中率", "【运行后填入】", "【运行后填入】", "反映推荐结果与用户画像兴趣标签的匹配情况。"],
        ["TF-IDF 相似度均值", "【运行后填入】", "【运行后填入】", "反映文本语义层面的兴趣匹配程度。"],
        ["平均互动量", "【运行后填入】", "【运行后填入】", "反映候选内容历史互动基础和内容热度。"],
        ["个性化提升倍率", "【运行后填入】", "—", "由个性化推荐平均分与时间流平均分计算得到。"],
    ],
    widths=[3.4, 3.2, 3.2, 5.8],
)
add_para(doc, "从统计结果看，若个性化推荐流在平均推荐分、标签命中率和 TF-IDF 相似度均值上高于时间倒序流，说明用户画像、标签匹配和文本相似度等因子对排序产生了有效影响；若平均互动量差异较小但个性化指标提升明显，则说明系统并非只按热门内容排序，而是在热门程度之外引入了用户兴趣特征。若个别测试中时间倒序流的互动量更高，则可解释为新近热门内容在短时间窗口内具有自然优势，此时应结合标签命中率和文本相似度综合判断推荐质量。")
fig_placeholder(doc, "放置 CompareView 推荐流与时间流对比统计截图，截图中应包含柱状图和五项统计指标。建议图题为“图 5-5 推荐流与时间流统计对比结果”。")
add_para(doc, "图 5-5 推荐流与时间流统计对比结果", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
add_para(doc, "为进一步验证推荐管道的工作机制，本文采用 PipelineFunnel 组件对候选规模变化进行统计。统计结果从全量候选开始，依次展示关注/其他作者候选拆分、负反馈过滤、评分截断和最终分发数量。该图能够直观说明推荐结果不是直接从数据库按时间读取，而是经过候选组织、过滤、打分和截断等步骤生成。")
three_line_table(
    doc,
    "表 5-7 推荐管道漏斗统计表",
    ["管道阶段", "统计数量", "作用说明"],
    [
        ["全量候选", "【运行后填入】", "数据库中可进入推荐管道的候选内容总量。"],
        ["关注/其他作者候选", "【运行后填入】", "体现 In-Network 与 Out-of-Network 双源候选构成。"],
        ["负反馈过滤后", "【运行后填入】", "剔除屏蔽作者、不感兴趣内容或强负反馈内容。"],
        ["评分截断后", "【运行后填入】", "按综合分数保留满足质量阈值的候选。"],
        ["最终分发", "【运行后填入】", "前端实际展示给用户的推荐列表数量。"],
    ],
    widths=[4.0, 3.0, 8.6],
)
fig_placeholder(doc, "放置 PipelineFunnel 漏斗图截图。建议图题为“图 5-6 推荐管道漏斗统计结果”。")
add_para(doc, "图 5-6 推荐管道漏斗统计结果", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
add_para(doc, "行为数据统计方面，本文采用 /api/analytics 接口对系统行为分布、用户阶段分布、内容分类和广告统计作汇总。行为分布可用于说明测试数据是否覆盖浏览、点赞、评论、转发、搜索、点踩和跳过等关键行为；用户阶段分布可用于说明系统是否具备冷启动用户、初级用户和活跃用户的差异化画像基础；广告统计中的展示次数、点击次数和 CTR 可用于说明广告分发模块形成了展示—点击—统计的闭环。")
fig_placeholder(doc, "放置 AnalyticsView 或后台统计看板截图，截图中应包含行为类型分布、用户阶段分布或广告 CTR。建议图题为“图 5-7 系统行为与广告统计结果”。")
add_para(doc, "图 5-7 系统行为与广告统计结果", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)

add_heading(doc, "5.1.4 AI 模块测试与分析", 2)
add_para(doc, "AI 模块测试应同时覆盖在线可用和异常降级两类场景。在线场景下，启动 Ollama 服务并确认推荐模型可用，管理员切换为 AI 推荐策略后，系统应能够根据用户画像和候选内容生成排序结果或推荐理由；异常场景下，关闭 Ollama 或模拟返回非法 JSON，系统应自动降级为传统推荐或基础互动排序，并保证首页信息流仍可访问。")
three_line_table(
    doc,
    "表 5-8 AI 模块测试与分析记录表",
    ["测试内容", "测试方法", "预期结果", "分析重点"],
    [
        ["AI 推荐在线", "启动 Ollama 后切换 AI 推荐策略", "返回 AI 排序或推荐理由", "观察 ranking、reasons 是否被正确解析。"],
        ["AI 推荐降级", "关闭 Ollama 后刷新推荐流", "系统返回传统推荐结果", "验证主链路不因模型不可用而中断。"],
        ["AI 自动打标", "对无标签帖子执行打标", "生成 1 至 3 个标签", "观察标签是否写入内容标签关系。"],
        ["AI 问答", "在 GrokView 输入系统功能问题", "返回与系统背景相关的回答", "验证 System Prompt 是否限制回答范围。"],
        ["推理耗时", "记录候选条数、Prompt 长度和响应时间", "候选越多耗时越高", "说明限制候选 Top 25 与缓存策略的必要性。"],
    ],
    widths=[3.0, 4.5, 3.8, 4.3],
)
add_para(doc, "由于大语言模型推理耗时受模型规模、硬件性能、候选条数和 Prompt 长度影响，本文不将 AI 推荐作为唯一主排序路径，而是采用传统排序兜底、AI 异步预计算和缓存命中的组合方式。测试记录中应注明模型名称、temperature、最大生成 token 数、候选条数和缓存状态，以保证实验过程可复现。")

add_heading(doc, "5.2 文献引用与正文对应关系优化", 2)
add_para(doc, "正文引用应做到“出现理论依据处有引用，工程实现描述处不过度引用”。推荐算法理论部分应优先引用任务书给出的核心文献；工程框架部分引用 Spring Boot、Vue、Java 相关文献即可；AI 部分引用大语言模型基础文献；广告和可视化等工程描述可结合任务书和项目实现说明，不宜堆砌无关引用。")
three_line_table(
    doc,
    "表 5-9 正文引用调整建议",
    ["正文位置", "建议引用", "引用理由"],
    [
        ["绪论中工业推荐系统发展", "[1]、[3]", "YouTube 深度推荐和 X/Twitter 开源推荐架构与本文信息流排序方向直接相关。"],
        ["TF-IDF 公式与内容相似度", "[4]", "Salton 与 Buckley 是 TF-IDF 经典文献，应与式（3-1）对应。"],
        ["协同过滤与矩阵分解", "[11]、[12]、[13]", "用于支撑协同过滤、用户—内容矩阵和推荐系统基础理论。"],
        ["AI 大语言模型原理", "[14]", "用于支撑 Prompt 调用、上下文学习和大模型辅助功能。"],
        ["前后端分离与 Java 实现", "[5]、[6]、[8]、[9]、[10]", "用于支撑 Spring Boot、Java、Vue 和工程实现基础。"],
        ["外文翻译文献", "新增至参考文献末尾", "若正文讨论会话感知或混合推荐，可补充引用 Bauer 与 Jannach 文章。"],
    ],
    widths=[4.4, 3.0, 8.2],
)
add_para(doc, "需要特别注意的是，若参考文献列表按任务书编号重新整理，则正文中所有引用编号必须同步调整。例如，TF-IDF 的公式说明应引用任务书参考文献 [4]，不能误引到协同过滤或广告推荐文献；X/Twitter 双源候选和推荐管道思想应引用 [3]；AI 大语言模型相关描述应引用 [14]。")

doc.save(OUT)
print(OUT)
