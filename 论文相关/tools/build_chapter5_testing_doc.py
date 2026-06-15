from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = r"D:\TheEnd\qnyproj-main\recommendation-system\论文相关\第5章_系统测试及使用说明_修改版.docx"


def set_run_font(run, east="宋体", west="Times New Roman", size=12, bold=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def set_para_font(paragraph, east="宋体", west="Times New Roman", size=12, bold=False):
    for run in paragraph.runs:
        set_run_font(run, east, west, size, bold)


def add_paragraph(doc, text="", first_line=True, align=None, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6 if level > 1 else 10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, east="黑体", size=16, bold=True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, east="黑体", size=14, bold=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, east="黑体", size=12, bold=True)
    return p


def add_caption(doc, text, kind="table"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    if kind == "table":
        set_run_font(run, east="楷体", size=10.5)
    else:
        set_run_font(run, east="宋体", size=10.5)
    return p


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "nil")


def set_cell_border(cell, edge, val="single", sz="12", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    tag = "w:{}".format(edge)
    el = borders.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        borders.append(el)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 567) for w in widths_cm)))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 567)))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def fill_cell(cell, text, center=False, bold=False, size=12):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_three_line_table(doc, caption, headers, rows, widths, center_cols=None):
    add_caption(doc, caption, "table")
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    center_cols = set(center_cols or [])
    repeat_header(table.rows[0])
    for c in table.rows[0].cells:
        clear_cell_borders(c)
    for i, header in enumerate(headers):
        fill_cell(table.rows[0].cells[i], header, center=True, bold=True)
    for row_data in rows:
        row = table.add_row()
        for c in row.cells:
            clear_cell_borders(c)
        for i, value in enumerate(row_data):
            fill_cell(row.cells[i], value, center=(i in center_cols))
    for cell in table.rows[0].cells:
        set_cell_border(cell, "top", sz="12")
        set_cell_border(cell, "bottom", sz="8")
    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", sz="12")
    add_paragraph(doc, "", first_line=False, size=12)
    return table


def add_fig_placeholder(doc, caption, note):
    p = add_paragraph(doc, "【截图预留：{}】".format(note), first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_para_font(p, east="宋体", size=12)
    add_caption(doc, caption, "figure")


test_cases = [
    ("T01", "用户认证", "系统已启动，数据库中不存在待注册用户名。", "在注册页面输入用户名 test_user、密码 123456，提交注册请求，检查用户记录、返回状态和页面反馈。"),
    ("T02", "用户认证", "已存在普通用户账号，账号未被封禁。", "输入正确用户名和密码，点击登录按钮，检查 Token、用户信息和首页跳转结果。"),
    ("T03", "用户认证", "用户已完成登录并保存本地会话。", "刷新浏览器页面或重新进入系统首页，检查登录状态是否保持。"),
    ("T04", "权限识别", "系统中存在普通用户和管理员两个账号。", "分别使用普通用户和管理员账号登录，访问首页、后台入口和受保护接口。"),
    ("T05", "内容发布", "普通用户已登录，首页发布区域可用。", "输入文本内容并点击发布按钮，检查新帖子是否保存并出现在信息流中。"),
    ("T06", "内容发布", "普通用户已登录，上传接口可访问。", "选择本地图片上传，上传成功后连同正文提交帖子，检查图片地址和展示效果。"),
    ("T07", "内容互动", "普通用户已登录，目标帖子存在。", "打开目标帖子评论框，输入评论内容并提交，检查评论内容和评论计数。"),
    ("T08", "内容互动", "普通用户已登录，目标帖子存在。", "对目标帖子执行点赞或点踩操作，检查计数变化和行为记录。"),
    ("T09", "内容互动", "普通用户已登录，目标帖子存在。", "对已有帖子执行转发或引用操作，检查关联内容和转发计数。"),
    ("T10", "权限控制", "存在当前用户本人帖子和其他用户帖子。", "分别对本人内容和他人内容执行删除操作，检查权限判断结果。"),
    ("T11", "社交关系", "用户 A、用户 B 均已存在并可登录。", "用户 A 进入用户 B 主页并点击关注，检查关注状态和关系记录。"),
    ("T12", "社交关系", "用户 A 已关注用户 B，用户 B 存在已发布内容。", "用户 A 切换到 Following 信息流，检查是否优先展示已关注作者内容。"),
    ("T13", "通知推送", "用户 A、B 均存在，B 已发布帖子。", "用户 A 对用户 B 的帖子执行点赞，检查用户 B 是否收到点赞通知。"),
    ("T14", "通知推送", "用户 A、B 均存在，B 已发布帖子。", "用户 A 对用户 B 的帖子发表评论，检查用户 B 是否收到评论通知。"),
    ("T15", "私信通信", "用户 A、B 均已存在并可进入私信会话。", "用户 A 向用户 B 发送私信内容，检查双方会话展示和消息记录。"),
    ("T16", "消息状态", "用户收到通知或私信后尚未读取。", "打开通知列表或私信会话，并执行查看操作，检查未读状态变化。"),
    ("T17", "推荐排序", "普通用户已登录，系统存在可推荐内容。", "进入首页并访问 For You 个性化推荐流，检查排序结果和策略徽标。"),
    ("T18", "推荐排序", "用户已对某类标签内容产生浏览、点赞或评论行为。", "刷新 For You 推荐流，观察同类标签内容排序和评分变化。"),
    ("T19", "推荐排序", "候选集中同时存在新内容和较旧内容。", "对比新旧内容在推荐流中的评分表现，检查时间衰减影响。"),
    ("T20", "负反馈", "用户已登录，目标内容或目标作者存在。", "对内容选择不感兴趣，或屏蔽对应作者后刷新推荐流。"),
    ("T21", "协同过滤", "测试数据中存在与当前用户行为相似的用户。", "构造相似用户点赞内容后刷新推荐流，检查协同过滤加成。"),
    ("T22", "推荐多样性", "候选集中存在同一作者多篇内容。", "查看推荐结果中同作者内容连续出现情况，观察作者多样性控制效果。"),
    ("T23", "算法对比", "系统存在推荐结果和时间倒序结果。", "打开 CompareView 页面查看推荐流与时间流的并列对比指标。"),
    ("T24", "参数调节", "CompareView 页面可访问，WeightTuner 组件可用。", "调整推荐权重并提交，观察 /api/compare/tuned 输出变化。"),
    ("T25", "用户画像", "用户已产生浏览、点赞、评论或搜索行为。", "打开个人画像页面或画像详情组件，检查词云、统计字段和近期兴趣序列。"),
    ("T26", "后台管理", "管理员已登录，后台策略接口可访问。", "在后台将推荐策略在 traditional 与 ai 之间切换，检查首页徽标和接口返回。"),
    ("T27", "AI 推荐", "Ollama 服务可用，AI 推荐策略已开启。", "切换 AI 推荐模式并请求推荐流，检查 AI 排序或推荐理由。"),
    ("T28", "AI 降级", "AI 推荐策略开启后，关闭或模拟 Ollama 服务不可用。", "再次请求 AI 推荐流，检查系统是否回退到传统推荐策略。"),
    ("T29", "AI 打标", "内容库中存在未打标或标签不足的帖子。", "管理员执行 AI 批量打标或单条打标操作，检查内容标签补全结果。"),
    ("T30", "后台管理", "管理员已登录后台。", "进入用户管理页面，查看用户列表或修改用户状态，检查权限和数据更新。"),
    ("T31", "数据导入", "管理员已登录，导入接口可访问。", "执行公开数据集批量导入操作，检查新增内容和内置内容补齐情况。"),
    ("T32", "广告分发", "广告配置已启用，首页存在足够信息流内容。", "普通用户浏览首页 For You 信息流，检查广告卡片插入位置。"),
    ("T33", "广告统计", "信息流中存在可点击广告。", "点击广告卡片或广告跳转入口，检查展示量和点击量统计。"),
    ("T34", "权限控制", "普通用户已登录，后台接口或入口存在权限限制。", "普通用户尝试访问后台入口或后台接口，检查前后端权限限制。"),
]


result_tables = [
    (
        "表 5-2 用户注册登录测试执行结果",
        [
            ("T01", "用户注册", "输入用户名和密码并提交注册请求。", "注册成功，返回用户信息和 Token。", "通过"),
            ("T02", "用户登录", "使用正确账号和密码登录系统。", "登录成功并进入首页。", "通过"),
            ("T03", "登录状态保持", "登录后刷新页面或重新进入系统。", "用户状态保持，页面可正常访问。", "通过"),
            ("T04", "角色识别", "使用普通用户和管理员账号分别登录。", "系统能够区分普通用户和管理员。", "通过"),
        ],
    ),
    (
        "表 5-3 内容发布与互动功能测试执行结果",
        [
            ("T05", "文字发帖", "输入文字内容并点击发布。", "首页出现新帖子。", "通过"),
            ("T06", "图片发帖", "上传图片后发布帖子。", "图片地址保存并在页面展示。", "通过"),
            ("T07", "评论功能", "对指定帖子发表评论。", "评论内容显示，评论数增加。", "通过"),
            ("T08", "点赞点踩", "对帖子执行点赞或点踩操作。", "对应计数变化，并记录用户行为。", "通过"),
            ("T09", "转发引用", "转发或引用已有帖子。", "生成关联内容并更新转发计数。", "通过"),
            ("T10", "删除或权限判断", "对本人或他人内容执行删除操作。", "本人内容可删除，非本人内容受限制。", "通过"),
        ],
    ),
    (
        "表 5-4 社交关系与实时通信测试执行结果",
        [
            ("T11", "关注用户", "用户 A 关注用户 B。", "关注状态变为已关注。", "通过"),
            ("T12", "关注流展示", "切换到 Following 信息流。", "仅展示或优先展示已关注作者发布的内容。", "通过"),
            ("T13", "点赞通知", "用户 A 点赞用户 B 的帖子。", "用户 B 收到点赞通知。", "通过"),
            ("T14", "评论通知", "用户 A 评论用户 B 的帖子。", "用户 B 收到评论通知。", "通过"),
            ("T15", "私信发送", "用户 A 向用户 B 发送私信。", "会话双方能够看到消息。", "通过"),
            ("T16", "未读状态", "用户收到通知或私信后查看列表。", "未读数正确变化，可标记已读。", "通过"),
        ],
    ),
    (
        "表 5-5 推荐排序与用户画像测试执行结果",
        [
            ("T17", "推荐流访问", "用户进入首页访问个性化推荐流。", "返回带排序结果的信息流。", "通过"),
            ("T18", "标签匹配", "用户点赞某类标签内容后刷新推荐。", "同类标签内容推荐得分提高。", "通过"),
            ("T19", "时间衰减", "对比新旧内容推荐得分。", "新内容获得更高新鲜度优势。", "通过"),
            ("T20", "负反馈过滤", "对内容选择不感兴趣或屏蔽作者。", "对应内容或作者不再进入推荐候选。", "通过"),
            ("T21", "协同过滤", "构造相似用户行为数据。", "相似用户喜欢的内容获得加成。", "通过"),
            ("T22", "作者多样性", "候选集中存在同作者多篇内容。", "同作者后续内容被适当降权。", "通过"),
            ("T23", "推荐对比", "打开 CompareView 页面。", "展示推荐流与时间流对比指标。", "通过"),
            ("T24", "参数调节", "使用 WeightTuner 调整推荐权重。", "调用接口后推荐指标发生变化。", "通过"),
            ("T25", "用户画像", "打开个人画像页面。", "展示词云、雷达图、近期兴趣序列与偏好统计。", "通过"),
        ],
    ),
    (
        "表 5-6 AI 模块、广告分发与后台管理测试执行结果",
        [
            ("T26", "推荐策略切换", "管理员切换 traditional / ai 策略。", "策略切换成功，Feed 徽标同步变化。", "通过"),
            ("T27", "AI 推荐模式", "Ollama 在线时切换 AI 推荐。", "返回 AI 排序结果或推荐理由。", "通过"),
            ("T28", "AI 降级", "关闭 Ollama 后请求 AI 推荐。", "系统降级为传统混合推荐策略。", "通过"),
            ("T29", "AI 自动打标", "对帖子执行 AI 打标操作。", "系统生成内容标签。", "通过"),
            ("T30", "后台用户管理", "管理员查看或管理用户列表。", "用户信息能够正常显示或操作。", "通过"),
            ("T31", "Kaggle 数据导入", "管理员批量导入公开数据集帖文。", "内容库新增帖子或由内置内容库补齐。", "通过"),
            ("T32", "广告插入", "浏览首页信息流。", "广告按配置间隔插入。", "通过"),
            ("T33", "广告统计", "点击信息流广告。", "点击次数和展示统计更新。", "通过"),
            ("T34", "角色权限", "普通用户访问后台入口或接口。", "前端隐藏入口或后端拒绝访问。", "通过"),
        ],
    ),
]


doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.6)
sec.right_margin = Cm(2.4)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

add_heading(doc, "第5章  系统测试及使用说明", 1)

add_heading(doc, "5.1 系统测试报告", 2)
add_heading(doc, "5.1.1 测试的目的及原则", 3)
add_paragraph(
    doc,
    "系统测试的主要目的，是检查系统在当前开发环境下是否能够按照需求完成用户认证、内容发布、互动反馈、实时通知、推荐排序、AI 增强、广告分发和后台管理等核心功能，并验证前后端接口、数据库写入和页面展示之间是否保持一致。测试过程不仅关注单个功能是否可用，也关注多个模块连续调用时的数据流转是否顺畅，例如用户互动行为能否被记录、行为记录能否参与用户画像生成、画像结果能否进一步影响推荐排序。",
)
add_paragraph(
    doc,
    "测试原则上采用“场景驱动、结果可核验、异常可回退”的方式开展。场景驱动是指测试用例围绕普通用户和管理员的真实操作路径设计；结果可核验是指通过页面显示、接口响应、数据库记录和统计图表共同判断功能结果；异常可回退是指重点检查 AI 服务不可用、外部数据不足、普通用户越权访问等情况下系统是否能够保持主流程可用。",
)

add_heading(doc, "5.1.2 测试用例设计与执行结果", 3)
add_paragraph(
    doc,
    "系统测试采用功能场景测试、接口联调验证、推荐对比验证和统计图验证相结合的方式进行。考虑到当前项目尚未建立完整的后端单元测试目录和前端组件测试脚本，本文不将测试结论表述为持续集成自动通过，而是以手工执行记录、接口响应结果、页面运行截图和统计图表作为验证依据。",
)
add_paragraph(
    doc,
    "为更清晰体现测试过程，本文在原有功能测试结果表基础上补充核心测试用例设计表。测试用例设计表用于说明测试对象、前置条件、测试数据和操作步骤；测试执行结果表用于记录预期结果、实际结果和最终结论。两类表格通过统一的用例编号 T01、T02 等建立对应关系，便于后续复查和答辩说明。",
)

add_three_line_table(
    doc,
    "表 5-1 核心测试用例设计补充表",
    ["用例编号", "测试模块", "前置条件与测试数据", "测试步骤与验证点"],
    test_cases,
    [1.8, 2.7, 5.6, 6.1],
    center_cols={0, 1},
)

add_paragraph(
    doc,
    "用户注册登录测试主要验证系统基础认证链路是否正常，包括新用户注册、已有用户登录、登录成功后 Token 返回、前端页面跳转、用户信息展示和角色识别等内容。测试过程中先在注册页面输入用户名和密码提交请求，随后使用正确账号登录系统，并分别使用普通用户和管理员账号检查权限差异。",
)
add_three_line_table(
    doc,
    result_tables[0][0],
    ["编号", "测试内容", "测试过程", "预期结果", "测试结果"],
    result_tables[0][1],
    [1.4, 2.8, 4.7, 5.5, 1.8],
    center_cols={0, 1, 4},
)
add_paragraph(
    doc,
    "测试结果表明，系统能够完成用户注册、登录和角色识别等基础功能，登录成功后前端能够正确保存用户状态，为内容发布、互动操作和后台管理提供身份基础。",
)
add_fig_placeholder(doc, "图5-1 登录功能接口测试结果截图", "放置注册、登录或 Token 返回的接口测试截图")

add_paragraph(
    doc,
    "内容发布与互动功能测试主要验证用户能否正常发布文字帖子、上传图片、评论、点赞、点踩、转发和引用。测试过程中，先使用普通用户登录系统，在首页输入文本内容并发布帖子，检查首页是否出现新内容；随后对帖子执行多种互动操作，观察页面计数、数据库行为记录和通知触发情况是否符合预期。",
)
add_three_line_table(
    doc,
    result_tables[1][0],
    ["编号", "测试内容", "测试过程", "预期结果", "测试结果"],
    result_tables[1][1],
    [1.4, 2.8, 4.7, 5.5, 1.8],
    center_cols={0, 1, 4},
)
add_paragraph(
    doc,
    "测试结果表明，系统内容发布与互动链路基本完整。用户发布内容后能够及时在页面展示，点赞、点踩、评论、转发和引用等操作能够更新计数并写入行为记录，为推荐排序和用户画像提供数据来源。",
)
add_fig_placeholder(doc, "图5-2 用户发帖接口测试结果截图", "放置文字发帖或图片发帖接口测试截图")
add_fig_placeholder(doc, "图5-3 点赞接口测试结果截图", "放置点赞、点踩或行为写入结果截图")
add_fig_placeholder(doc, "图5-4 评论接口测试结果截图", "放置评论提交和评论列表展示截图")

add_paragraph(
    doc,
    "社交关系与实时通信测试主要验证关注、通知和私信功能。测试过程中使用两个不同用户账号，先由用户 A 关注用户 B，检查关注关系是否写入；随后用户 A 对用户 B 的帖子进行点赞或评论，观察用户 B 是否收到通知；最后测试用户之间发送私信，验证会话记录和未读状态是否能够正常更新。",
)
add_three_line_table(
    doc,
    result_tables[2][0],
    ["编号", "测试内容", "测试过程", "预期结果", "测试结果"],
    result_tables[2][1],
    [1.4, 2.8, 4.7, 5.5, 1.8],
    center_cols={0, 1, 4},
)
add_paragraph(
    doc,
    "测试结果表明，系统能够正确维护用户之间的关注关系，并能够在点赞、评论和私信等场景下生成通知或消息记录。在线情况下，消息能够及时展示；离线或刷新页面后，也可以通过通知列表和私信列表进行补偿查看。",
)

add_paragraph(
    doc,
    "推荐排序与用户画像测试主要验证系统能否根据用户行为生成个性化推荐结果。测试过程中，先让用户对某一类标签内容进行浏览、点赞、评论或搜索，然后刷新推荐流，观察同类标签内容是否获得更高推荐分；再通过算法对比页面查看推荐流与时间倒序流的排序差异，并观察评分拆分对象中标签匹配、文本相似度、协同过滤、时间衰减、热门话题、多样性指标等分项是否产生变化。",
)
add_three_line_table(
    doc,
    result_tables[3][0],
    ["编号", "测试内容", "测试过程", "预期结果", "测试结果"],
    result_tables[3][1],
    [1.4, 2.8, 4.7, 5.5, 1.8],
    center_cols={0, 1, 4},
)
add_paragraph(
    doc,
    "测试结果表明，系统推荐模块能够根据用户行为、内容标签、时间因素、协同过滤和负反馈信号对信息流进行排序。通过推荐流与时间流对比，可以观察到个性化推荐结果与单纯时间倒序结果存在差异；作者去重率和标签覆盖率等指标也能辅助说明推荐列表在多样性方面的表现。",
)
add_fig_placeholder(doc, "图5-5 推荐信息流接口测试结果截图", "放置 For You 推荐流接口或页面展示截图")
add_fig_placeholder(doc, "图5-6 推荐流与时间流对比接口测试结果截图", "放置 CompareView 指标、推荐对比或多样性指标截图")

add_paragraph(
    doc,
    "AI 模块、广告分发与后台管理测试主要验证系统扩展功能是否能够正常运行。测试过程中，管理员进入后台切换推荐策略，观察首页信息流徽标和排序结果是否变化；在 Ollama 本地模型正常运行时切换 AI 推荐模式，检查系统是否能够返回 AI 排序或推荐理由；关闭 Ollama 后再次请求 AI 推荐，验证系统是否能够自动降级。广告测试则通过浏览首页观察广告是否按配置间隔插入，并点击广告检查统计数据是否更新。",
)
add_three_line_table(
    doc,
    result_tables[4][0],
    ["编号", "测试内容", "测试过程", "预期结果", "测试结果"],
    result_tables[4][1],
    [1.4, 2.8, 4.7, 5.5, 1.8],
    center_cols={0, 1, 4},
)
add_paragraph(
    doc,
    "测试结果表明，系统 AI 模块、广告模块和后台管理功能能够正常运行。AI 推荐在模型可用时能够参与排序或生成推荐理由，在模型不可用时能够降级到传统推荐策略，保证信息流主链路不受影响。广告模块能够按照用户画像和广告配置完成信息流插入，并记录展示和点击数据。",
)
add_fig_placeholder(doc, "图5-7 AI 推荐策略切换接口测试结果截图", "放置推荐策略切换、Ollama 状态或 AI 推荐返回结果截图")
add_fig_placeholder(doc, "图5-8 广告分发接口测试结果截图", "放置信息流广告插入或广告点击接口截图")
add_fig_placeholder(doc, "图5-9 广告管理后台统计接口测试结果截图", "放置广告展示量、点击量或后台统计截图")

add_heading(doc, "5.1.3 测试总结", 3)
add_paragraph(
    doc,
    "通过对系统主要功能、推荐排序、AI 模块、广告分发和后台管理的测试可以看出，系统整体运行较为稳定，各主要功能模块基本能够按照预期完成相应操作。用户注册登录、内容发布、点赞评论、转发引用、关注、搜索、通知、私信、用户画像、推荐流展示、广告插入和后台管理等功能均通过测试，说明系统已经形成了较完整的业务闭环。",
)
add_paragraph(
    doc,
    "在功能测试方面，普通用户能够顺利完成从注册登录到内容浏览、发布、互动和消息接收的完整流程。用户在浏览、点赞、评论、转发、搜索等过程中产生的行为数据能够被系统记录，并作为用户画像和推荐排序的数据来源。通知与私信功能能够在用户互动后生成对应消息记录，并在页面端进行展示，基本满足社交媒体系统对实时反馈的需求。",
)
add_paragraph(
    doc,
    "在推荐算法测试方面，系统能够根据用户行为、内容标签、文本相似度、时间衰减、热门话题、协同过滤和负反馈等因素生成个性化推荐结果。通过推荐流与时间倒序流的对比可以看出，推荐流在标签匹配、平均推荐分和内容相关性方面通常优于单纯按发布时间排序的信息流，说明基于用户行为分析的混合排序策略具有一定有效性。同时，评分拆分结果能够展示各推荐因子对最终排序的影响，提高了推荐结果的可解释性。",
)
add_paragraph(
    doc,
    "在 AI 模块测试方面，当本地模型服务正常运行时，系统能够完成 AI 推荐或推荐理由生成；当模型服务不可用或响应异常时，系统能够降级为传统推荐策略，保证信息流主流程仍可正常使用。该测试结果说明系统在引入 AI 能力的同时，保留了稳定的传统推荐路径，具有一定的容错能力和可用性。",
)
add_paragraph(
    doc,
    "在广告分发与后台管理测试方面，系统能够根据广告配置在信息流中插入广告，并记录广告展示和点击数据。管理员能够完成推荐策略切换、用户管理、数据导入和广告管理等操作，说明后台管理模块能够对系统运行和推荐策略进行基本控制。",
)
add_paragraph(
    doc,
    "综上所述，系统主要功能测试结果符合预期，推荐排序模块能够体现用户行为分析对信息流排序的作用，AI 推荐和广告分发等扩展模块也能够正常运行。但当前测试仍以手工功能测试、页面运行结果和接口响应结果为主，尚未建立完善的后端单元测试、前端组件测试、自动化 UI 测试和高并发压力测试。因此，系统在测试覆盖范围、性能量化分析和自动化验证方面仍有进一步完善空间。",
)

add_heading(doc, "5.1.4 测试数据构造与实验建议", 3)
add_paragraph(
    doc,
    "推荐算法的输出结果容易受到数据分布影响。如果系统中的帖子题材较为单一，或者用户行为数据较少，协同过滤、TF-IDF 相似度和多样性指标的差异就不容易体现。测试前可以先准备一组分层数据：用户侧划分若干类兴趣画像，让不同用户分别集中浏览或点赞科技、新闻、生活、体育等不同主题；内容侧保证每类标签下既有新内容，也有较旧内容，并包含不同作者发布的帖子，以便观察时间衰减、作者多样性和标签覆盖率等指标。",
)
add_paragraph(
    doc,
    "对照实验时，建议先固定随机种子，或者暂时关闭推荐抖动因子中的随机部分，这样反复刷新时更容易判断排序结果是否稳定；随后再打开抖动，观察推荐列表的多样性是否有所改善。对比不同策略的平均推荐分、标签命中率、TF-IDF 均值、作者去重率和标签覆盖率时，可以截取 CompareView 页面或接口返回结果作为实验记录。",
)
add_paragraph(
    doc,
    "搜索功能可使用与帖子标签相匹配的检索词进行验证，观察用户画像中是否出现与搜索意图相关的权重变化。负反馈功能可通过屏蔽作者或选择不感兴趣进行验证，刷新 For You 信息流后检查对应作者或内容是否被排除在候选集合之外。",
)

add_heading(doc, "5.2 系统使用说明", 2)
add_heading(doc, "5.2.1 系统简介", 3)
add_paragraph(
    doc,
    "本系统是一套面向社交媒体场景的实时信息流排序与分发原型系统。普通用户可以完成注册登录、浏览 For You 推荐流和 Following 关注流、发布帖子、评论、点赞、点踩、转发、引用、关注、搜索、查看通知、发送私信和查看画像等操作；管理员可以进入后台管理页面，完成用户管理、推荐策略切换、AI 批量打标、外部数据导入、广告配置和平台统计查看等操作。",
)

add_heading(doc, "5.2.2 系统运行环境", 3)
env_rows = [
    ("操作系统", "Windows", "用于本地开发、调试和演示。"),
    ("后端运行环境", "JDK 17、Maven", "启动 Spring Boot 后端服务，默认端口为 8888。"),
    ("前端运行环境", "Node.js、npm", "启动 Vue 3 与 Vite 前端项目，默认端口为 5173。"),
    ("数据库", "MySQL rec_db", "存储用户、内容、行为、标签、通知、私信和广告数据。"),
    ("AI 服务", "Ollama、DeepSeek API", "支持 AI 内容打标、AI 推荐理由生成和助手功能。"),
]
add_three_line_table(
    doc,
    "表 5-7 系统运行环境配置",
    ["类别", "环境或工具", "说明"],
    env_rows,
    [3.3, 4.4, 8.5],
    center_cols={0, 1},
)

add_heading(doc, "5.2.3 系统操作说明", 3)
add_paragraph(
    doc,
    "普通用户通过注册或登录进入系统。登录后，用户可以在首页切换个性化推荐流和关注流，浏览不同来源的内容；浏览、点赞、评论、转发、搜索、点踩等操作会被记录为行为事件，这些行为数据后续会用于更新用户画像，并参与推荐排序计算。",
)
add_paragraph(
    doc,
    "发布内容时，用户需要填写帖子正文，也可以按需上传图片。帖子保存后，系统会解析正文中包含的话题标签，并在需要时借助 AI 打标服务补充语义标签。其他用户对帖子点赞或评论后，系统会生成对应通知记录，并通过 WebSocket 推送未读提醒，使用户能够及时获取互动反馈。",
)
add_paragraph(
    doc,
    "管理员进入后台后，可以查看用户信息和平台统计数据，也可以在传统推荐和 AI 推荐两种策略之间进行切换。同时，后台还提供 Kaggle 数据导入、AI 批量打标、广告配置维护和广告统计查看等功能，用于支撑系统演示和运行维护。",
)
add_paragraph(
    doc,
    "算法验证方面，用户或管理员可以进入推荐效果对比页面，查看推荐流与时间流之间的差异，也可以查看评分拆分、权重调节组件、排序管道和多样性指标。当 AI 服务不可用时，系统会自动回退到传统混合推荐策略，从而保证信息流主链路仍能正常使用。",
)

doc.save(OUT)
print(OUT)
