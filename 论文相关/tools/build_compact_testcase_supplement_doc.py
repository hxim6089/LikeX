from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = r"D:\TheEnd\qnyproj-main\recommendation-system\论文相关\表5-6核心测试用例补充表_简洁版.docx"


def set_run_font(run, east="宋体", west="Times New Roman", size=12, bold=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, east="楷体", size=10.5)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_border(cell, edge, sz="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    node = borders.find(qn("w:" + edge))
    if node is None:
        node = OxmlElement("w:" + edge)
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), sz)
    node.set(qn("w:space"), "0")
    node.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(int(w * 567) for w in widths_cm)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 567)))


def fill_cell(cell, text, center=False, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=bold)


rows = [
    ("T01", "用户认证", "用户注册与登录", "完成用户创建、Token 返回、首页跳转与登录状态保持。"),
    ("T02", "权限控制", "普通用户与管理员角色识别", "系统能够区分不同角色，并限制普通用户访问后台功能。"),
    ("T03", "内容发布", "文字和图片帖子发布", "帖子保存成功，正文、图片地址和话题标签能够正常展示。"),
    ("T04", "内容互动", "评论、点赞、点踩和转发", "互动计数变化正确，相关行为记录写入数据库。"),
    ("T05", "社交关系", "关注用户与关注流展示", "关注状态更新成功，关注流能够展示已关注作者内容。"),
    ("T06", "通知私信", "通知生成与私信接收", "点赞、评论和私信操作能够生成消息记录并更新未读状态。"),
    ("T07", "行为采集", "浏览、互动和搜索行为记录", "系统能够记录 VIEW、LIKE、COMMENT、SEARCH 等行为数据。"),
    ("T08", "推荐排序", "个性化推荐流生成", "推荐流能够结合用户行为、标签和时间因素返回排序结果。"),
    ("T09", "推荐解释", "评分拆分与对比指标展示", "页面能够展示推荐分项、推荐流与时间流对比结果。"),
    ("T10", "用户画像", "兴趣标签与近期兴趣序列", "画像页能够展示词云、雷达图、偏好统计和兴趣变化。"),
    ("T11", "AI 模块", "AI 打标、AI 推荐与异常降级", "模型可用时返回 AI 结果，异常时回退传统推荐策略。"),
    ("T12", "广告分发", "广告插入与点击统计", "广告按配置插入信息流，并记录展示量和点击量。"),
    ("T13", "后台管理", "策略切换、用户管理和数据导入", "管理员能够完成推荐策略切换、用户维护和数据导入。"),
    ("T14", "数据统计", "行为类型分布与内容标签分布", "统计页能够展示平台行为分布、内容生态和标签覆盖情况。"),
]

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.6)
section.right_margin = Cm(2.4)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.font.size = Pt(12)

add_caption(doc, "表 5-6 核心测试用例补充表")
table = doc.add_table(rows=1, cols=4)
set_table_width(table, [1.8, 3.0, 4.6, 6.8])
repeat_header(table.rows[0])
headers = ["用例编号", "测试模块", "测试内容", "预期结果"]
for idx, text in enumerate(headers):
    clear_cell_borders(table.rows[0].cells[idx])
    fill_cell(table.rows[0].cells[idx], text, center=True, bold=True)

for data in rows:
    row = table.add_row()
    for c in row.cells:
        clear_cell_borders(c)
    for idx, text in enumerate(data):
        fill_cell(row.cells[idx], text, center=(idx in (0, 1)))

for cell in table.rows[0].cells:
    set_cell_border(cell, "top", sz="12")
    set_cell_border(cell, "bottom", sz="8")
for cell in table.rows[-1].cells:
    set_cell_border(cell, "bottom", sz="12")

doc.save(OUT)
print(OUT)
