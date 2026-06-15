from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = r"D:\TheEnd\qnyproj-main\recommendation-system\论文相关\表3-13_术语与缩写.docx"


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, east="楷体", size=10.5)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_border(cell, edge, sz="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    node = borders.find(qn("w:" + edge))
    if node is None:
        node = OxmlElement("w:" + edge)
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), sz)
    node.set(qn("w:space"), "0")
    node.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=65, start=80, bottom=65, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(int(w * 567) for w in widths_cm)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 567)))


def fill_cell(cell, text, center=False, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


headers = ["术语或符号", "含义说明"]
rows = [
    ("In-Network", "网内候选，主体来自当前用户关注对象发布的内容。"),
    ("Out-of-Network", "网外候选，主体来自全站公开内容，用于兴趣探索。"),
    ("TF-IDF", "词频—逆文档频率，用于构建文本向量并与兴趣向量计算相似度。"),
    ("CTR", "点击率，click/impression，用于广告质量评估。"),
    ("JWT", "JSON Web Token，一种紧凑的令牌格式，用于无状态会话。"),
    ("REST", "表述性状态转移风格的 HTTP API 设计范式。"),
    ("STOMP", "面向消息的文本协议，常与 WebSocket 组合实现订阅推送。"),
    ("DTO", "数据传输对象，封装接口响应字段。"),
    ("ScoreBreakdown", "推荐得分拆分对象，记录各分项贡献。"),
    ("DynamicWeights", "动态权重集合，由画像阶段输出并作用于打分公式。"),
    ("Ollama", "本地大模型运行环境，本课题用于打标与 AI 推荐。"),
    ("冷启动", "新用户或新内容缺少行为历史时的推荐难题。"),
    ("探索—利用", "Exploration-Exploitation，在尝试新内容与巩固已知兴趣之间权衡。"),
]

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.6)
section.right_margin = Cm(2.4)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.font.size = Pt(10.5)

add_caption(doc, "表 3-13 术语与缩写")
table = doc.add_table(rows=1, cols=2)
set_table_width(table, [4.5, 11.7])
repeat_header(table.rows[0])

for idx, text in enumerate(headers):
    clear_cell_borders(table.rows[0].cells[idx])
    fill_cell(table.rows[0].cells[idx], text, center=True, bold=True)

for data in rows:
    row = table.add_row()
    for c in row.cells:
        clear_cell_borders(c)
    fill_cell(row.cells[0], data[0], center=True)
    fill_cell(row.cells[1], data[1], center=False)

for cell in table.rows[0].cells:
    set_cell_border(cell, "top", sz="12")
    set_cell_border(cell, "bottom", sz="8")
for cell in table.rows[-1].cells:
    set_cell_border(cell, "bottom", sz="12")

doc.save(OUT)
print(OUT)
