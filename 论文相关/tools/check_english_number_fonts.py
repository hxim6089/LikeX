import io
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOC_PATH = Path(r"D:\SHARE\OneDrive\Desktop\毕业论文 初稿07.2.docx")


ASCII_RE = re.compile(r"[A-Za-z0-9]")


def rfonts_from_rpr(rpr):
    if rpr is None:
        return {}
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return {}
    return {
        "ascii": rfonts.get(qn("w:ascii")),
        "hAnsi": rfonts.get(qn("w:hAnsi")),
        "eastAsia": rfonts.get(qn("w:eastAsia")),
        "cs": rfonts.get(qn("w:cs")),
    }


def style_fonts(style):
    if style is None:
        return {}
    try:
        return rfonts_from_rpr(style.element.rPr)
    except Exception:
        return {}


def run_fonts(run):
    return rfonts_from_rpr(run._element.rPr)


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def effective_ascii_font(run, paragraph):
    direct = run_fonts(run)
    if direct.get("ascii") or direct.get("hAnsi"):
        return direct.get("ascii") or direct.get("hAnsi"), "run", direct
    sfonts = style_fonts(paragraph.style)
    if sfonts.get("ascii") or sfonts.get("hAnsi"):
        return sfonts.get("ascii") or sfonts.get("hAnsi"), "style", sfonts
    return None, "missing", {}


def main():
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    doc = Document(str(DOC_PATH))
    total = 0
    ok = 0
    missing = 0
    bad = []
    source_counter = Counter()
    font_counter = Counter()

    for p_idx, p in enumerate(iter_paragraphs(doc), start=1):
        for r in p.runs:
            text = r.text
            if not text or not ASCII_RE.search(text):
                continue
            total += 1
            font, source, detail = effective_ascii_font(r, p)
            source_counter[source] += 1
            font_counter[font or "<missing>"] += 1
            if font and font.lower().replace(" ", "") == "timesnewroman":
                ok += 1
            elif font is None:
                missing += 1
                bad.append((p_idx, text[:80], "<missing>", source, p.text[:140]))
            else:
                bad.append((p_idx, text[:80], font, source, p.text[:140]))

    print("DOC", DOC_PATH)
    print("ASCII_RUNS", total)
    print("TIMES_NEW_ROMAN_RUNS", ok)
    print("MISSING_EFFECTIVE_ASCII_FONT", missing)
    print("BAD_OR_UNKNOWN_RUNS", len(bad))
    print()
    print("FONT_COUNTER")
    for font, count in font_counter.most_common(20):
        print(f"{font}: {count}")
    print()
    print("SOURCE_COUNTER")
    for source, count in source_counter.items():
        print(f"{source}: {count}")
    print()
    print("SAMPLES_BAD_OR_UNKNOWN")
    for item in bad[:60]:
        p_idx, run_text, font, source, para_text = item
        print(f"P{p_idx} font={font} source={source} run={run_text!r} para={para_text!r}")


if __name__ == "__main__":
    main()
