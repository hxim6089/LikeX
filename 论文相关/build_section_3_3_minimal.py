from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path("thesis_072_source.docx")
OUT_DOCX = Path("3.3数据库设计_最小改动标注版.docx")
OUT_MD = Path("3.3局部ER图_PlantUML代码.md")


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_normal(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_font(r)
    return p


def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 14 if level == 1 else 12, True)
    return p


def add_marker(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, 10.5, True)
    return p


def copy_table(doc, src_table):
    rows = [[cell.text.strip() for cell in row.cells] for row in src_table.rows]
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri == 0:
                set_cell_shading(cells[ci], "F1F5F9")
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 or ci < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(val)
            set_font(r, 9.5, ri == 0)
    doc.add_paragraph()


def iter_blocks(doc):
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


table_notes = {
    "表 3-2 用户表设计": "用户表用于保存用户基础资料、登录认证信息和角色状态，是内容发布、行为记录、关注关系、消息通知及后台管理等功能的基础数据表。",
    "表 3-3 内容表设计": "内容表用于保存帖子、评论、转发和引用等内容数据，是信息流展示和推荐排序的核心数据表，其中浏览数、点赞数、评论数等字段用于提升列表展示和推荐打分效率。",
    "表 3-4 行为表设计": "行为表用于记录用户浏览、点赞、点踩、快速滑过和搜索等交互行为，为用户画像构建、兴趣分析和推荐排序提供基础行为信号。",
    "表 3-5 关注表设计概要": "关注表用于维护用户之间的有向关注关系，同时为推荐算法中的网内候选内容召回提供数据依据。",
    "表 3-6 标签实体与内容标签关联（逻辑结构）": "标签表和内容标签关联表用于表达内容与标签之间的多对多关系，支撑话题检索、热门标签统计、用户兴趣画像和推荐匹配。",
    "表 3-7 通知表设计概要": "通知表用于记录点赞、评论、关注、转发和引用等事件触发的系统提醒，并通过已读状态支持前端通知列表和未读提示。",
    "表 3-8 广告表设计概要": "广告表用于保存信息流广告的素材、定向标签、出价和展示点击统计，系统可结合用户画像和广告配置完成广告排序与插入。",
    "表 3-9 负反馈信号表（negative_signals）": "负反馈信号表用于保存用户对内容或作者的不感兴趣、屏蔽和静音等反向反馈，推荐流程会据此过滤候选内容。",
    "表 3-10 私信消息表（tb_message）": "私信消息表用于保存用户之间的点对点聊天记录，并与 WebSocket 通信机制配合支持实时消息、历史消息和未读状态展示。",
}


plantuml_md = r"""# 3.3 局部 E-R 图 PlantUML 代码

> 使用方式：复制对应代码到 PlantUML 渲染工具，导出 PNG 后插入到 Word 文档中标注的位置。

## 图 3-3（a）用户与内容局部 E-R 图

```plantuml
@startuml
hide circle
skinparam linetype ortho
skinparam shadowing false
skinparam entity {
  BackgroundColor White
  BorderColor #1f2937
}

entity "用户 User" as user {
  * id : BIGINT <<PK>>
  --
  username : VARCHAR
  handle : VARCHAR
  role : VARCHAR
  banned : BOOLEAN
}

entity "内容 Content" as content {
  * id : BIGINT <<PK>>
  --
  author_id : BIGINT <<FK>>
  parent_id : BIGINT
  repost_of_id : BIGINT
  quote_of_id : BIGINT
  content : TEXT
  image_url : VARCHAR
  category : VARCHAR
  created_at : DATETIME
  view_count : INT
  like_count : INT
  comment_count : INT
  repost_count : INT
}

user ||--o{ content : 发布
content ||--o{ content : 评论/回复
content ||--o{ content : 转发来源
content ||--o{ content : 引用来源
@enduml
```

## 图 3-3（b）行为画像与标签局部 E-R 图

```plantuml
@startuml
hide circle
skinparam linetype ortho
skinparam shadowing false
skinparam entity {
  BackgroundColor White
  BorderColor #1f2937
}

entity "用户 User" as user {
  * id : BIGINT <<PK>>
  --
  username : VARCHAR
  custom_weights : TEXT
}

entity "内容 Content" as content {
  * id : BIGINT <<PK>>
  --
  author_id : BIGINT
  content : TEXT
  category : VARCHAR
}

entity "行为 Behavior" as behavior {
  * id : BIGINT <<PK>>
  --
  user_id : BIGINT
  content_id : BIGINT
  type : VARCHAR
  duration : INT
  created_at : DATETIME
}

entity "标签 Tag" as tag {
  * id : BIGINT <<PK>>
  --
  name : VARCHAR
}

entity "内容标签 content_tags" as content_tags {
  * content_id : BIGINT <<FK>>
  * tag_id : BIGINT <<FK>>
}

entity "负反馈 NegativeSignal" as negative {
  * id : BIGINT <<PK>>
  --
  user_id : BIGINT
  target_type : VARCHAR
  target_id : BIGINT
  signal_type : VARCHAR
  created_at : DATETIME
}

user ||--o{ behavior : 产生
content ||--o{ behavior : 被交互
content ||--o{ content_tags : 绑定
tag ||--o{ content_tags : 关联
user ||--o{ negative : 发起
content ||--o{ negative : 可作为目标
@enduml
```

## 图 3-3（c）社交关系与消息通知局部 E-R 图

```plantuml
@startuml
hide circle
skinparam linetype ortho
skinparam shadowing false
skinparam entity {
  BackgroundColor White
  BorderColor #1f2937
}

entity "用户 User" as user {
  * id : BIGINT <<PK>>
  --
  username : VARCHAR
  handle : VARCHAR
  role : VARCHAR
}

entity "关注 Follow" as follow {
  * id : BIGINT <<PK>>
  --
  follower_id : BIGINT
  followee_id : BIGINT
  created_at : DATETIME
}

entity "通知 Notification" as notification {
  * id : BIGINT <<PK>>
  --
  recipient_id : BIGINT
  actor_id : BIGINT
  type : VARCHAR
  entity_id : BIGINT
  is_read : BOOLEAN
  created_at : DATETIME
}

entity "私信 Message" as message {
  * id : BIGINT <<PK>>
  --
  sender_id : BIGINT
  recipient_id : BIGINT
  content : TEXT
  is_read : BOOLEAN
  created_at : DATETIME
}

user ||--o{ follow : follower
user ||--o{ follow : followee
user ||--o{ notification : 接收
user ||--o{ notification : 触发
user ||--o{ message : 发送
user ||--o{ message : 接收
@enduml
```

## 图 3-3（d）广告投放局部 E-R 图

```plantuml
@startuml
hide circle
skinparam linetype ortho
skinparam shadowing false
skinparam entity {
  BackgroundColor White
  BorderColor #1f2937
}

entity "用户画像 UserProfile" as profile {
  * user_id : BIGINT
  --
  interest_tags : JSON
  behavior_stats : JSON
  match_rate : DOUBLE
}

entity "广告 Ad" as ad {
  * id : BIGINT <<PK>>
  --
  title : VARCHAR
  description : TEXT
  image_url : VARCHAR
  target_url : VARCHAR
  target_tags : VARCHAR
  category : VARCHAR
  bid_price : DOUBLE
  impression_count : INT
  click_count : INT
  active : BOOLEAN
}

entity "广告配置 AdConfig" as config {
  * id : BIGINT <<PK>>
  --
  enabled : BOOLEAN
  frequency : INT
  max_ads_per_page : INT
}

entity "信息流广告位 FeedAdSlot" as slot {
  * slot_index : INT
  --
  user_id : BIGINT
  ad_id : BIGINT
  ad_score : DOUBLE
}

profile ||--o{ slot : 画像匹配
ad ||--o{ slot : 投放展示
config ||--o{ slot : 控制频率
@enduml
```
"""


def main():
    src = Document(str(SOURCE))
    out = Document()
    sec = out.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)

    out.styles["Normal"].font.name = "宋体"
    out.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    out.styles["Normal"].font.size = Pt(10.5)

    in_sec = False
    pending_caption = None
    inserted_a = inserted_b = inserted_c = inserted_d = False

    for block in iter_blocks(src):
        if hasattr(block, "text"):
            text = block.text.strip()
            if text.startswith("3.3 数据库设计"):
                in_sec = True
                add_heading(out, text, level=1)
                continue
            if in_sec and (text.startswith("3.4 ") or text.startswith("3.4")):
                break
            if not in_sec:
                continue

            if text.startswith("图3-3") or text.startswith("图 3-3"):
                add_marker(out, "【保留原图位置】此处继续放置原“图3-3 系统 E-R 图”，作为数据库整体关系总览。")
                add_normal(out, text, first_line=False)
                if not inserted_a:
                    add_marker(out, "【新增图3-3（a）插入位置】用户与内容局部 E-R 图。建议放在用户表和内容表之前，用于说明用户发布内容以及内容评论、转发、引用等自关联关系。")
                    inserted_a = True
                continue

            if text.startswith("表 3-4") and not inserted_b:
                add_marker(out, "【新增图3-3（b）插入位置】行为画像与标签局部 E-R 图。建议放在行为表之前，用于说明用户、内容、行为、标签关联和负反馈信号之间的关系。")
                inserted_b = True
            if text.startswith("关注表主要用于") and not inserted_c:
                add_marker(out, "【新增图3-3（c）插入位置】社交关系与消息通知局部 E-R 图。建议放在关注表说明之前，用于说明用户、关注、通知和私信之间的关系。")
                inserted_c = True
            if text.startswith("表 3-8") and not inserted_d:
                add_marker(out, "【新增图3-3（d）插入位置】广告投放局部 E-R 图。建议放在广告表之前，用于说明广告、广告配置、用户画像匹配和信息流广告位之间的关系。")
                inserted_d = True

            if text.startswith("表 3-"):
                pending_caption = text
                if text in table_notes:
                    add_normal(out, table_notes[text])
                add_normal(out, text, first_line=False)
            elif text:
                add_normal(out, text)
        else:
            if in_sec:
                copy_table(out, block)

    out.save(OUT_DOCX)
    OUT_MD.write_text(plantuml_md, encoding="utf-8")
    print(OUT_DOCX.resolve())
    print(OUT_MD.resolve())


if __name__ == "__main__":
    main()
