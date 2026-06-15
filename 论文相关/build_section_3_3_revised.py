from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_DIR = Path("section_3_3_assets")
OUT_DIR.mkdir(exist_ok=True)


def font_path(name):
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return None


FONT = font_path("msyh.ttc")
BOLD = font_path("msyhbd.ttc") or FONT


def f(size, bold=False):
    return ImageFont.truetype(BOLD if bold else FONT, size)


def draw_center(draw, box, text, font, fill=(17, 24, 39), line_gap=8):
    x1, y1, x2, y2 = box
    max_w = x2 - x1 - 24
    lines = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            test = line + ch
            if draw.textbbox((0, 0), test, font=font)[2] <= max_w:
                line = test
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + line_gap


def round_rect(draw, box, fill, outline=(17, 24, 39), width=3, radius=14):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill=(17, 24, 39), width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 16
    pts = [
        (x2, y2),
        (x2 - size * math.cos(angle - 0.45), y2 - size * math.sin(angle - 0.45)),
        (x2 - size * math.cos(angle + 0.45), y2 - size * math.sin(angle + 0.45)),
    ]
    draw.polygon(pts, fill=fill)


def er_box(draw, x, y, w, h, title, fields, fill=(255, 255, 255)):
    round_rect(draw, (x, y, x + w, y + h), fill=fill, outline=(31, 41, 55), width=3, radius=12)
    draw.rounded_rectangle((x, y, x + w, y + 54), radius=12, fill=(241, 245, 249), outline=(31, 41, 55), width=3)
    draw.line((x, y + 54, x + w, y + 54), fill=(31, 41, 55), width=3)
    draw_center(draw, (x, y + 4, x + w, y + 52), title, f(25, True))
    yy = y + 75
    for field in fields:
        draw.text((x + 22, yy), field, font=f(18), fill=(31, 41, 55))
        yy += 31


def relation(draw, p1, p2, label="", left="1", right="N"):
    arrow(draw, p1, p2)
    if left:
        draw.text((p1[0] + 10, p1[1] - 30), left, font=f(18, True), fill=(17, 24, 39))
    if right:
        draw.text((p2[0] - 28, p2[1] - 30), right, font=f(18, True), fill=(17, 24, 39))
    if label:
        mx = (p1[0] + p2[0]) / 2
        my = (p1[1] + p2[1]) / 2
        draw.text((mx - 55, my - 36), label, font=f(17), fill=(55, 65, 81))


def save_diagram(name, title, boxes, rels, size=(1600, 900)):
    img = Image.new("RGB", size, "white")
    d = ImageDraw.Draw(img)
    d.rectangle((35, 35, size[0] - 35, size[1] - 35), outline=(226, 232, 240), width=2)
    draw_center(d, (0, 40, size[0], 100), title, f(34, True))
    for b in boxes:
        er_box(d, *b)
    for r in rels:
        relation(d, *r)
    path = OUT_DIR / name
    img.save(path, quality=96)
    return path


overall = save_diagram(
    "图3-3_系统总体ER图.png",
    "系统总体 E-R 关系总览",
    [
        (80, 160, 260, 220, "用户 User", ["PK id", "username", "role", "banned"]),
        (450, 145, 300, 250, "内容 Content", ["PK id", "FK author_id", "content", "计数字段"]),
        (870, 150, 230, 190, "标签 Tag", ["PK id", "name"]),
        (1240, 165, 250, 190, "广告 Ad", ["PK id", "targetTags", "bidPrice"]),
        (90, 520, 250, 190, "关注 Follow", ["PK id", "followerId", "followeeId"]),
        (455, 500, 285, 220, "负反馈", ["NegativeSignal", "targetType", "targetId", "signalType"]),
        (850, 500, 275, 210, "行为 Behavior", ["PK id", "userId", "contentId", "type"]),
        (1230, 500, 260, 200, "通知/私信", ["Notification", "Message", "isRead"]),
    ],
    [
        ((340, 260), (450, 260), "发布", "1", "N"),
        ((750, 260), (870, 245), "多对多", "N", "N"),
        ((210, 380), (210, 520), "关注", "1", "N"),
        ((340, 320), (455, 600), "负反馈", "1", "N"),
        ((600, 395), (600, 500), "过滤目标", "1", "N"),
        ((750, 335), (850, 600), "被互动", "1", "N"),
        ((1125, 610), (1230, 610), "触发", "N", "N"),
        ((1100, 245), (1240, 250), "匹配投放", "N", "N"),
    ],
)

user_content = save_diagram(
    "图3-3a_用户与内容局部ER图.png",
    "用户与内容局部 E-R 图",
    [
        (120, 200, 280, 250, "用户 User", ["PK id", "username", "role", "createdAt"]),
        (620, 170, 330, 300, "内容 Content", ["PK id", "FK author_id", "parent_id", "repost_of_id", "quote_of_id"]),
        (1120, 210, 280, 230, "内容统计字段", ["viewCount", "likeCount", "commentCount", "repostCount"]),
    ],
    [
        ((400, 310), (620, 310), "发布", "1", "N"),
        ((785, 470), (785, 615), "评论/转发/引用", "1", "N"),
        ((950, 310), (1120, 310), "冗余统计", "1", "1"),
    ],
)

behavior_tag = save_diagram(
    "图3-3b_行为画像与标签局部ER图.png",
    "行为画像与标签局部 E-R 图",
    [
        (90, 180, 260, 230, "用户 User", ["PK id", "username", "customWeights"]),
        (470, 160, 300, 260, "行为 Behavior", ["PK id", "FK userId", "FK contentId", "type", "duration"]),
        (900, 155, 290, 260, "内容 Content", ["PK id", "content", "category", "createdAt"]),
        (1270, 185, 230, 210, "标签 Tag", ["PK id", "name"]),
        (470, 560, 300, 210, "负反馈 NegativeSignal", ["FK user_id", "target_type", "target_id", "signal_type"]),
    ],
    [
        ((350, 295), (470, 295), "产生行为", "1", "N"),
        ((770, 295), (900, 295), "作用于内容", "N", "1"),
        ((1190, 285), (1270, 285), "内容标签", "N", "N"),
        ((350, 640), (470, 640), "反向偏好", "1", "N"),
        ((770, 640), (900, 350), "过滤候选", "N", "1"),
    ],
)

social_message = save_diagram(
    "图3-3c_社交关系与消息通知局部ER图.png",
    "社交关系与消息通知局部 E-R 图",
    [
        (120, 190, 270, 230, "用户 User", ["PK id", "username", "handle", "role"]),
        (610, 120, 300, 230, "关注 Follow", ["PK id", "followerId", "followeeId", "createdAt"]),
        (610, 460, 300, 230, "私信 Message", ["PK id", "sender_id", "recipient_id", "is_read"]),
        (1120, 285, 300, 240, "通知 Notification", ["PK id", "recipientId", "actorId", "type", "entityId"]),
    ],
    [
        ((390, 260), (610, 235), "关注/被关注", "1", "N"),
        ((390, 350), (610, 570), "发送/接收", "1", "N"),
        ((910, 235), (1120, 360), "触发提醒", "N", "N"),
        ((910, 570), (1120, 450), "消息提醒", "N", "N"),
    ],
)

ad_delivery = save_diagram(
    "图3-3d_广告投放局部ER图.png",
    "广告投放局部 E-R 图",
    [
        (100, 200, 280, 230, "用户画像", ["interestTags", "behaviorStats", "matchRate"]),
        (550, 160, 330, 300, "广告 Ad", ["PK id", "title", "targetTags", "bidPrice", "click/impression"]),
        (1030, 195, 310, 240, "广告配置 AdConfig", ["enabled", "frequency", "maxAdsPerPage"]),
        (550, 560, 330, 190, "信息流广告位", ["按间隔插入", "匹配标签", "统计曝光点击"]),
    ],
    [
        ((380, 315), (550, 315), "标签匹配", "1", "N"),
        ((880, 315), (1030, 315), "投放控制", "N", "1"),
        ((715, 460), (715, 560), "生成广告位", "N", "N"),
    ],
)


tables = {
    "表 3-2 用户表设计": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "用户 ID"],
        ["username", "VARCHAR", "用户名，唯一"],
        ["handle", "VARCHAR", "用户昵称标识，如 @user1"],
        ["avatarUrl", "VARCHAR", "头像地址"],
        ["bio", "VARCHAR", "个人简介"],
        ["password", "VARCHAR", "登录密码"],
        ["role", "VARCHAR", "USER 或 ADMIN"],
        ["banned", "BOOLEAN", "是否封禁"],
        ["customWeights", "TEXT", "用户自定义推荐权重 JSON"],
        ["createdAt", "DATETIME", "创建时间"],
    ],
    "表 3-3 内容表设计": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "帖子 ID"],
        ["author", "BIGINT", "作者 ID"],
        ["parentContent", "BIGINT", "父评论 ID"],
        ["content", "TEXT", "正文内容"],
        ["imageUrl", "VARCHAR", "图片地址"],
        ["category", "VARCHAR", "内容分类"],
        ["viewCount", "INT", "浏览数"],
        ["likeCount", "INT", "点赞数"],
        ["commentCount", "INT", "评论数"],
        ["dislikeCount", "INT", "点踩数"],
        ["repostCount", "INT", "转发或引用数"],
        ["createdAt", "DATETIME", "发布时间"],
    ],
    "表 3-4 行为表设计": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "行为 ID"],
        ["userId", "BIGINT", "用户 ID"],
        ["contentId", "BIGINT", "内容 ID，搜索等行为可为空"],
        ["type", "VARCHAR", "VIEW、LIKE、DISLIKE、SKIP、SEARCH 等"],
        ["duration", "INT", "浏览停留时长"],
        ["createdAt", "DATETIME", "行为发生时间"],
    ],
    "表 3-5 关注表设计概要": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "主键"],
        ["followerId", "BIGINT", "关注者用户 ID，逻辑外键指向 users.id"],
        ["followeeId", "BIGINT", "被关注者用户 ID"],
        ["createdAt", "DATETIME", "关注时间，可用于推荐中新关注加权"],
    ],
    "表 3-6 标签实体与内容标签关联（逻辑结构）": [
        ["对象", "说明"],
        ["tags", "标签表：id、name（唯一）、可选统计字段。"],
        ["content_tags", "关联表：contentId、tagId，联合唯一约束防止重复绑定。"],
    ],
    "表 3-7 通知表设计概要": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "主键"],
        ["recipientId", "BIGINT", "接收者"],
        ["actorId", "BIGINT", "触发者"],
        ["type", "VARCHAR", "LIKE、COMMENT、FOLLOW、REPOST、QUOTE 等"],
        ["entityId", "BIGINT", "关联帖子或内容 ID"],
        ["isRead", "BOOLEAN", "是否已读"],
        ["createdAt", "DATETIME", "通知时间"],
    ],
    "表 3-8 广告表设计概要": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "主键"],
        ["title、description、imageUrl", "VARCHAR/TEXT", "展示文案与素材"],
        ["targetUrl", "VARCHAR", "点击跳转地址"],
        ["advertiser", "VARCHAR", "广告主标识"],
        ["targetTags、category", "VARCHAR", "定向标签与类别"],
        ["bidPrice", "DOUBLE", "出价，用于排序加权"],
        ["impressionCount、clickCount", "INT", "展示与点击计数"],
        ["active", "BOOLEAN", "是否启用投放"],
    ],
    "表 3-9 负反馈信号表（negative_signals）": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "主键，自增"],
        ["user_id", "BIGINT", "发起负反馈的用户"],
        ["target_type", "枚举", "CONTENT（单条内容）或 AUTHOR（作者）"],
        ["target_id", "BIGINT", "目标内容 ID 或用户 ID"],
        ["signal_type", "枚举", "NOT_INTERESTED、BLOCK、MUTE"],
        ["created_at", "DATETIME", "创建时间"],
    ],
    "表 3-10 私信消息表（tb_message）": [
        ["字段", "类型", "说明"],
        ["id", "BIGINT", "主键"],
        ["sender_id", "BIGINT", "发送方用户 ID"],
        ["recipient_id", "BIGINT", "接收方用户 ID"],
        ["content", "TEXT", "消息正文"],
        ["created_at", "DATETIME", "发送时间"],
        ["is_read", "BOOLEAN", "是否已读"],
    ],
}


descriptions = {
    "表 3-2 用户表设计": "用户表用于存储系统用户的基础信息、登录认证信息和角色权限信息，是内容发布、行为记录、关注关系、通知消息和后台管理等功能的数据基础。表中通过 role 字段区分普通用户和管理员，通过 banned 字段记录账号封禁状态，并保留 customWeights 字段用于保存用户自定义推荐权重配置。",
    "表 3-3 内容表设计": "内容表用于保存用户发布的帖子、评论、转发和引用等内容数据，是系统信息流展示和推荐排序的核心数据表。表中除正文、图片、分类和发布时间外，还保存浏览数、点赞数、评论数、点踩数和转发数等统计字段，以便在列表展示和推荐评分时快速读取，提高查询效率。",
    "表 3-4 行为表设计": "行为表用于记录用户在系统中的各类交互行为，包括浏览、点赞、点踩、快速滑过和搜索等。该表为用户画像构建、兴趣偏好分析和推荐算法排序提供基础行为数据，其中 duration 字段可用于记录浏览停留时长，createdAt 字段用于分析行为发生时间和兴趣变化趋势。",
    "表 3-5 关注表设计概要": "关注表用于维护用户之间的有向关注关系，表示某一用户关注了另一用户。该表不仅支撑关注列表和粉丝列表展示，也为推荐算法中的网内候选内容召回提供依据，使当前用户关注对象发布的内容能够在信息流排序中获得一定优先级。",
    "表 3-6 标签实体与内容标签关联（逻辑结构）": "标签实体与内容标签关联结构用于描述内容和标签之间的多对多关系。tags 表保存标签名称等基础信息，content_tags 关联表用于记录帖子与标签之间的绑定关系，从而支持话题检索、热门标签统计、用户兴趣画像构建和个性化推荐匹配。",
    "表 3-7 通知表设计概要": "通知表用于记录系统中由用户互动触发的提醒消息，例如点赞、评论、关注、转发和引用等事件。该表通过接收者、触发者、通知类型和关联实体等字段描述通知来源与目标，并通过 isRead 字段维护消息已读状态，以支持前端通知列表和未读提醒功能。",
    "表 3-8 广告表设计概要": "广告表用于保存信息流广告的投放素材、跳转地址、广告主、定向标签、广告类别和出价等信息。系统在进行广告分发时，会结合用户画像、广告标签匹配度、出价和历史展示点击数据计算广告得分，从而实现信息流广告的插入和排序展示。",
    "表 3-9 负反馈信号表（negative_signals）": "负反馈信号表用于记录用户对内容或作者产生的不感兴趣、屏蔽、静音等反向反馈。推荐系统在生成候选内容后，会根据该表过滤指定内容或指定作者发布的内容，减少用户明确不想看到的信息继续进入推荐列表，从而提升推荐结果的可接受性。",
    "表 3-10 私信消息表（tb_message）": "私信消息表用于保存用户之间的点对点聊天记录，包括发送方、接收方、消息正文、发送时间和已读状态等信息。该表与 WebSocket 通信机制配合使用，可以支持用户之间的实时消息发送、历史消息查询和未读消息提醒。",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_table(doc, caption, data):
    add_caption(doc, caption)
    tbl = doc.add_table(rows=0, cols=len(data[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    widths = [4.0, 3.2, 8.0] if len(data[0]) == 3 else [4.5, 10.5]
    for ri, row in enumerate(data):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cells[ci], widths[ci])
            if ri == 0:
                set_cell_shading(cells[ci], "F1F5F9")
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 or ci < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(val)
            set_run_font(r, 9.5, bold=(ri == 0))
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.6)

styles = doc.styles
styles["Normal"].font.name = "宋体"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(10.5)

p = doc.add_paragraph()
p.style = doc.styles["Heading 1"]
r = p.add_run("3.3 数据库设计")
set_run_font(r, 16, True)

add_para(doc, "数据库设计主要围绕用户、内容、行为、社交关系以及推荐辅助数据等方面展开。系统中的核心实体主要包括用户、帖子、行为、标签、关注、通知、私信、负反馈、广告和广告配置等。用户实体与内容实体之间存在一对多关系，一个用户可以发布多条内容；用户实体与行为实体之间也存在一对多关系，用户在浏览、点赞、评论、搜索等操作过程中，会产生多条行为记录；同时，用户之间还可以通过关注关系建立联系，即一个用户可以关注多个其他用户。")
add_para(doc, "内容实体是系统中较重要的数据对象，一条内容可以关联多个标签，也可以触发多条通知记录；在评论、转发和引用等场景下，内容还可以作为其他内容的目标对象。广告相关数据则主要根据广告标签、广告类别、出价和历史展示效果等信息，参与信息流中的广告插入过程，系统会结合用户画像和广告配置，判断广告内容是否适合展示给当前用户。")
add_para(doc, "为使数据库结构说明更清晰，本文在保留系统总体 E-R 图的基础上，将总体关系进一步拆分为用户与内容、行为画像与标签、社交关系与消息通知、广告投放四类局部 E-R 图。总体 E-R 图用于说明系统数据实体之间的整体联系，局部 E-R 图则分别服务于后续数据表说明，便于读者理解各表在业务流程中的作用。")

doc.add_picture(str(overall), width=Cm(15.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc, "图 3-3 系统 E-R 图")

p = doc.add_paragraph()
r = p.add_run("3.3.1 用户与内容数据表")
set_run_font(r, 12, True)
doc.add_picture(str(user_content), width=Cm(14.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc, "图 3-3（a）用户与内容局部 E-R 图")
add_para(doc, descriptions["表 3-2 用户表设计"])
add_table(doc, "表 3-2 用户表设计", tables["表 3-2 用户表设计"])
add_para(doc, descriptions["表 3-3 内容表设计"])
add_table(doc, "表 3-3 内容表设计", tables["表 3-3 内容表设计"])

p = doc.add_paragraph()
r = p.add_run("3.3.2 行为画像与标签数据表")
set_run_font(r, 12, True)
doc.add_picture(str(behavior_tag), width=Cm(14.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc, "图 3-3（b）行为画像与标签局部 E-R 图")
add_para(doc, descriptions["表 3-4 行为表设计"])
add_table(doc, "表 3-4 行为表设计", tables["表 3-4 行为表设计"])
add_para(doc, descriptions["表 3-6 标签实体与内容标签关联（逻辑结构）"])
add_table(doc, "表 3-5 标签实体与内容标签关联（逻辑结构）", tables["表 3-6 标签实体与内容标签关联（逻辑结构）"])
add_para(doc, descriptions["表 3-9 负反馈信号表（negative_signals）"])
add_table(doc, "表 3-6 负反馈信号表（negative_signals）", tables["表 3-9 负反馈信号表（negative_signals）"])

p = doc.add_paragraph()
r = p.add_run("3.3.3 社交关系与消息通知数据表")
set_run_font(r, 12, True)
doc.add_picture(str(social_message), width=Cm(14.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc, "图 3-3（c）社交关系与消息通知局部 E-R 图")
add_para(doc, descriptions["表 3-5 关注表设计概要"])
add_table(doc, "表 3-7 关注表设计概要", tables["表 3-5 关注表设计概要"])
add_para(doc, descriptions["表 3-7 通知表设计概要"])
add_table(doc, "表 3-8 通知表设计概要", tables["表 3-7 通知表设计概要"])
add_para(doc, descriptions["表 3-10 私信消息表（tb_message）"])
add_table(doc, "表 3-9 私信消息表（tb_message）", tables["表 3-10 私信消息表（tb_message）"])

p = doc.add_paragraph()
r = p.add_run("3.3.4 广告投放数据表")
set_run_font(r, 12, True)
doc.add_picture(str(ad_delivery), width=Cm(14.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc, "图 3-3（d）广告投放局部 E-R 图")
add_para(doc, descriptions["表 3-8 广告表设计概要"])
add_table(doc, "表 3-10 广告表设计概要", tables["表 3-8 广告表设计概要"])

add_para(doc, "在设计上，内容表中的浏览数、点赞数和评论数等字段属于必要的冗余字段。如果每次列表展示或推荐打分都通过关联查询实时计算，会明显增加查询压力。因此系统选择直接保存这些统计字段，并在点赞、评论、转发等操作发生时同步维护计数，以换取常用页面读取和推荐评分的效率。")
add_para(doc, "行为表中的 duration 字段用于记录用户在不同内容上的停留时长，可以为后续画像分析提供参考。对于搜索行为 SEARCH 来说，contentId 可以为空，因为搜索行为更多体现的是用户主动输入关键词的意图，不一定对应某一条具体帖子。负反馈相关实体则主要用于屏蔽作者、隐藏帖子以及减少不感兴趣内容的出现。推荐流程在合并候选内容后，会先对这些负反馈数据进行过滤，避免继续对无效内容进行打分，也能减少不必要的计算开销。")

out = Path("3.3数据库设计_修改版.docx")
doc.save(out)
print(out.resolve())
